# FILE: src/mint_python/core/table.py
# VERSION: 0.0.0
# START_MODULE_CONTRACT
#   PURPOSE: Table model + 5 factory constructors + post-construction shaping +
#     .render(parent_doc) -> python-docx Table for the Pure Python Edition.
#     Phase-7 cells are str-based; richer Paragraph cells are deferred to a
#     Phase-2 wave (MP-CONTENT lands in parallel with this wave; do NOT couple
#     here to avoid the parallel race).
#   SCOPE: Public surface = Table, Cell, factory classmethods (from_list,
#     from_markdown, from_list_of_dicts, financial, comparison), shapers
#     (apply_style, set_column_widths, autofit), .render(), and the three
#     public exception types (TableRaggedRowsError, TableMarkdownParseError,
#     TableInvalidDictKeysError).
#   DEPENDS: MP-STYLE (Style only — read-only import), python-docx (Inches,
#     Pt as DocxPt, RGBColor, WD_ALIGN_PARAGRAPH). Stdlib: dataclasses,
#     logging, re. NOT depending on mint_python.core.content (parallel
#     sibling worker — race-collide).
#   LINKS: docs/development-plan.xml#MP-TABLE,
#     docs/verification-plan.xml#V-MP-TABLE,
#     docs/knowledge-graph.xml#MP-TABLE
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   Cell                          - frozen @dataclass; str value + Style/align/colspan/rowspan
#   Table                         - @dataclass; rows + shaping state + .render()
#   Table.from_list               - rectangular list-of-lists; optional header
#   Table.from_markdown           - pipe-table parser with explicit separator-line
#   Table.from_list_of_dicts      - strict same-key-set requirement
#   Table.financial               - right-aligned numerics + thousands separator
#   Table.comparison              - two-column side-by-side preset
#   Table.apply_style             - default cell style fluent setter
#   Table.set_column_widths       - per-column inch widths fluent setter
#   Table.autofit                 - python-docx autofit flag fluent setter
#   Table.render                  - emits BLOCK_RENDER_TABLE; returns docx.Table
#   TableError                    - base error
#   TableRaggedRowsError          - TABLE_RAGGED_ROWS
#   TableMarkdownParseError       - TABLE_MARKDOWN_PARSE_ERROR
#   TableInvalidDictKeysError     - TABLE_INVALID_DICT_KEYS
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: Wave-7-2 (MP-TABLE): initial implementation per V-MP-TABLE
#     scenarios 1-9. Cells are str-based; Phase-2 wave can refactor to
#     Paragraph once MP-CONTENT is available without racing this worker.
# END_CHANGE_SUMMARY

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from typing import Any, ClassVar, cast

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor
from docx.shared import Pt as DocxPt
from docx.table import Table as DocxTable

from mint_python.core.style import Style

logger = logging.getLogger("mint_python.core.table")


# ---------------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------------


class TableError(Exception):
    """Base for MP-TABLE errors."""


class TableRaggedRowsError(TableError):
    """Raised when from_list rows have differing column counts (TABLE_RAGGED_ROWS)."""


class TableMarkdownParseError(TableError):
    """Raised on malformed pipe-table syntax in from_markdown (TABLE_MARKDOWN_PARSE_ERROR)."""


class TableInvalidDictKeysError(TableError):
    """Raised when from_list_of_dicts rows have differing key sets (TABLE_INVALID_DICT_KEYS)."""


# ---------------------------------------------------------------------------
# Cell
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class Cell:
    """Frozen cell carrier.

    Phase-7 surface uses str values exclusively; richer types (Paragraph and
    friends) come in a Phase-2 wave once MP-CONTENT is settled. Non-str
    values are coerced via ``str(value)`` at construction time so callers
    can pass numbers in lists without bookkeeping.
    """

    value: str
    style: Style | None = None
    align: str | None = None  # "left" | "center" | "right" | None=inherit
    colspan: int = 1
    rowspan: int = 1

    def __post_init__(self) -> None:
        # Coerce non-str values to str. Use object.__setattr__ because the
        # dataclass is frozen.
        if not isinstance(self.value, str):
            object.__setattr__(self, "value", str(self.value))


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


_NUMERIC_RE = re.compile(r"^-?\d+(\.\d+)?$")
_VALID_ALIGN: frozenset[str] = frozenset({"left", "center", "right"})
_ALIGN_MAP: dict[str, int] = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _promote_cell(value: Any) -> Cell:
    """Promote a raw cell input into a Cell instance.

    Pass-through when already a Cell; otherwise wrap with str() coercion.
    """
    if isinstance(value, Cell):
        return value
    return Cell(value=str(value))


def _is_numeric_like(value: Any) -> bool:
    """Return True for int/float (excluding bool) and numeric-looking strings."""
    if isinstance(value, bool):
        return False
    if isinstance(value, (int, float)):
        return True
    if isinstance(value, str):
        return bool(_NUMERIC_RE.match(value.strip()))
    return False


def _format_numeric(value: Any) -> str:
    """Format a numeric-like value with a thousands separator.

    int-like -> ``f"{n:,.0f}"``; float-like -> ``f"{n:,.2f}"``. Strings are
    parsed first; the source representation determines int-vs-float.
    """
    if isinstance(value, int) and not isinstance(value, bool):
        return f"{value:,.0f}"
    if isinstance(value, float):
        return f"{value:,.2f}"
    if isinstance(value, str):
        s = value.strip()
        if "." in s:
            return f"{float(s):,.2f}"
        return f"{int(s):,.0f}"
    # Should not be reached when guarded by _is_numeric_like.
    return str(value)


# ---------------------------------------------------------------------------
# Table
# ---------------------------------------------------------------------------


@dataclass
class Table:
    """Table model with §3.3 factories + .render(parent_doc) -> docx.Table."""

    _rows: list[list[Cell]]
    _has_header: bool = True
    _column_widths: list[float] | None = None  # inches
    _autofit: bool = False
    _row_styles: dict[int, Style] = field(default_factory=dict)
    _column_styles: dict[int, Style] = field(default_factory=dict)
    _table_style: Style | None = None  # default cell style

    # Surface module-level errors as class attributes for convenient access.
    RaggedRowsError: ClassVar[type[TableRaggedRowsError]] = TableRaggedRowsError
    MarkdownParseError: ClassVar[type[TableMarkdownParseError]] = TableMarkdownParseError
    InvalidDictKeysError: ClassVar[type[TableInvalidDictKeysError]] = TableInvalidDictKeysError

    # ------------------------------------------------------------------ #
    # Factory constructors
    # ------------------------------------------------------------------ #

    @classmethod
    def from_list(cls, rows: list[list[Any]], header: bool = True) -> Table:
        """Build from a list-of-lists. First row is treated as header by default.

        Validates rectangular shape: every row must have the same column
        count as row 0. Raises :class:`TableRaggedRowsError` naming the
        first divergent row index when violated.
        """
        if not rows:
            # Empty table is permitted; downstream render emits a zero-row,
            # zero-col docx table. _has_header is forced False so the
            # render-time header treatment is a no-op.
            return cls(_rows=[], _has_header=False)
        expected_cols = len(rows[0])
        promoted: list[list[Cell]] = []
        for idx, row in enumerate(rows):
            if len(row) != expected_cols:
                raise TableRaggedRowsError(
                    f"row {idx} has {len(row)} columns, expected {expected_cols} "
                    f"(matching row 0)"
                )
            promoted.append([_promote_cell(v) for v in row])
        return cls(_rows=promoted, _has_header=header)

    @classmethod
    def from_markdown(cls, text: str) -> Table:
        """Parse a pipe-table block.

        Format: first non-blank line is the header row "| a | b |"; the
        IMMEDIATELY following non-blank line MUST be the separator row
        "|---|---|" with at least three dashes per column. Subsequent
        non-blank lines are data rows. Outer pipes are trimmed; cells are
        split on "|" and stripped.

        Raises :class:`TableMarkdownParseError` on malformed input,
        embedding the offending line index (1-based) and a short reason.
        """
        if not isinstance(text, str):
            raise TableMarkdownParseError(
                f"expected str input, got {type(text).__name__}"
            )

        # Collect (1-based original line index, stripped content) for non-blank lines.
        lines: list[tuple[int, str]] = []
        for i, raw in enumerate(text.splitlines(), start=1):
            stripped = raw.strip()
            if stripped:
                lines.append((i, stripped))

        if len(lines) < 2:
            raise TableMarkdownParseError(
                "markdown table requires at least a header line and a separator line "
                f"(got {len(lines)} non-blank line(s))"
            )

        def _split_pipe(line_no: int, line: str) -> list[str]:
            if "|" not in line:
                raise TableMarkdownParseError(
                    f"line {line_no}: expected pipe-delimited cells, got {line!r}"
                )
            inner = line
            if inner.startswith("|"):
                inner = inner[1:]
            if inner.endswith("|"):
                inner = inner[:-1]
            return [c.strip() for c in inner.split("|")]

        header_line_no, header_line = lines[0]
        sep_line_no, sep_line = lines[1]

        header_cells = _split_pipe(header_line_no, header_line)

        sep_cells = _split_pipe(sep_line_no, sep_line)
        if len(sep_cells) != len(header_cells):
            raise TableMarkdownParseError(
                f"line {sep_line_no}: separator has {len(sep_cells)} cells, "
                f"header (line {header_line_no}) has {len(header_cells)}"
            )
        sep_re = re.compile(r"^:?-{3,}:?$")
        for cell in sep_cells:
            if not sep_re.match(cell):
                raise TableMarkdownParseError(
                    f"line {sep_line_no}: separator cell {cell!r} is not '---' (>=3 dashes)"
                )

        rows_raw: list[list[Any]] = [list(header_cells)]
        for line_no, line in lines[2:]:
            data_cells = _split_pipe(line_no, line)
            if len(data_cells) != len(header_cells):
                raise TableMarkdownParseError(
                    f"line {line_no}: data row has {len(data_cells)} cells, "
                    f"expected {len(header_cells)} (matching header)"
                )
            rows_raw.append(list(data_cells))

        return cls.from_list(rows_raw, header=True)

    @classmethod
    def from_list_of_dicts(cls, rows: list[dict[str, Any]]) -> Table:
        """Build from a list of dicts that share an identical key set.

        Column order is the insertion order of the FIRST dict. Strict
        equality is required (===, not subset). Raises
        :class:`TableInvalidDictKeysError` naming the divergent keys
        (set diff between row 0 and the first divergent row).
        """
        if not rows:
            return cls(_rows=[], _has_header=False)
        if not isinstance(rows[0], dict):
            raise TableInvalidDictKeysError(
                f"row 0 is not a dict: got {type(rows[0]).__name__}"
            )
        # Insertion order of the first dict drives column order.
        columns: list[str] = list(rows[0].keys())
        column_set: set[str] = set(columns)
        for idx, r in enumerate(rows[1:], start=1):
            if not isinstance(r, dict):
                raise TableInvalidDictKeysError(
                    f"row {idx} is not a dict: got {type(r).__name__}"
                )
            keys = set(r.keys())
            if keys != column_set:
                missing = sorted(column_set - keys)
                extra = sorted(keys - column_set)
                raise TableInvalidDictKeysError(
                    f"row {idx} key set diverges from row 0: "
                    f"missing={missing}, extra={extra}"
                )
        # Build rectangular rows.
        out_rows: list[list[Any]] = [list(columns)]
        for r in rows:
            out_rows.append([r[k] for k in columns])
        return cls.from_list(out_rows, header=True)

    @classmethod
    def financial(cls, rows: list[list[Any]]) -> Table:
        """Numeric cells right-aligned with thousands separator.

        Builds via :meth:`from_list` (rectangular validation), then for
        every numeric cell (``int``, ``float``, or ``str`` matching
        ``r"^-?\\d+(\\.\\d+)?$"``):
          * apply right alignment (``Cell.align="right"``)
          * reformat with thousands separator (int-like ``f"{n:,.0f}"``,
            float-like ``f"{n:,.2f}"``)
        Non-numeric cells keep their raw value AND default alignment (no
        silent coercion). The header row, if present, is left untouched.
        """
        # Build first via from_list to validate shape; then reshape numerics.
        if not rows:
            return cls.from_list(rows)
        # We need access to the original (un-promoted) values to detect
        # numeric ints vs strings; promote-and-walk loses int/float type.
        # So we re-run validation here without promotion.
        expected_cols = len(rows[0])
        for idx, r in enumerate(rows):
            if len(r) != expected_cols:
                raise TableRaggedRowsError(
                    f"row {idx} has {len(r)} columns, expected {expected_cols} "
                    f"(matching row 0)"
                )

        promoted: list[list[Cell]] = []
        for idx, row in enumerate(rows):
            cells: list[Cell] = []
            for raw in row:
                # Header row (idx 0) is left as-is per spec — no numeric
                # coercion of header labels.
                if idx == 0:
                    cells.append(_promote_cell(raw))
                    continue
                if isinstance(raw, Cell):
                    # Inspect Cell.value (already str-coerced) for numeric shape.
                    if _is_numeric_like(raw.value):
                        cells.append(
                            Cell(
                                value=_format_numeric(raw.value),
                                style=raw.style,
                                align="right",
                                colspan=raw.colspan,
                                rowspan=raw.rowspan,
                            )
                        )
                    else:
                        cells.append(raw)
                    continue
                if _is_numeric_like(raw):
                    cells.append(Cell(value=_format_numeric(raw), align="right"))
                else:
                    cells.append(_promote_cell(raw))
            promoted.append(cells)
        return cls(_rows=promoted, _has_header=True)

    @classmethod
    def comparison(
        cls, left_label: str, right_label: str, rows: list[list[Any]]
    ) -> Table:
        """Two-column side-by-side preset with header [left_label, right_label]."""
        for idx, r in enumerate(rows):
            if len(r) != 2:
                raise TableRaggedRowsError(
                    f"comparison row {idx} has {len(r)} columns, expected 2"
                )
        full = [[left_label, right_label]] + [list(r) for r in rows]
        return cls.from_list(full, header=True)

    # ------------------------------------------------------------------ #
    # Fluent shapers
    # ------------------------------------------------------------------ #

    def apply_style(self, style: Style) -> Table:
        """Set the default cell style; returns self for chaining."""
        self._table_style = style
        return self

    def set_column_widths(self, widths: list[float]) -> Table:
        """Set column widths in inches; len must equal column count."""
        col_count = len(self._rows[0]) if self._rows else 0
        if len(widths) != col_count:
            raise ValueError(
                f"column_widths length {len(widths)} != column count {col_count}"
            )
        self._column_widths = list(widths)
        return self

    def autofit(self) -> Table:
        """Enable python-docx autofit; returns self for chaining."""
        self._autofit = True
        return self

    # ------------------------------------------------------------------ #
    # Render
    # ------------------------------------------------------------------ #

    # START_BLOCK_RENDER_TABLE
    def render(self, parent_doc: DocxDocument) -> DocxTable:
        """Render the Table into a python-docx Table on ``parent_doc``.

        Emits ``[MP-Table][render][BLOCK_RENDER_TABLE] rows=N cols=M`` at INFO
        BEFORE the python-docx mutation. Applies cell -> row -> column ->
        table style precedence. Returns the python-docx Table for downstream
        chaining.
        """
        row_count = len(self._rows)
        col_count = len(self._rows[0]) if self._rows else 0

        logger.info(
            "[MP-Table][render][BLOCK_RENDER_TABLE] rows=%d cols=%d",
            row_count,
            col_count,
        )

        if row_count == 0 or col_count == 0:
            # python-docx requires positive rows/cols for add_table; fall
            # back to a 1x1 placeholder so callers that render an empty
            # Table get a tangible (if vacuous) result.
            docx_table = cast(DocxTable, parent_doc.add_table(rows=1, cols=1))
            docx_table.autofit = self._autofit
            return docx_table

        docx_table = cast(DocxTable, parent_doc.add_table(rows=row_count, cols=col_count))
        docx_table.autofit = self._autofit

        # Apply column widths if requested.
        if self._column_widths is not None:
            for i, w in enumerate(self._column_widths):
                docx_table.columns[i].width = Inches(w)
                # Also set cell-level width — Word respects per-cell width
                # values when columns are auto-fit.
                for row in docx_table.rows:
                    row.cells[i].width = Inches(w)

        for r_idx, row in enumerate(self._rows):
            docx_row = docx_table.rows[r_idx]
            for c_idx, cell in enumerate(row):
                docx_cell = docx_row.cells[c_idx]
                # Resolve effective style by precedence: cell > row > column > table.
                effective_style: Style | None = (
                    cell.style
                    or self._row_styles.get(r_idx)
                    or self._column_styles.get(c_idx)
                    or self._table_style
                )
                # Resolve effective alignment: explicit cell.align wins;
                # otherwise fall back to the effective style's alignment
                # (when valid for paragraph alignment).
                effective_align: str | None = cell.align
                if (
                    effective_align is None
                    and effective_style is not None
                    and effective_style.alignment in _VALID_ALIGN
                ):
                    effective_align = effective_style.alignment

                # python-docx creates an empty paragraph in each new cell;
                # write our text via that paragraph (rather than .text=) so
                # we can attach run-level font formatting.
                paragraph = docx_cell.paragraphs[0]
                # Wipe any existing runs in case python-docx pre-seeds.
                for existing_run in list(paragraph.runs):
                    existing_run.text = ""
                run = paragraph.add_run(cell.value)

                if effective_style is not None:
                    run.font.name = effective_style.font
                    run.font.size = DocxPt(effective_style.size_pt)
                    run.font.bold = effective_style.bold
                    run.font.italic = effective_style.italic
                    # color_hex is post-load guaranteed literal "#RRGGBB".
                    hex_value = effective_style.color_hex.lstrip("#")
                    run.font.color.rgb = RGBColor.from_string(hex_value)

                if effective_align is not None and effective_align in _ALIGN_MAP:
                    paragraph.paragraph_format.alignment = _ALIGN_MAP[effective_align]

                # Header treatment: bold the first row when _has_header.
                # Cell-level explicit style wins (only set bold when no
                # explicit style override is in play and no run-level bold
                # was set by the effective style).
                if self._has_header and r_idx == 0 and effective_style is None:
                    run.font.bold = True

        return docx_table

    # END_BLOCK_RENDER_TABLE


__all__ = [
    "Cell",
    "Table",
    "TableError",
    "TableInvalidDictKeysError",
    "TableMarkdownParseError",
    "TableRaggedRowsError",
]
