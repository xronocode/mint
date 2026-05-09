# FILE: src/mint_python/core/document.py
# VERSION: 1.3.0
# START_MODULE_CONTRACT
#   PURPOSE: Document facade per handover §3.1 — full API surface (cover,
#     sections, TOC, header/footer, save, validate, fix, to_pdf) for the Pure
#     Python Edition. Save uses python-docx as the backbone with a small
#     lxml drop-down for the Word TOC field. validate() and fix() delegate
#     to MP-VALIDATE and MP-FIX via temp-file save. inject_grace delegates to
#     MP-GRACE via temp-file save. to_pdf delegates to Gotenberg HTTP API.
#   SCOPE: Public surface = Document (@dataclass facade) + 5 errors
#     (DocumentError, DocumentFormatUnsupportedError, DocumentPresetNotFoundError,
#     DocumentSaveIOError, GotenbergError) + PhaseGuardNotImplementedError.
#     Sibling-only deps: MP-STYLE (load_preset; STYLE_PRESET_NOT_FOUND),
#     MP-SECTION (Section), MP-VALIDATE (run_checks, SeverityMode,
#     ValidationReport), MP-FIX (fix as mp_fix, FixReport). save() emits exactly
#     one BLOCK_SAVE_DOCX before serialization; pins core.xml
#     dcterms:created/modified to DOCUMENT_FIXED_TIMESTAMP so two saves of the
#     same Document produce the same fingerprint hash.
#   DEPENDS: python-docx (1.1.x), lxml (5.x), httpx, mint_python.core.style,
#     mint_python.core.section, mint_python.validate, mint_python.fix,
#     mint_python.grace (lazy via inject_grace).
#     NO sibling import of MP-CONTENT/MP-TABLE — we reach those types only
#     transitively via Section.render.
#   LINKS: docs/development-plan.xml#MP-DOCUMENT,
#     docs/verification-plan.xml#V-MP-DOCUMENT,
#     docs/knowledge-graph.xml#MP-DOCUMENT
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   Document                            - @dataclass facade
#   Document.with_style_preset          - registry-name preset loader
#   Document.with_style_from            - JSON-path preset loader
#   Document.add_cover                  - cover-page builder
#   Document.add_section                - append a MP-SECTION
#   Document.add_toc                    - Word TOC field (lxml drop-down)
#   Document.set_header                 - document-wide header text
#   Document.set_footer                 - document-wide footer text
#   Document.save                       - serialize to .docx; emits BLOCK_SAVE_DOCX
#   Document.inject_grace               - delegate to MP-GRACE.bootstrap via temp-file save
#   Document.validate                   - Phase-9: validates via MP-VALIDATE
#   Document.fix                        - Phase-9: auto-fixes via MP-FIX
#   Document.to_pdf                     - Gotenberg PDF render; emits BLOCK_RENDER_PDF
#   _resolve_severity_mode - str → SeverityMode helper
#   DocumentError                       - base error
#   DocumentFormatUnsupportedError      - format != 'docx'
#   DocumentPresetNotFoundError         - wraps STYLE_PRESET_NOT_FOUND
#   DocumentSaveIOError                 - save() OSError; partial output unlinked
#   GotenbergError                      - Gotenberg HTTP failure (non-200 or unreachable)
#   PhaseGuardNotImplementedError       - Phase-N stub (NotImplementedError subclass)
#   DOCUMENT_FIXED_TIMESTAMP            - 1980-01-01 datetime; idempotency anchor
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: Wave-11-2 — unstub to_pdf via Gotenberg HTTP integration
#   PRIOR: Wave-11-1 — unstub inject_grace via temp-file delegation to MP-GRACE
#   PRIOR: Wave-9-4 — unstub Document.validate and Document.fix via
#     temp-file delegation to MP-VALIDATE + MP-FIX.
#   PRIOR: Wave-7-4 (MP-Document): initial implementation per V-MP-DOCUMENT
#     scenarios 1-9 + BLOCK_SAVE_DOCX + BLOCK_PHASE_GUARD trace assertions.
# END_CHANGE_SUMMARY

from __future__ import annotations

import contextlib
import logging
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from types import SimpleNamespace
from typing import Any, Literal

from docx import Document as DocxDocumentFactory
from docx.document import Document as DocxDocument
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from lxml import etree

from mint_python.core.section import Section
from mint_python.core.style import STYLE_PRESET_NOT_FOUND, load_preset
from mint_python.fix import FixReport
from mint_python.fix import fix as mp_fix
from mint_python.validate import SeverityMode, ValidationReport, run_checks

logger = logging.getLogger("mint_python.core.document")

# Idempotency anchor: pinning core.xml dcterms:created/modified to a fixed
# datetime makes two save() calls produce byte-stable style XML and therefore
# a stable mint.fingerprint hash. 1980-01-01 is the ZIP-format epoch; chosen
# specifically so timestamp leakage breaking idempotency is impossible.
DOCUMENT_FIXED_TIMESTAMP: datetime = datetime(1980, 1, 1, 0, 0, 0)


# ---------------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------------


class DocumentError(Exception):
    """Base for MP-DOCUMENT errors."""


class DocumentFormatUnsupportedError(DocumentError):
    """Raised when format is not 'docx' (DOCUMENT_FORMAT_UNSUPPORTED).

    Phase-7 ships only the .docx backbone; PPTX/XLSX are deferred to later
    phases per handover §6.
    """


class DocumentPresetNotFoundError(DocumentError):
    """Raised when ``with_style_preset(name)`` lookup fails.

    Wraps :class:`mint_python.core.style.STYLE_PRESET_NOT_FOUND` so the
    MP-DOCUMENT public API stays decoupled from MP-STYLE error symbols.
    """


class DocumentSaveIOError(DocumentError):
    """Raised when ``save()`` encounters an I/O failure.

    Partial output (the half-written .docx) is unlinked before re-raising so
    callers can retry without manual cleanup.
    """


class GotenbergError(DocumentError):
    """Raised when Gotenberg returns non-200 or is unreachable."""


class PhaseGuardNotImplementedError(NotImplementedError):
    """Phase-N stub touched on a not-yet-implemented method.

    Subclass of :class:`NotImplementedError` so callers using
    ``except NotImplementedError`` keep working; carries the target phase
    + delegation target in the message for diagnosability.
    """


# ---------------------------------------------------------------------------
# Helper
# ---------------------------------------------------------------------------


def _resolve_severity_mode(level: str) -> SeverityMode:
    """Resolve a human-readable level to SeverityMode."""
    if level in ("lenient",):
        return SeverityMode.LENIENT
    if level in ("strict",):
        return SeverityMode.STRICT
    return SeverityMode.AUDIT


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------


@dataclass
class Document:
    """Top-level document facade — fluent surface mirroring handover §3.1.

    Construction::

        Document(format='docx', title='Annual Report') \
            .with_style_preset('alga_corporate') \
            .add_cover(title='Annual Report', subtitle='2026') \
            .add_toc(max_level=3) \
            .add_section(Section('Overview', level=1).add_paragraph('...')) \
            .set_header('Confidential') \
            .set_footer('Page') \
            .save('out.docx')

    All fluent setters return ``self`` so callers compose the whole
    document in a single expression. ``save(path)`` is the only method
    that performs disk I/O.

    Idempotency contract: two calls to ``save(p_a)`` and ``save(p_b)``
    on the same Document produce equal :func:`mint.fingerprint.fingerprint`
    hashes. python-docx populates ``core.xml`` ``dcterms:created`` /
    ``dcterms:modified`` with ``datetime.now()`` during save; without
    pinning, every save produces a different fingerprint. We pin both
    fields to :data:`DOCUMENT_FIXED_TIMESTAMP` (1980-01-01) at the start
    of ``save()`` to guarantee stability. See V-MP-DOCUMENT scenario-8.

    Stub (``to_pdf``) emits ``[MP-Document][stub][BLOCK_PHASE_GUARD]`` at
    INFO BEFORE raising :class:`PhaseGuardNotImplementedError` so plan
    consumers that accidentally schedule a deferred capability in Phase-7
    fail loud + observable.
    """

    format: Literal["docx"]
    title: str
    _preset: SimpleNamespace | None = None
    _preset_name: str | None = None
    _cover: dict[str, Any] | None = None
    _sections: list[Section] = field(default_factory=list)
    _toc: dict[str, Any] | None = None
    _header: str | None = None
    _footer: str | None = None

    def __post_init__(self) -> None:
        if self.format != "docx":
            raise DocumentFormatUnsupportedError(
                f"format must be 'docx' in Phase-7, got {self.format!r}. "
                f"Valid formats: {{'docx'}}. PPTX/XLSX deferred to later phases "
                f"per handover §6."
            )

    # ------------------------------------------------------------------ #
    # Fluent style binding
    # ------------------------------------------------------------------ #

    def with_style_preset(self, name: str) -> Document:
        """Bind a registry preset (e.g. ``'alga_corporate'``); returns self.

        Wraps :class:`STYLE_PRESET_NOT_FOUND` as
        :class:`DocumentPresetNotFoundError` so MP-DOCUMENT callers don't
        leak MP-STYLE error symbols.
        """
        try:
            self._preset = load_preset(name)
        except STYLE_PRESET_NOT_FOUND as exc:
            raise DocumentPresetNotFoundError(str(exc)) from exc
        self._preset_name = name
        return self

    def with_style_from(self, path: Path | str) -> Document:
        """Bind a preset from a JSON file path; returns self.

        ``_preset_name`` is set to ``Path(path).stem`` for trace-payload
        consistency. The preset's declared ``"name"`` field already drives
        :func:`load_preset`'s own ``BLOCK_LOAD_PRESET`` payload, so the
        stem fallback here is only for our local debugging hint.
        """
        self._preset = load_preset(path=Path(path))
        self._preset_name = Path(path).stem
        return self

    # ------------------------------------------------------------------ #
    # Fluent content builders
    # ------------------------------------------------------------------ #

    def add_cover(
        self,
        title: str,
        subtitle: str | None = None,
        logo: Path | str | None = None,
    ) -> Document:
        """Set cover-page metadata; rendered at save() time."""
        self._cover = {
            "title": title,
            "subtitle": subtitle,
            "logo": Path(logo) if logo else None,
        }
        return self

    def add_section(self, section: Section) -> Document:
        """Append a Section; returns self for fluent chaining."""
        self._sections.append(section)
        return self

    def add_toc(self, max_level: int = 3) -> Document:
        """Insert a Word TOC field (auto-update on first open).

        Rendered at save() time via the lxml drop-down — python-docx has no
        first-class TOC API. ``max_level`` controls the ``\\o "1-N"`` switch
        in the field instruction, capping which heading levels appear.
        """
        self._toc = {"max_level": max_level}
        return self

    def set_header(self, text: str) -> Document:
        """Set document-wide running header text."""
        self._header = text
        return self

    def set_footer(self, text: str) -> Document:
        """Set document-wide running footer text."""
        self._footer = text
        return self

    # ------------------------------------------------------------------ #
    # save()
    # ------------------------------------------------------------------ #

    def save(self, path: Path | str) -> Path:
        """Serialize the document to ``path`` (.docx).

        Trace contract: emits exactly one
        ``[MP-Document][save][BLOCK_SAVE_DOCX] section_count=N table_count=M
        output_path=...`` INFO record BEFORE the python-docx ``save()`` call.
        Section/table counts are computed from the in-memory model so the
        marker payload is meaningful even if the actual write fails.

        Idempotency: byte-equality of the resulting ZIP is NOT guaranteed
        (python-docx attribute order varies across runs), but the
        :func:`mint.fingerprint.fingerprint` hash MUST match across calls.
        We pin ``core.xml`` ``dcterms:created`` and ``dcterms:modified`` to
        :data:`DOCUMENT_FIXED_TIMESTAMP` to prevent timestamp leakage from
        breaking that fingerprint. See V-MP-DOCUMENT forbidden-5 + scenario-8.

        On :class:`OSError`, any partial output at ``path`` is unlinked
        before re-raising as :class:`DocumentSaveIOError` so callers can
        retry without manual cleanup.

        Returns the resolved output ``Path`` for convenient chaining.
        """
        out = Path(path)

        # Cover counts as a "section" in the trace payload sense per the
        # MP-DOCUMENT contract; we surface "intended structural blocks" so
        # consumers can sanity-check before opening the file.
        section_count = len(self._sections) + (1 if self._cover else 0)
        table_count = sum(
            1
            for s in self._sections
            for b in s._blocks
            if hasattr(b, "_rows")  # duck-type detect MP-TABLE.Table
        )

        # START_BLOCK_SAVE_DOCX
        logger.info(
            "[MP-Document][save][BLOCK_SAVE_DOCX] "
            "section_count=%d table_count=%d output_path=%s",
            section_count,
            table_count,
            str(out),
        )
        # END_BLOCK_SAVE_DOCX

        try:
            doc = DocxDocumentFactory()
            self._render_cover(doc)
            self._render_toc(doc)
            for section in self._sections:
                section.render(doc)
            self._render_header_footer(doc)
            self._pin_core_xml_timestamps(doc)
            doc.save(str(out))
        except OSError as exc:
            # Best-effort cleanup; if the unlink itself fails we still
            # re-raise the original I/O error below.
            if out.exists():
                with contextlib.suppress(OSError):
                    out.unlink()
            raise DocumentSaveIOError(f"save failed: {exc}") from exc

        return out

    # ------------------------------------------------------------------ #
    # Phase-N stubs
    # ------------------------------------------------------------------ #

    def inject_grace(self, *args: Any, **kwargs: Any) -> Any:
        """Inject GRACE manifest + instructions via MP-GRACE.

        Delegation pattern: saves to temp, calls MP-GRACE.bootstrap.
        """
        import tempfile

        from mint_python.grace import bootstrap as grace_bootstrap

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
            tmp_path = Path(tf.name)
        try:
            self.save(tmp_path)
            rules = kwargs.get("rules")
            output_path = kwargs.get("output_path")
            return grace_bootstrap(tmp_path, rules=rules, output_path=output_path)
        finally:
            tmp_path.unlink(missing_ok=True)

    def validate(self, level: str = "lenient") -> ValidationReport:
        """Validate the document via MP-VALIDATE (pure Python).

        Delegation pattern: saves to a temp .docx, then calls
        MP-VALIDATE.run_checks on the saved file. The temp file is
        cleaned up after the call returns.
        """
        import tempfile
        from pathlib import Path

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
            tmp_path = Path(tf.name)
        try:
            self.save(tmp_path)
            severity = _resolve_severity_mode(level)
            return run_checks(tmp_path, severity)
        finally:
            tmp_path.unlink(missing_ok=True)

    def fix(self, strategy: str = "safe_first") -> FixReport:
        """Auto-fix the document via MP-FIX (pure Python).

        Delegation pattern: saves to a temp .docx, then calls
        MP-FIX.fix on the saved file. The temp file is cleaned up.
        """
        import tempfile
        from pathlib import Path

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
            tmp_path = Path(tf.name)
        try:
            self.save(tmp_path)
            return mp_fix(tmp_path)
        finally:
            tmp_path.unlink(missing_ok=True)

    def to_pdf(
        self,
        output_path: str | Path | None = None,
        host: str = "localhost",
        port: int = 3002,
    ) -> Path:
        """Convert saved document to PDF via Gotenberg.

        Saves the current document to a temp .docx, sends it to Gotenberg,
        and writes the resulting PDF to output_path. If output_path is None,
        uses a temp path under /tmp.
        """
        import tempfile

        import httpx

        output = (
            Path(output_path)
            if output_path
            else Path(f"/tmp/mint_pdf_output_{uuid.uuid4().hex[:8]}.pdf")
        )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
            tmp_docx = Path(tf.name)

        try:
            self.save(tmp_docx)
            url = f"http://{host}:{port}/forms/libreoffice/convert"

            with open(tmp_docx, "rb") as f:
                files = {
                    "files": (
                        tmp_docx.name,
                        f,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                }
                response = httpx.post(url, files=files, timeout=60.0)

            if response.status_code != 200:
                raise GotenbergError(
                    f"Gotenberg returned {response.status_code}: {response.text[:200]}"
                )

            # START_BLOCK_RENDER_PDF
            output.write_bytes(response.content)
            logger.info(
                "[MP-Document][to_pdf][BLOCK_RENDER_PDF] "
                "Rendered PDF: output=%s size=%d bytes",
                output,
                len(response.content),
            )
            # END_BLOCK_RENDER_PDF
            return output
        finally:
            tmp_docx.unlink(missing_ok=True)

    # ------------------------------------------------------------------ #
    # Internal rendering helpers
    # ------------------------------------------------------------------ #

    def _render_cover(self, doc: DocxDocument) -> None:
        """Render the cover page using python-docx primitives.

        Implementation note: we use python-docx primitives end-to-end here
        — ``add_heading`` for the title, ``add_paragraph`` for the
        subtitle, ``add_picture`` for the optional logo, and
        ``WD_BREAK.PAGE`` for the page break. We deliberately do NOT drop
        down to lxml for the cover: python-docx covers everything we need
        and the lxml route would be strictly more brittle.
        """
        if self._cover is None:
            return

        # Title — rendered as Heading 0 (Title) so Word recognizes it as a
        # cover-style heading. python-docx maps level=0 to the built-in
        # "Title" style which is what Word expects on a cover page.
        title_text: str = self._cover["title"]
        doc.add_heading(title_text, level=0)

        # Optional subtitle — plain paragraph below the title.
        subtitle: str | None = self._cover["subtitle"]
        if subtitle:
            doc.add_paragraph(subtitle)

        # Optional logo image.
        logo: Path | None = self._cover["logo"]
        if logo is not None:
            doc.add_picture(str(logo))

        # Page break separating cover from body content.
        break_para = doc.add_paragraph()
        break_run = break_para.add_run()
        break_run.add_break(WD_BREAK.PAGE)

    def _render_toc(self, doc: DocxDocument) -> None:
        """Inject a Word TOC field via lxml drop-down.

        python-docx has no first-class TOC support — the closest you can
        get is constructing the OOXML field nodes by hand. The injected
        structure is::

            <w:p>
              <w:r>
                <w:fldChar w:fldCharType="begin"/>
                <w:instrText xml:space="preserve"> TOC \\o "1-N" \\h \\z \\u </w:instrText>
                <w:fldChar w:fldCharType="end"/>
              </w:r>
            </w:p>

        Word resolves the field on first open (or when the user presses
        F9), populating the actual entries. The ``\\h`` switch makes
        entries hyperlinks; ``\\z`` hides tab leaders in web layout;
        ``\\u`` uses the document's outline level when no TOC entries
        are explicitly marked.
        """
        if self._toc is None:
            return

        max_level: int = self._toc["max_level"]
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        fld_char_begin = etree.SubElement(run._element, qn("w:fldChar"))
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = etree.SubElement(run._element, qn("w:instrText"))
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '

        fld_char_end = etree.SubElement(run._element, qn("w:fldChar"))
        fld_char_end.set(qn("w:fldCharType"), "end")

    def _render_header_footer(self, doc: DocxDocument) -> None:
        """Apply document-wide header/footer text via python-docx primitives.

        python-docx exposes ``doc.sections[0].header.paragraphs[0].text`` —
        this is the cleanest path; no lxml drop-down required. The first
        section's header/footer apply to the entire document by default
        (unless a section-break overrides them, which we don't emit).
        """
        if self._header is None and self._footer is None:
            return

        if not doc.sections:  # pragma: no cover
            # Defensive: a brand-new python-docx Document always has a
            # default section, so this branch is theoretically unreachable.
            return

        section = doc.sections[0]
        if self._header is not None:
            section.header.paragraphs[0].text = self._header
        if self._footer is not None:
            section.footer.paragraphs[0].text = self._footer

    def _pin_core_xml_timestamps(self, doc: DocxDocument) -> None:
        """Pin core.xml ``dcterms:created`` / ``dcterms:modified`` for idempotency.

        Without this pin, python-docx populates both fields with
        ``datetime.now()`` during ``save()``; two saves of the same
        Document then yield different fingerprint hashes. See V-MP-DOCUMENT
        forbidden-5 + scenario-8.
        """
        cp = doc.core_properties
        cp.created = DOCUMENT_FIXED_TIMESTAMP
        cp.modified = DOCUMENT_FIXED_TIMESTAMP


__all__ = [
    "DOCUMENT_FIXED_TIMESTAMP",
    "Document",
    "DocumentError",
    "DocumentFormatUnsupportedError",
    "DocumentPresetNotFoundError",
    "DocumentSaveIOError",
    "GotenbergError",
    "PhaseGuardNotImplementedError",
]
