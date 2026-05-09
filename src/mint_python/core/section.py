# FILE: src/mint_python/core/section.py
# VERSION: 0.3.0
# START_MODULE_CONTRACT
#   PURPOSE: Fluent Section node per handover §3.2 — heading + ordered list of
#     content blocks (Paragraph, Table, Image, Chart). Phase-8 unstubs
#     add_chart: appends a Chart block to _blocks and returns self for fluent
#     chaining. Section.render walks block.render(parent_doc) so each Chart
#     block emits its own BLOCK_RENDER_CHART trace marker.
#   SCOPE: Public surface = Section, SectionError, SectionLevelOutOfRangeError,
#     PhaseGuardNotImplementedError. Sibling-only deps: MP-CONTENT (Paragraph,
#     Image), MP-TABLE (Table), MP-CHART (Chart). Section.render(parent_doc)
#     emits a heading paragraph via parent_doc.add_heading(title, level) then
#     iterates blocks calling each .render(parent_doc); children own their own
#     BLOCK markers.
#   DEPENDS: mint_python.core.content (Paragraph, Image), mint_python.core.table
#     (Table), mint_python.core.chart (Chart), python-docx (1.1.x; Document type only).
#   LINKS: docs/development-plan.xml#MP-SECTION,
#     docs/verification-plan.xml#V-MP-SECTION,
#     docs/knowledge-graph.xml#MP-SECTION
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   Section                          - @dataclass; title + level (1..6) + blocks
#   Section.add_paragraph            - str shorthand or Paragraph; returns self
#   Section.add_table                - Table; returns self
#   Section.add_image                - Image; returns self
#   Section.add_chart                - appends Chart to _blocks; returns self for fluent chain
#   Section.add_list                 - appends List to _blocks; returns self for fluent chain
#   Section.add_callout              - appends Callout to _blocks; returns self for fluent chain
#   Section.render                   - heading + ordered block render
#   SectionError                     - base error
#   SectionLevelOutOfRangeError      - SECTION_LEVEL_OUT_OF_RANGE
#   PhaseGuardNotImplementedError    - PHASE_GUARD_NOT_IMPLEMENTED (NotImplementedError subclass)
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: v0.3.0 — add_callout / Callout block support. Section._blocks
#     widened to include MP-CALLOUT.Callout. Sibling-only dep on
#     mint_python.core.callout.
#   PRIOR: v0.2.0 — add_list / List block support.
#   PRIOR: Wave-8-2 (MP-SECTION): unstub add_chart.
#   PRIOR: Wave-7-3 (MP-SECTION): initial implementation per V-MP-SECTION
#     scenarios 1-6 + BLOCK_PHASE_GUARD trace assertion.
# END_CHANGE_SUMMARY

from __future__ import annotations

from dataclasses import dataclass, field

from docx.document import Document as DocxDocument

from mint_python.core.callout import Callout
from mint_python.core.chart import Chart
from mint_python.core.content import Image, Paragraph
from mint_python.core.list_block import List
from mint_python.core.table import Table

# ---------------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------------


class SectionError(Exception):
    """Base for MP-SECTION errors."""


class SectionLevelOutOfRangeError(SectionError):
    """Raised when Section level is not in 1..6 (SECTION_LEVEL_OUT_OF_RANGE).

    python-docx supports level=0..9 on add_heading, but the MP-SECTION contract
    constrains us to 1..6 (matching docx Heading 1..Heading 6 style names) so
    that style-resolution downstream stays predictable.
    """


class PhaseGuardNotImplementedError(NotImplementedError):
    """Phase-N stub touched on a not-yet-implemented method.

    Subclass of :class:`NotImplementedError` so callers using ``except
    NotImplementedError`` keep working; carries the target phase + method name
    in the message for diagnosability.
    """


# ---------------------------------------------------------------------------
# Section
# ---------------------------------------------------------------------------


@dataclass
class Section:
    """Section node holding a heading + ordered content blocks.

    Construction::

        Section("Overview", level=1) \
            .add_paragraph("intro line") \
            .add_table(my_table) \
            .add_image(my_image)

    The fluent setters return ``self`` so callers compose the whole section in
    a single expression. ``level`` is validated at construction time (1..6).

    Phase-8: add_chart is a concrete fluent setter that appends a Chart block
    to _blocks. Section.render iterates _blocks calling block.render(parent_doc)
    so each Chart emits its own BLOCK_RENDER_CHART marker — there is no
    Section-side render marker for charts.
    """

    title: str
    level: int  # 1..6
    _blocks: list[Paragraph | Table | Image | Chart | List | Callout] = field(
        default_factory=list
    )

    def __post_init__(self) -> None:
        if not (1 <= self.level <= 6):
            raise SectionLevelOutOfRangeError(
                f"Section level must be in 1..6, got {self.level}"
            )

    def add_paragraph(
        self, text_or_paragraph: str | Paragraph
    ) -> Section:
        """Append a paragraph block; accepts str shorthand or Paragraph instance.

        - ``str`` is wrapped via ``Paragraph(text)`` so callers don't have to
          import the sibling type for the common case.
        - A ``Paragraph`` instance is appended as-is, preserving its existing
          style + run list. We deliberately do NOT accept a ``style`` kwarg
          here: the str path can be styled by constructing the Paragraph
          explicitly, which avoids the ambiguity of "what wins, the explicit
          Paragraph.style or the section-level override?".
        """
        if isinstance(text_or_paragraph, str):
            self._blocks.append(Paragraph(text_or_paragraph))
        else:
            self._blocks.append(text_or_paragraph)
        return self

    def add_table(self, table: Table) -> Section:
        """Append a Table block; returns self for fluent chaining."""
        self._blocks.append(table)
        return self

    def add_image(self, image: Image) -> Section:
        """Append an Image block; returns self for fluent chaining."""
        self._blocks.append(image)
        return self

    def add_chart(self, chart: Chart) -> Section:
        """Append a Chart to this Section's blocks; returns self for fluent chaining.

        Phase-8: concrete impl. Pre-Phase-8 raised NotImplementedError + emitted
        BLOCK_PHASE_GUARD. Section.render walks block.render(parent_doc) so a
        Chart block emits its own BLOCK_RENDER_CHART trace marker; no
        Section-side marker.
        """
        self._blocks.append(chart)
        return self

    def add_list(self, list_block: List) -> Section:
        """Append a List block; returns self for fluent chaining.

        The List instance owns its own kind (bullet/numbered/checklist), nesting
        level, and items. List.render emits its own BLOCK_RENDER_LIST trace
        marker — no Section-side marker.
        """
        self._blocks.append(list_block)
        return self

    def add_callout(self, callout: Callout) -> Section:
        """Append a Callout block (info / warning / code); returns self.

        The Callout instance owns its kind, body text, and optional title.
        Callout.render emits its own BLOCK_RENDER_CALLOUT trace marker.
        """
        self._blocks.append(callout)
        return self

    def render(self, parent_doc: DocxDocument) -> None:
        """Emit heading + ordered child block renders into ``parent_doc``.

        python-docx ``Document.add_heading(text, level=N)`` maps level 1..6 to
        the built-in ``Heading 1`` .. ``Heading 6`` style. We don't emit our
        own BLOCK marker here — children (Paragraph / Table / Image) own
        their own BLOCK_RENDER_CONTENT / BLOCK_RENDER_TABLE markers and
        emitting an extra section-level marker would over-instrument the
        trace. Callers who need section boundaries can reconstruct them
        from caplog by pairing add_heading payloads with their level.
        """
        parent_doc.add_heading(self.title, level=self.level)
        for block in self._blocks:
            block.render(parent_doc)


__all__ = [
    "PhaseGuardNotImplementedError",
    "Section",
    "SectionError",
    "SectionLevelOutOfRangeError",
]
