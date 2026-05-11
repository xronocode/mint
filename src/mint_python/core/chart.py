# FILE: src/mint_python/core/chart.py
# VERSION: 1.2.0
# START_MODULE_CONTRACT
#   PURPOSE: Chart class with 7 factory constructors (bar, line, stacked_bar,
#     pie, heatmap, waterfall, gantt) plus from_matplotlib / from_seaborn /
#     from_plotly channels. from_plotly converts plotly Figures to PNG via
#     plotly.io.to_image + kaleido and wraps via the from_matplotlib pattern.
#     Renders matplotlib Figures to PNG @ 300 DPI and embeds via python-docx
#     add_picture. Corporate theming is applied per-construction via
#     `matplotlib.rc_context(...)` so global rcParams never leak
#     (V-MP-CHART forbidden-5).
#   SCOPE: Public surface = Chart (dataclass with factories + render),
#     ChartError, ChartInvalidDataError, ChartFigureRenderFailedError,
#     PhaseGuardNotImplementedError. PNG bytes are cached on the instance;
#     module never writes to filesystem (V-MP-CHART forbidden-1).
#   DEPENDS: mint_python.core.style (Style — for color/font theming),
#     matplotlib >= 3.7 (hard dep; backend forced to 'Agg' at import per
#     V-MP-CHART forbidden-2), python-docx 1.1.x (Inches + add_picture).
#     seaborn is OPTIONAL and lazy-imported inside Chart.from_seaborn body.
#   LINKS: docs/development-plan.xml#MP-CHART,
#     docs/verification-plan.xml#V-MP-CHART,
#     docs/knowledge-graph.xml#MP-CHART
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   Chart                              - dataclass; chart_type tag + cached PNG bytes
#   Chart.bar / line / stacked_bar     - composed via ax.bar / ax.plot
#   Chart.pie / heatmap                - ax.pie + ax.imshow with colorbar
#   Chart.waterfall / gantt            - composed bars with cumulative offsets
#   Chart.from_matplotlib              - wraps a user-supplied Figure
#   Chart.from_seaborn                 - lazy-imports seaborn; aliases to from_matplotlib
#   Chart.from_plotly                  - wraps a plotly Figure via to_image (kaleido)
#   Chart.render                       - emits w:drawing via add_picture
#   _apply_corporate_theme             - builds rcParams override dict from Style
#   _CORPORATE_RC                      - chartjunk-free rcParams baseline
#   ChartError / ChartInvalidDataError / ChartFigureRenderFailedError
#   PhaseGuardNotImplementedError
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: Phase-12 — unstub from_plotly via plotly.io.to_image + kaleido.
#     Method now does lazy import of plotly, calls fig.to_image(format="png",
#     scale=2), and wraps PNG bytes in Chart. Imports kaleido hint on failure.
#     Previous: Wave-8-1 (MP-CHART) initial implementation.
# END_CHANGE_SUMMARY

# `from __future__ import annotations` MUST precede all runtime imports per
# Python's grammar (it is a compile-time directive, not executable code). The
# matplotlib backend pin (V-MP-CHART forbidden-2) is therefore the FIRST
# RUNTIME executable block below, before any pyplot import.
from __future__ import annotations

# CRITICAL: matplotlib backend MUST be set to 'Agg' BEFORE any pyplot import.
# V-MP-CHART forbidden-2 invariant. Verified by
# tests.unit._mp_helpers.assert_matplotlib_backend_is_agg().
import matplotlib

matplotlib.use("Agg")  # must precede pyplot import
import matplotlib.pyplot as plt  # noqa: I001  - sort order locked by V-MP-CHART forbidden-2 (matplotlib first)
import io
import logging
from dataclasses import dataclass
from typing import Any

from docx.shared import Inches

from mint_python.core.style import Style

logger = logging.getLogger("mint_python.core.chart")

_DPI: int = 300  # PNG resolution for chart embed


# ---------------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------------


class ChartError(Exception):
    """Base for MP-CHART errors."""


class ChartInvalidDataError(ChartError):
    """Raised when factory data shape is invalid (e.g. labels/values length mismatch)."""


class ChartFigureRenderFailedError(ChartError):
    """Raised when matplotlib savefig raises during PNG bytes capture."""


class PhaseGuardNotImplementedError(NotImplementedError):
    """Phase-N stub touched on a not-yet-implemented method."""


# ---------------------------------------------------------------------------
# Corporate theme rcParams baseline
# ---------------------------------------------------------------------------

# Chartjunk-free baseline applied per-construct via matplotlib.rc_context().
# V-MP-CHART forbidden-5: this dict is consumed inside `with rc_context(...):`
# blocks so global rcParams never leak between tests / Chart instances.
_CORPORATE_RC: dict[str, Any] = {
    "axes.spines.top": False,    # chartjunk: remove top spine
    "axes.spines.right": False,  # chartjunk: remove right spine
    "axes.grid": True,
    "grid.alpha": 0.3,
    "grid.linewidth": 0.5,
    "axes.titlesize": 11,
    "axes.labelsize": 10,
    "xtick.labelsize": 9,
    "ytick.labelsize": 9,
    "legend.fontsize": 9,
    "font.size": 10,
}


def _apply_corporate_theme(style: Style | None) -> dict[str, Any]:
    """Build an rcParams override dict from the active Style preset.

    Returns a dict suitable for ``with matplotlib.rc_context(rc):`` so callers
    can wrap figure construction without leaking rcParams globally
    (V-MP-CHART forbidden-5). When ``style`` is None, only the chartjunk-free
    baseline is returned; when supplied, font.family + axis colors are
    overridden from the preset's typography fields.
    """
    rc: dict[str, Any] = dict(_CORPORATE_RC)
    if style is not None:
        rc["font.family"] = [style.font, "DejaVu Sans"]  # fallback if Inter missing
        rc["axes.labelcolor"] = style.color_hex
        rc["axes.titlecolor"] = style.color_hex
        rc["xtick.color"] = style.color_hex
        rc["ytick.color"] = style.color_hex
    return rc


# ---------------------------------------------------------------------------
# Chart dataclass
# ---------------------------------------------------------------------------


@dataclass
class Chart:
    """Renderable chart wrapper around a cached PNG byte string.

    Construction is via the 7 factory classmethods (bar, line, stacked_bar,
    pie, heatmap, waterfall, gantt) or the 3 from_* channels (from_matplotlib,
    from_seaborn, from_plotly). All factories produce PNG bytes via matplotlib
    @ 300 DPI inside a ``matplotlib.rc_context`` block; bytes are cached on
    the instance (no filesystem writes — V-MP-CHART forbidden-1).

    Render embeds the PNG in a python-docx Document via add_picture; width
    is converted to EMU through ``docx.shared.Inches(width_inches)`` (NOT a
    hand-multiply, which would lose precision and trip VF-014 inv-3).

    Attributes:
        chart_type: Tag for trace markers; one of "bar", "line", "stacked_bar",
            "pie", "heatmap", "waterfall", "gantt", "matplotlib", "seaborn".
        _png_bytes: Cached PNG bytes captured at factory time.
        width_inches: Embed width in inches (default 5.0).
        height_inches: Embed height in inches; None lets python-docx preserve
            aspect ratio.
        caption: Optional caption text (consumed by Section.add_chart in
            Wave-8-2).
        vector: Accepted for forward-compat; Phase-8 always emits PNG @ 300
            DPI. True SVG/EMF deferred to Phase-8.5.
    """

    chart_type: str
    _png_bytes: bytes
    width_inches: float = 5.0
    height_inches: float | None = None
    caption: str | None = None
    vector: bool = True

    # ------------------------------------------------------------------
    # Render
    # ------------------------------------------------------------------

    def render(self, parent_doc: Any) -> Any:
        """Embed the cached PNG into ``parent_doc`` via add_picture.

        EMU PRECISION: width is converted via ``docx.shared.Inches`` —
        NOT hand-multiplied — so VF-014 inv-3 ``width.emu == round(inches *
        914400)`` holds without rounding drift.
        """
        emu_width = round(self.width_inches * 914400)
        # START_BLOCK_RENDER_CHART
        logger.info(
            "[MP-Chart][render][BLOCK_RENDER_CHART] chart_type=%s emu_width=%d",
            self.chart_type,
            emu_width,
        )
        # END_BLOCK_RENDER_CHART
        kwargs: dict[str, Any] = {"width": Inches(self.width_inches)}
        if self.height_inches is not None:
            kwargs["height"] = Inches(self.height_inches)
        return parent_doc.add_picture(io.BytesIO(self._png_bytes), **kwargs)

    # ------------------------------------------------------------------
    # Internal: figure -> PNG bytes
    # ------------------------------------------------------------------

    @staticmethod
    def _fig_to_png_bytes(fig: Any) -> bytes:
        """Save ``fig`` to an in-memory PNG buffer @ _DPI; return raw bytes.

        Wraps savefig errors in ChartFigureRenderFailedError so callers see a
        domain-specific exception rather than the matplotlib internal.
        """
        buf = io.BytesIO()
        try:
            fig.savefig(buf, format="png", dpi=_DPI, bbox_inches="tight")
        except Exception as exc:
            raise ChartFigureRenderFailedError(
                f"matplotlib savefig failed: {exc}"
            ) from exc
        return buf.getvalue()

    # ==================================================================
    # Factory: bar
    # ==================================================================

    @classmethod
    def bar(
        cls,
        labels: list[str],
        values: list[float],
        *,
        title: str | None = None,
        caption: str | None = None,
        width_inches: float = 5.0,
        height_inches: float | None = None,
        style: Style | None = None,
    ) -> Chart:
        """Bar chart from parallel labels/values arrays."""
        if len(labels) != len(values):
            raise ChartInvalidDataError(
                f"labels and values length mismatch: "
                f"len(labels)={len(labels)}, len(values)={len(values)}"
            )
        rc = _apply_corporate_theme(style)
        # START_BLOCK_BUILD_CHART_BAR
        logger.info(
            "[MP-Chart][bar][BLOCK_BUILD_CHART] chart_type=bar data_shape=%dx1",
            len(labels),
        )
        # END_BLOCK_BUILD_CHART_BAR
        with matplotlib.rc_context(rc):
            fig, ax = plt.subplots(
                figsize=(width_inches, height_inches or 3.0), dpi=_DPI
            )
            ax.bar(labels, values)
            if title:
                ax.set_title(title)
            png_bytes = cls._fig_to_png_bytes(fig)
        plt.close(fig)
        return cls(
            chart_type="bar",
            _png_bytes=png_bytes,
            width_inches=width_inches,
            height_inches=height_inches,
            caption=caption,
        )

    # ==================================================================
    # Factory: line
    # ==================================================================

    @classmethod
    def line(
        cls,
        labels: list[str],
        values: list[float],
        *,
        title: str | None = None,
        caption: str | None = None,
        width_inches: float = 5.0,
        height_inches: float | None = None,
        style: Style | None = None,
    ) -> Chart:
        """Line chart from parallel labels/values arrays."""
        if len(labels) != len(values):
            raise ChartInvalidDataError(
                f"labels and values length mismatch: "
                f"len(labels)={len(labels)}, len(values)={len(values)}"
            )
        rc = _apply_corporate_theme(style)
        # START_BLOCK_BUILD_CHART_LINE
        logger.info(
            "[MP-Chart][line][BLOCK_BUILD_CHART] chart_type=line data_shape=%dx1",
            len(labels),
        )
        # END_BLOCK_BUILD_CHART_LINE
        with matplotlib.rc_context(rc):
            fig, ax = plt.subplots(
                figsize=(width_inches, height_inches or 3.0), dpi=_DPI
            )
            ax.plot(labels, values)
            if title:
                ax.set_title(title)
            png_bytes = cls._fig_to_png_bytes(fig)
        plt.close(fig)
        return cls(
            chart_type="line",
            _png_bytes=png_bytes,
            width_inches=width_inches,
            height_inches=height_inches,
            caption=caption,
        )

    # ==================================================================
    # Factory: stacked_bar
    # ==================================================================

    @classmethod
    def stacked_bar(
        cls,
        labels: list[str],
        series: dict[str, list[float]],
        *,
        title: str | None = None,
        caption: str | None = None,
        width_inches: float = 5.0,
        height_inches: float | None = None,
        style: Style | None = None,
    ) -> Chart:
        """Stacked bar chart: ``series`` is a dict[name -> list[float]]."""
        for name, vals in series.items():
            if len(vals) != len(labels):
                raise ChartInvalidDataError(
                    f"series '{name}' length mismatch: "
                    f"len(values)={len(vals)}, len(labels)={len(labels)}"
                )
        rc = _apply_corporate_theme(style)
        # START_BLOCK_BUILD_CHART_STACKED_BAR
        logger.info(
            "[MP-Chart][stacked_bar][BLOCK_BUILD_CHART] "
            "chart_type=stacked_bar data_shape=%dx%d",
            len(labels),
            len(series),
        )
        # END_BLOCK_BUILD_CHART_STACKED_BAR
        with matplotlib.rc_context(rc):
            fig, ax = plt.subplots(
                figsize=(width_inches, height_inches or 3.0), dpi=_DPI
            )
            cumulative = [0.0] * len(labels)
            for name, vals in series.items():
                ax.bar(labels, vals, bottom=cumulative, label=name)
                cumulative = [c + v for c, v in zip(cumulative, vals, strict=True)]
            if series:
                ax.legend()
            if title:
                ax.set_title(title)
            png_bytes = cls._fig_to_png_bytes(fig)
        plt.close(fig)
        return cls(
            chart_type="stacked_bar",
            _png_bytes=png_bytes,
            width_inches=width_inches,
            height_inches=height_inches,
            caption=caption,
        )

    # ==================================================================
    # Factory: pie
    # ==================================================================

    @classmethod
    def pie(
        cls,
        labels: list[str],
        values: list[float],
        *,
        title: str | None = None,
        caption: str | None = None,
        width_inches: float = 5.0,
        height_inches: float | None = None,
        style: Style | None = None,
    ) -> Chart:
        """Pie chart with autopct labels."""
        if len(labels) != len(values):
            raise ChartInvalidDataError(
                f"labels and values length mismatch: "
                f"len(labels)={len(labels)}, len(values)={len(values)}"
            )
        rc = _apply_corporate_theme(style)
        # START_BLOCK_BUILD_CHART_PIE
        logger.info(
            "[MP-Chart][pie][BLOCK_BUILD_CHART] chart_type=pie data_shape=%dx1",
            len(labels),
        )
        # END_BLOCK_BUILD_CHART_PIE
        with matplotlib.rc_context(rc):
            fig, ax = plt.subplots(
                figsize=(width_inches, height_inches or 3.0), dpi=_DPI
            )
            ax.pie(values, labels=labels, autopct="%.1f%%")
            if title:
                ax.set_title(title)
            png_bytes = cls._fig_to_png_bytes(fig)
        plt.close(fig)
        return cls(
            chart_type="pie",
            _png_bytes=png_bytes,
            width_inches=width_inches,
            height_inches=height_inches,
            caption=caption,
        )

    # ==================================================================
    # Factory: heatmap
    # ==================================================================

    @classmethod
    def heatmap(
        cls,
        matrix: list[list[float]],
        row_labels: list[str],
        col_labels: list[str],
        *,
        title: str | None = None,
        caption: str | None = None,
        width_inches: float = 5.0,
        height_inches: float | None = None,
        style: Style | None = None,
    ) -> Chart:
        """Heatmap via imshow + colorbar.

        ``matrix`` shape MUST equal (len(row_labels), len(col_labels)).
        """
        if len(matrix) != len(row_labels):
            raise ChartInvalidDataError(
                f"matrix row count mismatch: "
                f"len(matrix)={len(matrix)}, len(row_labels)={len(row_labels)}"
            )
        for i, row in enumerate(matrix):
            if len(row) != len(col_labels):
                raise ChartInvalidDataError(
                    f"matrix row {i} column count mismatch: "
                    f"len(row)={len(row)}, len(col_labels)={len(col_labels)}"
                )
        rc = _apply_corporate_theme(style)
        # START_BLOCK_BUILD_CHART_HEATMAP
        logger.info(
            "[MP-Chart][heatmap][BLOCK_BUILD_CHART] "
            "chart_type=heatmap data_shape=%dx%d",
            len(row_labels),
            len(col_labels),
        )
        # END_BLOCK_BUILD_CHART_HEATMAP
        with matplotlib.rc_context(rc):
            fig, ax = plt.subplots(
                figsize=(width_inches, height_inches or 3.0), dpi=_DPI
            )
            im = ax.imshow(matrix, aspect="auto")
            ax.set_xticks(range(len(col_labels)))
            ax.set_xticklabels(col_labels)
            ax.set_yticks(range(len(row_labels)))
            ax.set_yticklabels(row_labels)
            fig.colorbar(im, ax=ax)
            if title:
                ax.set_title(title)
            png_bytes = cls._fig_to_png_bytes(fig)
        plt.close(fig)
        return cls(
            chart_type="heatmap",
            _png_bytes=png_bytes,
            width_inches=width_inches,
            height_inches=height_inches,
            caption=caption,
        )

    # ==================================================================
    # Factory: waterfall
    # ==================================================================

    @classmethod
    def waterfall(
        cls,
        labels: list[str],
        deltas: list[float],
        *,
        title: str | None = None,
        caption: str | None = None,
        width_inches: float = 5.0,
        height_inches: float | None = None,
        style: Style | None = None,
    ) -> Chart:
        """Waterfall chart composed via cumulative-sum offset bars.

        Each bar is annotated with its signed delta. The cumulative sum
        across all deltas is annotated under the final bar (V-MP-CHART
        scenario-6 sum-of-deltas evidence).
        """
        if len(labels) != len(deltas):
            raise ChartInvalidDataError(
                f"labels and deltas length mismatch: "
                f"len(labels)={len(labels)}, len(deltas)={len(deltas)}"
            )
        rc = _apply_corporate_theme(style)
        # START_BLOCK_BUILD_CHART_WATERFALL
        logger.info(
            "[MP-Chart][waterfall][BLOCK_BUILD_CHART] "
            "chart_type=waterfall data_shape=%dx1",
            len(labels),
        )
        # END_BLOCK_BUILD_CHART_WATERFALL
        with matplotlib.rc_context(rc):
            fig, ax = plt.subplots(
                figsize=(width_inches, height_inches or 3.0), dpi=_DPI
            )
            cumulative = 0.0
            bottoms: list[float] = []
            for d in deltas:
                if d >= 0:
                    bottoms.append(cumulative)
                else:
                    bottoms.append(cumulative + d)
                cumulative += d
            heights = [abs(d) for d in deltas]
            ax.bar(labels, heights, bottom=bottoms)
            # Annotate each bar with its signed delta.
            running = 0.0
            for i, d in enumerate(deltas):
                running += d
                ax.annotate(
                    f"{d:+g}",
                    xy=(i, running),
                    ha="center",
                    va="bottom",
                    fontsize=8,
                )
            # Sum-of-deltas annotation under last bar.
            total = sum(deltas)
            ax.annotate(
                f"sum={total:+g}",
                xy=(len(labels) - 1, 0),
                xytext=(0, -20),
                textcoords="offset points",
                ha="center",
                va="top",
                fontsize=8,
            )
            if title:
                ax.set_title(title)
            png_bytes = cls._fig_to_png_bytes(fig)
        plt.close(fig)
        return cls(
            chart_type="waterfall",
            _png_bytes=png_bytes,
            width_inches=width_inches,
            height_inches=height_inches,
            caption=caption,
        )

    # ==================================================================
    # Factory: gantt
    # ==================================================================

    @classmethod
    def gantt(
        cls,
        tasks: list[tuple[str, float, float]],
        *,
        title: str | None = None,
        caption: str | None = None,
        width_inches: float = 5.0,
        height_inches: float | None = None,
        style: Style | None = None,
    ) -> Chart:
        """Gantt chart from list of (name, start, duration) tuples.

        Tasks are stacked vertically with name labels on the y-axis.
        """
        if not tasks:
            raise ChartInvalidDataError("tasks list must be non-empty")
        for i, t in enumerate(tasks):
            if not isinstance(t, tuple) or len(t) != 3:
                raise ChartInvalidDataError(
                    f"task {i} must be a 3-tuple (name, start, duration); got {t!r}"
                )
        rc = _apply_corporate_theme(style)
        # START_BLOCK_BUILD_CHART_GANTT
        logger.info(
            "[MP-Chart][gantt][BLOCK_BUILD_CHART] chart_type=gantt data_shape=%dx1",
            len(tasks),
        )
        # END_BLOCK_BUILD_CHART_GANTT
        with matplotlib.rc_context(rc):
            fig, ax = plt.subplots(
                figsize=(width_inches, height_inches or 3.0), dpi=_DPI
            )
            names = [t[0] for t in tasks]
            for y, (_, start, duration) in enumerate(tasks):
                ax.barh(y, duration, left=start)
            ax.set_yticks(range(len(tasks)))
            ax.set_yticklabels(names)
            ax.invert_yaxis()  # first task at top
            if title:
                ax.set_title(title)
            png_bytes = cls._fig_to_png_bytes(fig)
        plt.close(fig)
        return cls(
            chart_type="gantt",
            _png_bytes=png_bytes,
            width_inches=width_inches,
            height_inches=height_inches,
            caption=caption,
        )

    # ==================================================================
    # from_matplotlib
    # ==================================================================

    @classmethod
    def from_matplotlib(
        cls,
        fig: Any,
        *,
        caption: str | None = None,
        width_inches: float = 5.0,
        height_inches: float | None = None,
        vector: bool = True,
    ) -> Chart:
        """Wrap a user-supplied matplotlib Figure as a Chart.

        The ``vector`` kwarg is accepted for forward-compat and stored on the
        resulting instance; Phase-8 ALWAYS emits PNG @ 300 DPI. True SVG/EMF
        is deferred to Phase-8.5.
        """
        # START_BLOCK_BUILD_CHART_FROM_MATPLOTLIB
        logger.info(
            "[MP-Chart][from_matplotlib][BLOCK_BUILD_CHART] "
            "chart_type=matplotlib data_shape=user-figure",
        )
        # END_BLOCK_BUILD_CHART_FROM_MATPLOTLIB
        png_bytes = cls._fig_to_png_bytes(fig)
        plt.close(fig)
        return cls(
            chart_type="matplotlib",
            _png_bytes=png_bytes,
            width_inches=width_inches,
            height_inches=height_inches,
            caption=caption,
            vector=vector,
        )

    # ==================================================================
    # from_seaborn (lazy import)
    # ==================================================================

    @classmethod
    def from_seaborn(cls, fig: Any, **kwargs: Any) -> Chart:
        """Wrap a seaborn-built Figure (which is a matplotlib Figure).

        seaborn is OPTIONAL — this method lazy-imports inside the body so
        ``import mint_python.core.chart`` does not pull seaborn into
        sys.modules (V-MP-CHART forbidden-3). When seaborn is missing,
        ImportError is re-raised with a ``pip install seaborn`` hint.
        """
        try:
            import seaborn as sns  # noqa: F401  - lazy per V-MP-CHART forbidden-3
        except ImportError as exc:
            raise ImportError(
                "seaborn is optional — install with `pip install seaborn` "
                "to use Chart.from_seaborn"
            ) from exc
        # seaborn returns a matplotlib Figure; alias to from_matplotlib.
        chart = cls.from_matplotlib(fig, **kwargs)
        chart.chart_type = "seaborn"  # override for trace clarity
        return chart

    # ==================================================================
    # from_plotly (Phase-12 — concrete via plotly.io.to_image + kaleido)
    # ==================================================================

    @classmethod
    def from_plotly(cls, fig: Any, **kwargs: Any) -> Chart:
        """Wrap a plotly Figure.

        Converts plotly Figure to PNG bytes via plotly.io.to_image (requires
        kaleido package: ``pip install kaleido``). Then wraps via
        from_matplotlib pattern — stores PNG bytes in Chart._png_bytes.

        Raises ImportError with hint if plotly or kaleido is missing.
        """
        try:
            import plotly  # noqa: F401
        except ImportError as exc:
            raise ImportError(
                "plotly is optional — install with `pip install plotly kaleido` "
                "to use Chart.from_plotly"
            ) from exc

        try:
            png_bytes = fig.to_image(format="png", scale=2)
        except Exception as exc:
            raise ImportError(
                "kaleido is required for plotly PNG export — "
                "install with `pip install kaleido`"
            ) from exc

        caption = kwargs.pop("caption", None)
        width_inches = kwargs.pop("width_inches", None)
        height_inches = kwargs.pop("height_inches", None)

        chart = cls(
            chart_type="plotly",
            _png_bytes=png_bytes,
            caption=caption,
            width_inches=width_inches if width_inches is not None else 6.0,
            height_inches=height_inches if height_inches is not None else 4.0,
        )
        logger.info(
            "[MP-Chart][from_plotly][BLOCK_BUILD_CHART_PLOTLY] "
            "chart_type=plotly size=%d bytes",
            len(png_bytes),
        )
        return chart
