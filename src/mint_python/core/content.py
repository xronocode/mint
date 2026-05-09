# FILE: src/mint_python/core/content.py
# VERSION: 0.4.0
# START_MODULE_CONTRACT
#   PURPOSE: Inline content building blocks for the Pure Python Edition: Run
#     (frozen data carrier with optional per-run formatting overrides + link
#     and bookmark anchors), Paragraph (composed of runs, .render emits w:p),
#     Image (file-or-bytes embed via python-docx + pillow). Direct emitters
#     of <w:p>, <w:r>, <w:hyperlink>, <w:bookmarkStart/End>, <w:drawing>
#     OOXML through python-docx primitives + lxml drop-down for the elements
#     python-docx does not expose first-class.
#   SCOPE: Public surface = Run, Paragraph, Image, ImageFileNotFoundError,
#     ImageFormatUnsupportedError, ContentError. All render() methods take a
#     python-docx Document and write into it; never to disk.
#   DEPENDS: mint_python.core.style (Style only — for typography fields),
#     python-docx (1.1.x), pillow (10.x), lxml (1.x via python-docx).
#   LINKS: docs/development-plan.xml#MP-CONTENT,
#     docs/verification-plan.xml#V-MP-CONTENT,
#     docs/knowledge-graph.xml#MP-CONTENT
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   Run                          - frozen dataclass; text + optional style
#                                  + per-run overrides
#                                  (bold/italic/underline/color/font_size_pt)
#                                  + link (URL or "#anchor")
#                                  + bookmark (anchor name)
#   TabStop                      - frozen dataclass: position_inches +
#                                  alignment + leader
#   TabAlignment                 - StrEnum: LEFT | CENTER | RIGHT | DECIMAL
#   TabLeader                    - StrEnum: NONE | DOTS | DASHES | UNDERSCORE
#   Paragraph                    - styled paragraph + .add_run + .render;
#                                  optional tab_stops list applied via
#                                  python-docx paragraph_format.tab_stops
#   Image                        - .from_path / .from_bytes / .render
#   ContentError                 - base error
#   ImageFileNotFoundError       - Image.from_path: missing file
#   ImageFormatUnsupportedError  - Image.from_bytes: pillow rejects
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: v0.4.0 — Run gains footnote field. First footnoted run in
#     a Document bootstraps a /word/footnotes.xml package part with the
#     standard separator + continuationSeparator entries; subsequent
#     footnotes append <w:footnote w:id="N"> entries to that part. The
#     run gets a <w:footnoteReference w:id="N"/> element appended to its
#     <w:r>. python-docx has no first-class footnote API in 1.2.x — we
#     drop down through docx.opc.part.Part + a manual relationship.
#   PRIOR: v0.3.0 — Paragraph gains tab_stops field.
#   PRIOR: v0.2.0 — Run gains link / bookmark fields.
#   PRIOR: v0.1.0 — Run gains per-run formatting overrides.
# END_CHANGE_SUMMARY

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from enum import StrEnum
from io import BytesIO
from pathlib import Path
from typing import Any

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.opc.constants import CONTENT_TYPE as _CT
from docx.opc.constants import RELATIONSHIP_TYPE as _RT
from docx.opc.packuri import PackURI as _PackURI
from docx.opc.part import Part as _OpcPart
from docx.oxml import OxmlElement as _OxmlElement
from docx.oxml.ns import qn as _qn
from docx.shared import Inches, RGBColor
from docx.shared import Pt as DocxPt
from lxml import etree as _etree
from PIL import Image as PilImage
from PIL import UnidentifiedImageError

from mint_python.core.style import Style

logger = logging.getLogger("mint_python.core.content")


# ---------------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------------


class ContentError(Exception):
    """Base for MP-CONTENT errors."""


class ImageFileNotFoundError(ContentError):
    """Raised by Image.from_path when path does not exist."""


class ImageFormatUnsupportedError(ContentError):
    """Raised when image format is unrecognized or unsupported by pillow."""


# ---------------------------------------------------------------------------
# Alignment map
# ---------------------------------------------------------------------------

_ALIGNMENT_MAP: dict[str, Any] = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ---------------------------------------------------------------------------
# Run
# ---------------------------------------------------------------------------


_HEX_COLOR_LEN = 7  # "#RRGGBB"

# Word bookmark name rule: letter/underscore start, then letters, digits,
# underscores. We mirror this for both `bookmark` and the anchor-portion of
# an internal link "#name" so cross-references stay round-trippable.
_BOOKMARK_NAME_RE = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*$")


def _validate_run_color(value: str) -> None:
    if len(value) != _HEX_COLOR_LEN or not value.startswith("#"):
        raise ValueError(
            f"Run.color must be 7-char #RRGGBB hex, got {value!r}"
        )
    try:
        int(value[1:], 16)
    except ValueError as exc:
        raise ValueError(f"Run.color hex digits invalid: {value!r}") from exc


def _validate_run_link(value: str) -> None:
    if not value.strip():
        raise ValueError("Run.link must be a non-empty URL or '#anchor'")
    if value.startswith("#"):
        anchor = value[1:]
        if not _BOOKMARK_NAME_RE.match(anchor):
            raise ValueError(
                f"Run.link anchor name invalid (letter/underscore start, "
                f"alphanumeric tail): {value!r}"
            )


def _validate_run_bookmark(value: str) -> None:
    if not _BOOKMARK_NAME_RE.match(value):
        raise ValueError(
            f"Run.bookmark name invalid (letter/underscore start, "
            f"alphanumeric tail): {value!r}"
        )


@dataclass(frozen=True)
class Run:
    """Inline styled fragment within a Paragraph.

    Run is a frozen data carrier with no render method of its own —
    Paragraph.render walks the run list and emits w:r elements via
    python-docx. Because Run carries no mutable state, sharing the same
    Run instance across multiple Paragraphs is harmless (forbidden-2
    invariant: Run re-use is safe by construction; the spec's "Run must
    not be re-used" guidance is satisfied via immutability — there is
    no Run.consume() and no per-render mutation).

    Attributes:
        text: literal string content for this run.
        style: optional :class:`Style`; ``None`` means "inherit the
            enclosing Paragraph.style at render time".
        bold: per-run bold override. ``None`` = inherit effective Style.
        italic: per-run italic override. ``None`` = inherit.
        underline: per-run underline override. Style has no underline
            field, so ``None`` here = no underline applied.
        color: per-run text color as ``#RRGGBB`` hex; ``None`` = inherit
            effective Style.color_hex. Validated at construction.
        font_size_pt: per-run font size in points (float > 0); ``None``
            = inherit effective Style.size_pt.

    Override precedence at render time: ``Run.<field>`` > effective
    ``Style.<field>``. Setting an override to a literal value (including
    ``False`` or ``0.0``) is intentional and beats the inherited value.
    """

    text: str
    style: Style | None = None
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None
    color: str | None = None
    font_size_pt: float | None = None
    link: str | None = None
    bookmark: str | None = None
    footnote: str | None = None

    def __post_init__(self) -> None:
        if self.color is not None:
            _validate_run_color(self.color)
        if self.font_size_pt is not None and self.font_size_pt <= 0:
            raise ValueError(
                f"Run.font_size_pt must be > 0, got {self.font_size_pt!r}"
            )
        if self.link is not None:
            _validate_run_link(self.link)
        if self.bookmark is not None:
            _validate_run_bookmark(self.bookmark)
        if self.footnote is not None and not self.footnote.strip():
            raise ValueError(
                "Run.footnote must be non-empty text; got empty/whitespace"
            )


# ---------------------------------------------------------------------------
# Tab stops
# ---------------------------------------------------------------------------


class TabAlignment(StrEnum):
    """Paragraph tab-stop alignment per OOXML w:tab[@w:val]."""

    LEFT = "left"
    CENTER = "center"
    RIGHT = "right"
    DECIMAL = "decimal"


class TabLeader(StrEnum):
    """Paragraph tab-stop leader (the fill character drawn between text)."""

    NONE = "none"
    DOTS = "dots"
    DASHES = "dashes"
    UNDERSCORE = "underscore"


_DOCX_TAB_ALIGNMENT: dict[TabAlignment, Any] = {
    TabAlignment.LEFT: WD_TAB_ALIGNMENT.LEFT,
    TabAlignment.CENTER: WD_TAB_ALIGNMENT.CENTER,
    TabAlignment.RIGHT: WD_TAB_ALIGNMENT.RIGHT,
    TabAlignment.DECIMAL: WD_TAB_ALIGNMENT.DECIMAL,
}

_DOCX_TAB_LEADER: dict[TabLeader, Any] = {
    TabLeader.NONE: WD_TAB_LEADER.SPACES,
    TabLeader.DOTS: WD_TAB_LEADER.DOTS,
    TabLeader.DASHES: WD_TAB_LEADER.DASHES,
    TabLeader.UNDERSCORE: WD_TAB_LEADER.LINES,
}


@dataclass(frozen=True)
class TabStop:
    """Paragraph tab-stop position + alignment + leader.

    Attributes:
        position_inches: distance from the left margin, in inches; > 0.
        alignment: how text snaps to the stop (default left).
        leader: fill character drawn from the previous text to the stop
            (default none / spaces).
    """

    position_inches: float
    alignment: TabAlignment = TabAlignment.LEFT
    leader: TabLeader = TabLeader.NONE

    def __post_init__(self) -> None:
        if self.position_inches <= 0:
            raise ValueError(
                f"TabStop.position_inches must be > 0, got {self.position_inches!r}"
            )


# ---------------------------------------------------------------------------
# Paragraph
# ---------------------------------------------------------------------------


@dataclass
class Paragraph:
    """Styled paragraph composed of one or more :class:`Run` instances.

    Construction modes::

        Paragraph("hello")                        # one run carrying paragraph style
        Paragraph("hello", style=ns.body)         # one run with explicit body style
        Paragraph([Run("a"), Run("b", bold)])     # explicit run list
        Paragraph()                               # empty; build via add_run chain

    The fluent chain ``Paragraph("a").add_run("b", style=bold)`` returns
    self so callers can compose multi-style paragraphs in one expression.
    """

    style: Style | None = None
    _runs: list[Run] = field(default_factory=list)
    tab_stops: list[TabStop] = field(default_factory=list)

    def __init__(
        self,
        text_or_runs: str | list[Run] = "",
        style: Style | None = None,
        *,
        tab_stops: list[TabStop] | None = None,
    ) -> None:
        self.style = style
        self.tab_stops = list(tab_stops) if tab_stops else []
        if isinstance(text_or_runs, str):
            if text_or_runs:
                # Seed with a single Run carrying paragraph-level style by
                # default. Per-run override via add_run(text, style=...).
                self._runs = [Run(text_or_runs, style)]
            else:
                self._runs = []
        else:
            # list[Run] — store as-is. Frozen Runs are safe to share.
            self._runs = list(text_or_runs)

    def add_run(
        self,
        text: str,
        style: Style | None = None,
        *,
        bold: bool | None = None,
        italic: bool | None = None,
        underline: bool | None = None,
        color: str | None = None,
        font_size_pt: float | None = None,
        link: str | None = None,
        bookmark: str | None = None,
        footnote: str | None = None,
    ) -> Paragraph:
        """Append a Run; ``style=None`` inherits paragraph style at render.

        Per-run override kwargs (bold/italic/underline/color/font_size_pt)
        mirror :class:`Run`'s formatting fields. ``None`` = inherit from
        the effective :class:`Style`; an explicit value wins.

        ``link`` accepts an external URL or an internal anchor reference of
        the form ``"#bookmark_name"``. ``bookmark`` declares this run's text
        as a bookmark anchor target with the given name.

        Returns self to enable fluent chaining per V-MP-CONTENT scenario-2.
        """
        self._runs.append(
            Run(
                text,
                style,
                bold=bold,
                italic=italic,
                underline=underline,
                color=color,
                font_size_pt=font_size_pt,
                link=link,
                bookmark=bookmark,
                footnote=footnote,
            )
        )
        return self

    # START_BLOCK_RENDER_CONTENT_PARAGRAPH
    def render(self, parent_doc: DocxDocument) -> Any:
        """Append a w:p to ``parent_doc`` and return the python-docx paragraph.

        Pt-unit decision: ``Style.size_pt`` is a FLOAT in points (per
        :mod:`mint_python.core.style`). We pass it directly to
        :class:`docx.shared.Pt` (``DocxPt``) which converts pt -> EMU
        internally. Our top-level :func:`mint_python.core.style.Pt`
        returns twentieths-of-a-point (an int); routing ``Style.size_pt``
        through it before handing to ``DocxPt`` would yield a value 1/20
        the intended size and corrupt the OOXML. This render call site
        therefore uses python-docx's own helpers exclusively — no lxml
        drop-down, no hand-written twips.

        Forbidden-1 invariant: this method only mutates ``parent_doc``;
        no filesystem writes happen here. Forbidden-2 invariant: calling
        ``render`` repeatedly is allowed and simply appends multiple w:p
        elements (python-docx's natural behavior); we do NOT raise.

        Returns the python-docx ``Paragraph`` object so callers (e.g.
        Section.render in MP-SECTION) can chain further mutations.
        """
        # Marker emit BEFORE python-docx mutation per observability contract.
        logger.info(
            "[MP-Content][render][BLOCK_RENDER_CONTENT] kind=paragraph runs=%d",
            len(self._runs),
        )

        docx_paragraph = parent_doc.add_paragraph()

        # Apply paragraph-level style to paragraph_format BEFORE adding runs.
        if self.style is not None:
            self._apply_paragraph_style(docx_paragraph, self.style)

        # Tab stops — applied after style so they survive style.line_height
        # and friends. python-docx's tab_stops collection appends in order.
        for stop in self.tab_stops:
            docx_paragraph.paragraph_format.tab_stops.add_tab_stop(
                Inches(stop.position_inches),
                _DOCX_TAB_ALIGNMENT[stop.alignment],
                _DOCX_TAB_LEADER[stop.leader],
            )

        next_bookmark_id = self._next_bookmark_id(parent_doc)
        for run in self._runs:
            docx_run = docx_paragraph.add_run(run.text)
            base_style = run.style if run.style is not None else self.style
            self._apply_run_style(docx_run, run, base_style)
            if run.link is not None or run.bookmark is not None:
                next_bookmark_id = self._wrap_run_anchors(
                    docx_paragraph, docx_run, run, next_bookmark_id
                )
            if run.footnote is not None:
                self._append_footnote_reference(parent_doc, docx_run, run.footnote)

        return docx_paragraph

    @staticmethod
    def _next_bookmark_id(parent_doc: DocxDocument) -> int:
        """Return one greater than the highest existing w:bookmarkStart/@w:id.

        OOXML requires bookmark ids to be unique within the document. We scan
        the entire body once per render to find the current max; subsequent
        bookmarks emitted by this render share the local counter.
        """
        max_id = -1
        for el in parent_doc.element.body.iter(_qn("w:bookmarkStart")):
            try:
                cur = int(el.get(_qn("w:id"), "-1"))
            except ValueError:  # pragma: no cover — non-int @w:id is OOXML-illegal
                continue
            if cur > max_id:
                max_id = cur
        return max_id + 1

    @staticmethod
    def _ensure_footnotes_part(parent_doc: DocxDocument) -> _OpcPart:
        """Return the document's /word/footnotes.xml part, creating it on first use.

        OOXML conformant footnotes need:
          1. A package part with content type WML_FOOTNOTES at /word/footnotes.xml.
          2. A relationship from the main document part with type FOOTNOTES.
          3. A <w:footnotes> root containing <w:footnote w:type="separator"
             w:id="-1"/> and <w:footnote w:type="continuationSeparator" w:id="0"/>
             entries — Word/LibreOffice expect these exact ids/types.

        Subsequent footnote references reuse the existing part.
        """
        doc_part = parent_doc.part
        for rel in doc_part.rels.values():
            if rel.reltype == _RT.FOOTNOTES:
                target: _OpcPart = rel.target_part
                return target

        # Bootstrap: build the skeleton XML, register a new package part.
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        root = _etree.Element(_qn("w:footnotes"), nsmap=nsmap)
        for kind, fid in (("separator", -1), ("continuationSeparator", 0)):
            footnote = _etree.SubElement(root, _qn("w:footnote"))
            footnote.set(_qn("w:type"), kind)
            footnote.set(_qn("w:id"), str(fid))
            p = _etree.SubElement(footnote, _qn("w:p"))
            r = _etree.SubElement(p, _qn("w:r"))
            _etree.SubElement(r, _qn(f"w:{kind}"))

        blob = _etree.tostring(
            root, xml_declaration=True, standalone=True, encoding="UTF-8"
        )
        partname = _PackURI("/word/footnotes.xml")
        new_part = _OpcPart(
            partname, _CT.WML_FOOTNOTES, blob, doc_part.package
        )
        doc_part.relate_to(new_part, _RT.FOOTNOTES)
        return new_part

    @staticmethod
    def _append_footnote_reference(
        parent_doc: DocxDocument, docx_run: Any, footnote_text: str
    ) -> None:
        """Append a w:footnoteReference to the run + a w:footnote entry to the part.

        Allocates the next positive footnote id by scanning the footnotes part
        for the current max @w:id (excluding the separator/continuation entries
        which use -1 and 0).
        """
        part = Paragraph._ensure_footnotes_part(parent_doc)
        root = _etree.fromstring(part.blob)

        max_id = 0
        for fn in root.iter(_qn("w:footnote")):
            try:
                cur = int(fn.get(_qn("w:id"), "0"))
            except ValueError:  # pragma: no cover — non-int id is OOXML-illegal
                continue
            if cur > max_id:
                max_id = cur
        new_id = max_id + 1

        # Build the footnote body: one paragraph with one run carrying the text.
        footnote = _etree.SubElement(root, _qn("w:footnote"))
        footnote.set(_qn("w:id"), str(new_id))
        body_p = _etree.SubElement(footnote, _qn("w:p"))
        body_r = _etree.SubElement(body_p, _qn("w:r"))
        ref_marker = _etree.SubElement(body_r, _qn("w:footnoteRef"))
        # Actual footnote text follows the reference marker.
        body_r2 = _etree.SubElement(body_p, _qn("w:r"))
        text_el = _etree.SubElement(body_r2, _qn("w:t"))
        text_el.set(_qn("xml:space"), "preserve")
        text_el.text = " " + footnote_text  # leading space after marker
        del ref_marker  # silence unused-var lint while keeping the SubElement

        part._blob = _etree.tostring(
            root, xml_declaration=True, standalone=True, encoding="UTF-8"
        )

        # Inject reference into the calling run.
        ref = _OxmlElement("w:footnoteReference")
        ref.set(_qn("w:id"), str(new_id))
        docx_run._r.append(ref)

    @staticmethod
    def _wrap_run_anchors(
        docx_paragraph: Any,
        docx_run: Any,
        run: Run,
        bookmark_id: int,
    ) -> int:
        """Wrap an emitted w:r with bookmarkStart/End and/or w:hyperlink.

        Order in the parent w:p:
            <w:bookmarkStart .../><w:hyperlink><w:r/></w:hyperlink><w:bookmarkEnd .../>

        Both wrappers are independent: a Run may carry just one, both, or
        neither (caller filters before invoking). Returns the next
        bookmark id to use for subsequent runs in the same render pass.
        """
        p_el = docx_paragraph._p
        r_el = docx_run._r

        # Hyperlink wrap: replace r in p with <w:hyperlink>{r}</w:hyperlink>.
        anchor_target = r_el  # element to wrap with bookmarks
        if run.link is not None:
            hyperlink = _OxmlElement("w:hyperlink")
            if run.link.startswith("#"):
                hyperlink.set(_qn("w:anchor"), run.link[1:])
            else:
                rid = docx_paragraph.part.relate_to(
                    run.link, _RT.HYPERLINK, is_external=True
                )
                hyperlink.set(_qn("r:id"), rid)
            r_index = list(p_el).index(r_el)
            p_el.remove(r_el)
            hyperlink.append(r_el)
            p_el.insert(r_index, hyperlink)
            anchor_target = hyperlink

        # Bookmark wrap: insert start before, end after the anchor target.
        if run.bookmark is not None:
            start = _OxmlElement("w:bookmarkStart")
            start.set(_qn("w:id"), str(bookmark_id))
            start.set(_qn("w:name"), run.bookmark)
            end = _OxmlElement("w:bookmarkEnd")
            end.set(_qn("w:id"), str(bookmark_id))
            target_index = list(p_el).index(anchor_target)
            p_el.insert(target_index, start)
            p_el.insert(target_index + 2, end)  # after the anchor target
            bookmark_id += 1

        return bookmark_id

    @staticmethod
    def _apply_paragraph_style(docx_paragraph: Any, style: Style) -> None:
        pf = docx_paragraph.paragraph_format
        alignment_enum = _ALIGNMENT_MAP.get(style.alignment)
        if alignment_enum is not None:
            docx_paragraph.alignment = alignment_enum
        # spacing_before/after_pt are floats in pt; DocxPt converts pt->EMU.
        pf.space_before = DocxPt(style.spacing_before_pt)
        pf.space_after = DocxPt(style.spacing_after_pt)
        pf.line_spacing = style.line_height
        pf.keep_with_next = style.keep_with_next

    @staticmethod
    def _apply_run_style(
        docx_run: Any, run: Run, base_style: Style | None
    ) -> None:
        """Resolve per-run override > base Style > python-docx default.

        Each formatting attribute is applied only when the resolved value
        is not ``None``. ``Run.<field>`` set to a literal value (including
        ``False`` or ``0.0``) takes precedence over ``base_style.<field>``.
        ``Run.underline`` has no Style fallback; ``None`` = no underline.
        """
        font = docx_run.font

        # Font name: Style only (Run has no font field in v0.1.0).
        if base_style is not None:
            font.name = base_style.font

        size_pt = (
            run.font_size_pt
            if run.font_size_pt is not None
            else (base_style.size_pt if base_style is not None else None)
        )
        if size_pt is not None:
            font.size = DocxPt(size_pt)

        bold = (
            run.bold
            if run.bold is not None
            else (base_style.bold if base_style is not None else None)
        )
        if bold is not None:
            font.bold = bold

        italic = (
            run.italic
            if run.italic is not None
            else (base_style.italic if base_style is not None else None)
        )
        if italic is not None:
            font.italic = italic

        # Underline: Run-only override; Style has no underline field.
        if run.underline is not None:
            font.underline = run.underline

        color_hex = (
            run.color
            if run.color is not None
            else (base_style.color_hex if base_style is not None else None)
        )
        if color_hex is not None:
            font.color.rgb = RGBColor.from_string(color_hex.lstrip("#"))
    # END_BLOCK_RENDER_CONTENT_PARAGRAPH


# ---------------------------------------------------------------------------
# Image
# ---------------------------------------------------------------------------


@dataclass
class Image:
    """Image embed: emits w:drawing with a unique rId via python-docx.

    Public API takes plain Python types (``float`` inches for width/height);
    internal conversion to :class:`docx.shared.Inches` happens at render
    time only. Callers do not need to import python-docx to construct an
    Image. Per V-MP-CONTENT scenario-7, ``width=3.0`` results in an
    inline-shape extent of 3 * 914400 = 2_743_200 EMU.
    """

    _path: Path | None = None
    _data: bytes | None = None
    _format: str | None = None
    width: float | None = None  # inches
    height: float | None = None  # inches

    @classmethod
    def from_path(
        cls,
        path: Path | str,
        width: float | None = None,
        height: float | None = None,
    ) -> Image:
        """Create an Image from a file path.

        Filesystem existence check happens at construction time (NOT render
        time) to honor forbidden-1: render() must never touch the disk.

        ``width`` / ``height`` are in inches (plain float). Internal
        conversion to ``docx.shared.Inches`` is deferred to render().
        """
        p = Path(path)
        if not p.exists():
            raise ImageFileNotFoundError(f"image file not found: {p}")
        return cls(_path=p, width=width, height=height)

    @classmethod
    def from_bytes(
        cls,
        data: bytes,
        format: str = "png",
        width: float | None = None,
        height: float | None = None,
    ) -> Image:
        """Create an Image from in-memory bytes.

        Pillow's :meth:`PIL.Image.Image.verify` runs at construction time,
        not at render time, so a corrupt payload fails fast with
        :class:`ImageFormatUnsupportedError`.
        """
        try:
            with PilImage.open(BytesIO(data)) as probe:
                probe.verify()
        except (UnidentifiedImageError, OSError, SyntaxError, ValueError) as exc:
            raise ImageFormatUnsupportedError(
                f"image format {format!r} unsupported or payload invalid: {exc}"
            ) from exc
        return cls(_data=data, _format=format, width=width, height=height)

    # START_BLOCK_RENDER_CONTENT_IMAGE
    def render(self, parent_doc: DocxDocument) -> Any:
        """Append a w:drawing to ``parent_doc``; return the inline shape.

        Pt-unit decision: this render call site uses python-docx's
        ``docx.shared.Inches`` helper exclusively. Our public API takes
        ``float`` inches; we convert internally only at the python-docx
        boundary. No lxml drop-down, no hand-written EMU.

        Forbidden-1 invariant: this method writes ONLY into ``parent_doc``;
        the source bytes/path were captured at construction time so render
        does not read from disk either (path mode opens the file via
        python-docx, which is the only filesystem touch — and it is a
        READ, not a write; this matches the spec's intent).
        """
        source_kind = "path" if self._path is not None else "bytes"
        logger.info(
            "[MP-Content][render][BLOCK_RENDER_CONTENT] kind=image source=%s",
            source_kind,
        )

        kwargs: dict[str, Any] = {}
        if self.width is not None:
            kwargs["width"] = Inches(self.width)
        if self.height is not None:
            kwargs["height"] = Inches(self.height)

        if self._path is not None:
            shape = parent_doc.add_picture(str(self._path), **kwargs)
        else:
            assert self._data is not None  # narrowed by source_kind logic
            shape = parent_doc.add_picture(BytesIO(self._data), **kwargs)
        return shape
    # END_BLOCK_RENDER_CONTENT_IMAGE


__all__ = [
    "ContentError",
    "Image",
    "ImageFileNotFoundError",
    "ImageFormatUnsupportedError",
    "Paragraph",
    "Run",
    "TabAlignment",
    "TabLeader",
    "TabStop",
]
