# FILE: src/mint_python/core/callout.py
# VERSION: 0.1.0
# START_MODULE_CONTRACT
#   PURPOSE: Callout block — visually distinct info / warning / code box
#     rendered as a single-cell table with kind-specific shading + border.
#     Sibling block of Paragraph/Table/Image/Chart/List inside a Section.
#   SCOPE: Public surface = Callout, CalloutKind, CalloutError. render()
#     mutates a python-docx Document; never writes to disk.
#   DEPENDS: python-docx (1.1.x).
#   LINKS: docs/development-plan.xml#MP-CALLOUT,
#     docs/verification-plan.xml#V-MP-CALLOUT,
#     docs/knowledge-graph.xml#MP-CALLOUT
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   CalloutKind     - StrEnum: INFO | WARNING | CODE
#   Callout         - @dataclass; text + kind + optional title; .render
#   CalloutError    - base error
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: v0.1.0 — initial implementation. Single-cell table render
#     with kind-specific shading + left-accent border. Code kind switches
#     to monospace (Courier New).
# END_CHANGE_SUMMARY

from __future__ import annotations

import logging
from dataclasses import dataclass
from enum import StrEnum
from typing import Any

from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger("mint_python.core.callout")


class CalloutError(Exception):
    """Base for MP-CALLOUT errors."""


class CalloutKind(StrEnum):
    """Visual kind of a Callout block.

    INFO:    blue accent / light-blue background.
    WARNING: amber accent / light-amber background.
    CODE:    grey accent / off-white background, Courier New monospace.
    """

    INFO = "info"
    WARNING = "warning"
    CODE = "code"


# kind -> (border_hex, fill_hex, monospace, default_title)
_KIND_TOKENS: dict[CalloutKind, tuple[str, str, bool, str]] = {
    CalloutKind.INFO: ("2E75B6", "EBF5FB", False, "Info"),
    CalloutKind.WARNING: ("E8A838", "FFF8E1", False, "Warning"),
    CalloutKind.CODE: ("DDDDDD", "F5F5F5", True, ""),
}


@dataclass
class Callout:
    """Styled boxed block — info, warning, or code listing.

    Construction::

        Callout("Save your work before continuing.", kind=CalloutKind.WARNING)
        Callout("def foo(): ...", kind=CalloutKind.CODE)
        Callout("Background note.", kind=CalloutKind.INFO, title="Note")

    Render strategy: emit a 1x1 python-docx Table whose cell carries:
      - direct w:shd shading with the kind's fill color
      - thick left border + thin top/right/bottom borders in the kind's
        accent color (creates the "card with left rail" look)
      - inner cell margins for breathing room
      - optional bold title paragraph
      - body paragraph(s); CODE kind switches to Courier New
    """

    text: str
    kind: CalloutKind = CalloutKind.INFO
    title: str | None = None

    # START_BLOCK_RENDER_CALLOUT
    def render(self, parent_doc: DocxDocument) -> Any:
        """Append a callout box to ``parent_doc``; return the python-docx Table.

        Forbidden-1 invariant: render() only mutates ``parent_doc``; no
        filesystem writes.
        """
        border_hex, fill_hex, monospace, default_title = _KIND_TOKENS[self.kind]

        logger.info(
            "[MP-Callout][render][BLOCK_RENDER_CALLOUT] kind=%s has_title=%s",
            self.kind.value,
            self.title is not None or bool(default_title),
        )

        table = parent_doc.add_table(rows=1, cols=1)
        table.autofit = False
        cell = table.cell(0, 0)

        _set_cell_shading(cell, fill_hex)
        _set_cell_borders(cell, border_hex)
        _set_cell_margins(cell, top=120, bottom=120, left=200, right=200)

        # Drop the python-docx-seeded empty paragraph; we manage paragraphs
        # ourselves so title/body styling is predictable.
        cell._tc.remove(cell.paragraphs[0]._p)

        title_text = self.title if self.title is not None else default_title
        if title_text:
            tp = cell.add_paragraph()
            tr = tp.add_run(title_text)
            tr.bold = True
            tr.font.color.rgb = RGBColor.from_string(border_hex)

        bp = cell.add_paragraph()
        br = bp.add_run(self.text)
        if monospace:
            br.font.name = "Courier New"
            br.font.size = Pt(10)

        # Vertically center cell content for short callouts.
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        return table
    # END_BLOCK_RENDER_CALLOUT


# ---------------------------------------------------------------------------
# Cell styling helpers (lxml drop-down; python-docx doesn't expose w:shd /
# tcBorders / tcMar at the cell level)
# ---------------------------------------------------------------------------


def _set_cell_shading(cell: Any, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_borders(cell: Any, accent_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    # Thick left border = the kind's accent rail.
    for side, size, color in (
        ("left", "24", accent_hex),
        ("top", "4", accent_hex),
        ("right", "4", accent_hex),
        ("bottom", "4", accent_hex),
    ):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)  # eighths of a point: 24 = 3pt, 4 = 0.5pt
        b.set(qn("w:color"), color)
        b.set(qn("w:space"), "0")
        borders.append(b)
    tc_pr.append(borders)


def _set_cell_margins(
    cell: Any, *, top: int, bottom: int, left: int, right: int
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (
        ("top", top), ("bottom", bottom), ("left", left), ("right", right)
    ):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr.append(tc_mar)


# Suppress unused-import warning for Inches — kept for potential future use
# in column-width sizing without re-importing across the module.
_ = Inches


__all__ = [
    "Callout",
    "CalloutError",
    "CalloutKind",
]
