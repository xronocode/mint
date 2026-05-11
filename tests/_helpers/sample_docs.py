# FILE: tests/_helpers/sample_docs.py
"""sample_docs — controller pre-flight fixture for Phase-16 Wave-16-1.

Shared helpers for the five W1 workers (MP-MCP-VALIDATE, MP-MCP-FIX,
MP-FINGERPRINT, MP-EXTRACT, MP-DOC-BUNDLE) so each can build, validate,
fingerprint, extract, or re-render a small bank of canonical docx fixtures
without redefining the construction logic locally.

Usage:
    from tests._helpers.sample_docs import (
        minimal_docx_bytes,
        valid_memo_docx_bytes,
        broken_styles_docx_bytes,
        no_styles_xml_docx_bytes,
        not_a_zip_bytes,
    )

    # Common pattern in a test:
    def test_validate_on_valid_doc(tmp_path):
        path = tmp_path / "doc.docx"
        path.write_bytes(valid_memo_docx_bytes())
        report = mint_python.validate.validate(str(path))
        assert report.passed

The factory functions are deterministic and memoized per process — the
bytes returned by two calls of the same factory in the same test session
are byte-identical, so fingerprint / hash equality assertions are stable.

Constraint-8: this helper imports only stdlib + python-docx + lxml; it
MUST NOT import from src/mint/. Workers under src/mint_python/ can safely
import it via tests._helpers.sample_docs.
"""

from __future__ import annotations

import io
import zipfile
from functools import lru_cache
from pathlib import Path

_FIXTURES_DIR = Path(__file__).parent.parent / "fixtures"


# ---------------------------------------------------------------------------
# Canonical fixture bytes — each factory returns a fresh bytes object so
# tests that mutate the path don't leak across the suite.
# ---------------------------------------------------------------------------


@lru_cache(maxsize=1)
def _minimal_valid_template() -> bytes:
    """Read tests/fixtures/minimal_valid.docx once and cache the bytes."""
    path = _FIXTURES_DIR / "minimal_valid.docx"
    return path.read_bytes()


def minimal_docx_bytes() -> bytes:
    """Minimal valid .docx — no GRACE, no styles.xml customizations.

    Useful as a 'foreign docx' baseline for validate / fingerprint / extract
    tools that must handle docs they didn't generate.
    """
    return _minimal_valid_template()


@lru_cache(maxsize=1)
def _valid_memo_template() -> bytes:
    """Build a small but realistic memo docx via python-docx, cache bytes."""
    from docx import Document

    doc = Document()
    doc.add_heading("Memo: Phase-16 Pre-flight", level=1)
    doc.add_paragraph("From: controller")
    doc.add_paragraph("To: workers")
    doc.add_paragraph("Date: 2026-05-11")
    doc.add_paragraph("Subject: Wave-16-1 sample fixture")
    doc.add_paragraph("")
    doc.add_paragraph(
        "This memo exists so the five Wave-16-1 workers share an "
        "identical structural baseline for round-trip / parity tests."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def valid_memo_docx_bytes() -> bytes:
    """A small memo docx built with python-docx; valid OOXML structure."""
    return _valid_memo_template()


@lru_cache(maxsize=1)
def _no_styles_xml_template() -> bytes:
    """Construct a docx zip whose word/styles.xml entry is absent.

    Used to drive MissingStyleXmlError paths in MP-FINGERPRINT.
    """
    src = _valid_memo_template()
    buf = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(src), "r") as src_zf,
        zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as dst_zf,
    ):
        for info in src_zf.infolist():
            if info.filename == "word/styles.xml":
                continue
            dst_zf.writestr(info, src_zf.read(info.filename))
    return buf.getvalue()


def no_styles_xml_docx_bytes() -> bytes:
    """A docx zip without word/styles.xml.

    Drives V-MP-FINGERPRINT scenario-3 (MissingStyleXmlError).
    """
    return _no_styles_xml_template()


@lru_cache(maxsize=1)
def _broken_styles_template() -> bytes:
    """Construct a docx whose word/styles.xml is unparseable XML."""
    src = _valid_memo_template()
    buf = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(src), "r") as src_zf,
        zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as dst_zf,
    ):
        for info in src_zf.infolist():
            if info.filename == "word/styles.xml":
                dst_zf.writestr(
                    info.filename,
                    b"<not-a-real-styles-xml>broken<<<>>>>",
                )
            else:
                dst_zf.writestr(info, src_zf.read(info.filename))
    return buf.getvalue()


def broken_styles_docx_bytes() -> bytes:
    """A docx whose word/styles.xml is malformed.

    Drives V-MP-MCP-VALIDATE scenario-2 (broken-doc validation report)
    and similar broken-structure paths in other tools.
    """
    return _broken_styles_template()


def not_a_zip_bytes() -> bytes:
    """Bytes that fail BadZipFile when opened with zipfile.

    Drives INVALID_DOCUMENT branches for tools that open the path via
    zipfile (validate / fix / fingerprint / extract / edit).
    """
    return b"This is not a zip archive. Just plain text."


# ---------------------------------------------------------------------------
# Path-traversal sentinel — shared rejection target for safe_doc guards.
# ---------------------------------------------------------------------------


def path_traversal_sentinel() -> str:
    """A path-traversal pattern that safe_doc MUST reject.

    Returns a fixed string callers can use to assert INVALID_DOCUMENT /
    traversal-rejected scenarios without each test reinventing the form.
    """
    return "../../etc/passwd"


# ---------------------------------------------------------------------------
# Helpers for writing fixture bytes to tmp_path in the most common pattern.
# ---------------------------------------------------------------------------


def write_to_tmp(tmp_path: Path, name: str, content: bytes) -> Path:
    """Convenience: write content under tmp_path / name and return the path.

    Standardizes the per-test pattern so callers don't repeat the same
    three-line write-bytes / return-path dance.
    """
    target = tmp_path / name
    target.write_bytes(content)
    return target
