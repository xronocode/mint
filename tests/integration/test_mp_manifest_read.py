# FILE: tests/integration/test_mp_manifest_read.py
# VERSION: 0.1.0
# START_MODULE_CONTRACT
#   PURPOSE: Phase-15 Wave-15-1 — V-MP-MANIFEST-READ scenarios 1-6 +
#     VF-018 invariant tests covering MP-MANIFEST-READ
#     (mint_read_grace_manifest tool). Round-trip from create_document,
#     multi-manifest most-recent-wins selection, structured tool errors
#     on no-manifest / invalid-doc / malformed-XML, log markers, and
#     the read-only / zip-mode-pinned invariants.
#   SCOPE: Integration tests against the live MP-GRACE injection +
#     MP-MANIFEST-READ tool + FastMCP server registration. Uses the
#     controller-provided fixtures (zip_byte_snapshot, caplog_at_info,
#     tmp_docx_path) consumed via tests/integration/conftest.py.
#   DEPENDS: pytest, pytest-asyncio, mint_python.mcp.document
#     (_run_pipeline + create_document for fixture docx production),
#     mint_python.mcp.manifest (the unit-under-test),
#     mint_python.grace (bootstrap re-injection for scenario-2 / 5),
#     tests._helpers.fake_mcp_context.FakeMCPContext.
#   LINKS: docs/verification-plan.xml#V-MP-MANIFEST-READ,
#     docs/verification-plan.xml#VF-018
# END_MODULE_CONTRACT
from __future__ import annotations

import logging
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest

from mint_python.grace import GRACE_NS
from mint_python.mcp import manifest as manifest_module
from mint_python.mcp.document import _run_pipeline
from mint_python.mcp.manifest import (
    CANONICAL_KEYS,
    InvalidDocument,
    ManifestNotFound,
    ManifestParseError,
    mint_read_grace_manifest,
)
from tests._helpers.fake_mcp_context import FakeMCPContext

MEMO_FIXTURES = Path(__file__).parent.parent / "fixtures" / "memo_poc"


def _read_intent_full() -> str:
    return (MEMO_FIXTURES / "intent_full.txt").read_text(encoding="utf-8")


@pytest.fixture(autouse=True)
def _isolate_output_dir(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """Hermetic output dir per test — same pattern as test_mp_doc_generic.py."""
    monkeypatch.setenv("MINT_MEMO_DIR", str(tmp_path / "doc_out"))


@pytest.fixture(autouse=True)
def _isolate_templates_dir(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    """tmp_path-isolated templates/ snapshot — insulates from local
    update_template-authored siblings sitting in repo's templates/."""
    fixtures = tmp_path / "templates"
    fixtures.mkdir()
    repo_templates = Path(__file__).parent.parent.parent / "templates"
    for name in ("memo.yaml",):
        (fixtures / name).write_text(
            (repo_templates / name).read_text(encoding="utf-8"),
            encoding="utf-8",
        )
    from mint_python.mcp import document as document_module
    from mint_python.templates import registry as reg_module

    monkeypatch.setattr(document_module, "_TEMPLATES_DIR", fixtures)
    monkeypatch.setattr(reg_module, "_TEMPLATES_DIR", fixtures)
    from mint_python.templates.registry import reset_default_registry

    reset_default_registry()
    yield fixtures
    reset_default_registry()


async def _create_memo_docx() -> Path:
    """Helper: produce a fresh GRACE-injected memo docx via the live
    pipeline. Returns the saved path. Heuristic extraction fills all
    required fields from intent_full.txt so no elicit calls fire."""
    intent = _read_intent_full()
    ctx = FakeMCPContext(answers={})
    result = await _run_pipeline(
        intent=intent, doc_type="memo", source_md=intent, ctx=ctx
    )
    assert result["status"] == "complete"
    return Path(result["path"])


def _build_manifest_xml(
    *,
    audit_id: str,
    timestamp: str,
    template: str = "memo.yaml",
    template_version: str = "1.0",
) -> str:
    """Construct a GRACE-shaped manifest XML body used by scenarios that
    need to inject a second manifest part directly (scenario-2 + 5)."""
    root = ET.Element(f"{{{GRACE_NS}}}manifest")
    root.set("xmlns:grace", GRACE_NS)
    struct = ET.SubElement(root, f"{{{GRACE_NS}}}documentStructure")
    fmt = ET.SubElement(struct, f"{{{GRACE_NS}}}format")
    fmt.text = "docx"
    fp = ET.SubElement(root, f"{{{GRACE_NS}}}fingerprint")
    fp.text = "fingerprint-not-applicable-for-fixture"
    instr = ET.SubElement(root, f"{{{GRACE_NS}}}instructions")
    rules = [
        f"audit_id={audit_id}",
        "generated_by=MP-DOC-GENERIC",
        f"generated_at={timestamp}",
        "fields_elicited=(none)",
        f"template={template}",
        f"template_version={template_version}",
        "preset=klawd",
    ]
    for i, rule_text in enumerate(rules):
        rule_el = ET.SubElement(instr, f"{{{GRACE_NS}}}rule")
        rule_el.set("index", str(i))
        rule_el.text = rule_text
    return ET.tostring(root, encoding="unicode", xml_declaration=True)


def _inject_extra_manifest_part(docx: Path, part_name: str, xml_text: str) -> None:
    """Append `grace/<part_name>` to an existing docx zip.

    Used ONLY by scenario fixtures that need to craft multi-manifest /
    malformed-manifest test artifacts. The unit-under-test never opens
    a docx in append mode (VF-018 forbidden-1).
    """
    with zipfile.ZipFile(docx, mode="a") as zf:
        zf.writestr(part_name, xml_text)


def _replace_manifest_part(docx: Path, replacement_xml: str) -> str:
    """Rewrite the existing grace/manifest_*.xml part with `replacement_xml`.

    Returns the part name that was replaced. Implemented by extracting +
    re-zipping (zipfile can't atomically replace an entry in place).
    Used by scenario-5 (malformed-manifest)."""
    tmp = docx.with_suffix(".tmp.docx")
    replaced_name = ""
    with (
        zipfile.ZipFile(docx, mode="r") as zf_in,
        zipfile.ZipFile(tmp, mode="w", compression=zipfile.ZIP_DEFLATED) as zf_out,
    ):
        for info in zf_in.infolist():
            data = zf_in.read(info.filename)
            if (
                info.filename.startswith("grace/")
                and info.filename.endswith(".xml")
                and not replaced_name
            ):
                replaced_name = info.filename
                zf_out.writestr(info.filename, replacement_xml)
            else:
                zf_out.writestr(info, data)
    shutil.move(str(tmp), str(docx))
    return replaced_name


# --------------------------------------------------------------------------- #
# Scenario-1 — round-trip: create_document → mint_read_grace_manifest
# --------------------------------------------------------------------------- #


@pytest.mark.asyncio
async def test_scenario_1_round_trip_create_then_read_manifest() -> None:
    """create_document(doc_type='memo') → mint_read_grace_manifest(path)
    returns the canonical 10-key dict with values matching what was
    injected (audit_id, fields_elicited list, template, template_version,
    timestamp ISO 8601, namespace)."""
    docx = await _create_memo_docx()
    ctx = FakeMCPContext(answers={})

    result = await mint_read_grace_manifest(str(docx), ctx=ctx)

    # Canonical key set — exactly the 12 keys (Phase-17 W17-0 added
    # preset_version + lang for extended provenance), no extras / missing.
    assert set(result.keys()) == set(CANONICAL_KEYS)
    # audit_id round-trip — UUID4 string injected during create_document.
    assert isinstance(result["audit_id"], str) and len(result["audit_id"]) >= 32
    assert result["template"] == "memo.yaml"
    assert result["template_version"] == "1.0"
    assert result["namespace"] == "urn:mint:grace:2026:manifest"
    assert isinstance(result["fields_elicited"], list)
    # generated_at is recorded via datetime.now(tz=UTC).isoformat() — ISO 8601.
    assert result["timestamp"].startswith("20")  # 20XX-...
    assert "T" in result["timestamp"]
    # Fingerprint is the SHA-256 hex of the pre-injection docx bytes.
    assert isinstance(result["fingerprint"], str) and len(result["fingerprint"]) == 64
    # Raw instructions list preserves the full key=value lines.
    assert any(line.startswith("audit_id=") for line in result["instructions"])
    assert any(line.startswith("preset=klawd") for line in result["instructions"])
    # model_identity / template_author absent for the standard memo path.
    assert result["model_identity"] is None
    assert result["template_author"] is None


# --------------------------------------------------------------------------- #
# Scenario-2 — multiple manifest parts: most-recent-wins by timestamp
# --------------------------------------------------------------------------- #


@pytest.mark.asyncio
async def test_scenario_2_multiple_manifest_parts_returns_most_recent() -> None:
    """When N>1 grace/manifest_*.xml parts exist, the tool returns the
    one with the highest parsed `generated_at` timestamp — NOT the first
    in namelist order. Inverts namelist order vs timestamp order to make
    the assertion meaningful."""
    docx = await _create_memo_docx()
    # The pipeline already injected ONE manifest. We add two more, named
    # in INVERTED chronology: `manifest_v1.xml` carries the EARLIER
    # timestamp; `manifest_v2.xml` carries the LATER timestamp. Both are
    # added AFTER the auto-injected part so the namelist order is
    # [auto, v1, v2]. The auto-injected one has its own timestamp from
    # _run_pipeline; pick T_LATER so v2 (sorted after auto by timestamp
    # too) is the unambiguous winner.
    early = "2020-01-01T00:00:00+00:00"
    late = "2099-12-31T23:59:59+00:00"
    early_audit = "11111111-1111-4111-8111-111111111111"
    late_audit = "22222222-2222-4222-8222-222222222222"
    _inject_extra_manifest_part(
        docx,
        "grace/manifest_v1.xml",
        _build_manifest_xml(audit_id=early_audit, timestamp=early),
    )
    _inject_extra_manifest_part(
        docx,
        "grace/manifest_v2.xml",
        _build_manifest_xml(audit_id=late_audit, timestamp=late),
    )

    ctx = FakeMCPContext(answers={})
    result = await mint_read_grace_manifest(str(docx), ctx=ctx)

    # The latest (v2) entry wins — assert by audit_id and timestamp.
    assert result["audit_id"] == late_audit
    assert result["timestamp"] == late


# --------------------------------------------------------------------------- #
# Scenario-3 — no GRACE manifest part → MANIFEST_NOT_FOUND
# --------------------------------------------------------------------------- #


@pytest.mark.asyncio
async def test_scenario_3_no_grace_part_raises_manifest_not_found(
    tmp_path: Path,
) -> None:
    """Docx produced WITHOUT GRACE injection raises MANIFEST_NOT_FOUND
    structured error (not a stack trace). Message names document_path."""
    # Produce a minimal docx via python-docx — bypasses MINT entirely so
    # no GRACE manifest is injected. python-docx is a soft dep; if not
    # installed, fall back to a hand-rolled minimal OOXML zip.
    docx_path = tmp_path / "no_grace.docx"
    try:
        from docx import Document as PyDocxDocument
    except ImportError:  # pragma: no cover — python-docx is in dev deps
        # Hand-roll a minimal OOXML zip with no grace/ entries.
        with zipfile.ZipFile(docx_path, mode="w") as zf:
            zf.writestr(
                "[Content_Types].xml",
                '<?xml version="1.0"?><Types xmlns='
                '"http://schemas.openxmlformats.org/package/2006/content-types"/>',
            )
            zf.writestr(
                "word/document.xml",
                '<?xml version="1.0"?><document/>',
            )
    else:
        d = PyDocxDocument()
        d.add_paragraph("no GRACE here")
        d.save(str(docx_path))

    ctx = FakeMCPContext(answers={})
    with pytest.raises(ManifestNotFound, match="MANIFEST_NOT_FOUND") as exc_info:
        await mint_read_grace_manifest(str(docx_path), ctx=ctx)
    # document_path surfaces in the message so the connected model can
    # echo it back to the user.
    assert str(docx_path) in str(exc_info.value) or docx_path.name in str(exc_info.value)


# --------------------------------------------------------------------------- #
# Scenario-4 — invalid document / path traversal
# --------------------------------------------------------------------------- #


@pytest.mark.asyncio
async def test_scenario_4_a_invalid_zip_raises_invalid_document(
    tmp_path: Path,
) -> None:
    """A non-zip file at the requested path raises INVALID_DOCUMENT (NOT
    a BadZipFile traceback)."""
    bogus = tmp_path / "not_a_zip.docx"
    bogus.write_bytes(b"this is plain text, not a zip archive")

    ctx = FakeMCPContext(answers={})
    with pytest.raises(InvalidDocument, match="INVALID_DOCUMENT"):
        await mint_read_grace_manifest(str(bogus), ctx=ctx)


@pytest.mark.asyncio
async def test_scenario_4_a2_missing_file_raises_invalid_document(
    tmp_path: Path,
) -> None:
    """A path that doesn't exist (or isn't a regular file) raises
    INVALID_DOCUMENT with a clear message — distinct from BadZipFile and
    from path traversal."""
    missing = tmp_path / "does_not_exist.docx"
    ctx = FakeMCPContext(answers={})
    with pytest.raises(InvalidDocument, match="INVALID_DOCUMENT"):
        await mint_read_grace_manifest(str(missing), ctx=ctx)


@pytest.mark.asyncio
async def test_scenario_4_b_path_traversal_rejected(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """`../../etc/passwd`-shaped paths are rejected by safe_doc BEFORE any
    zipfile open. Verified via a ZipFile.__init__ sentinel that records
    every call — the sentinel must NOT trip on the traversal path."""
    zip_calls: list[tuple] = []
    real_init = zipfile.ZipFile.__init__

    def _sentinel_init(self, file, *args, **kwargs):  # type: ignore[no-untyped-def]
        zip_calls.append((file, args, kwargs))
        return real_init(self, file, *args, **kwargs)

    monkeypatch.setattr(zipfile.ZipFile, "__init__", _sentinel_init)

    ctx = FakeMCPContext(answers={})
    with pytest.raises(InvalidDocument, match="INVALID_DOCUMENT"):
        await mint_read_grace_manifest("../../etc/passwd", ctx=ctx)

    # safe_doc must reject before any ZipFile is constructed.
    assert zip_calls == [], (
        f"path traversal opened a zip before rejection: {zip_calls!r}"
    )


# --------------------------------------------------------------------------- #
# Scenario-5 — malformed manifest XML → MANIFEST_PARSE_ERROR
# --------------------------------------------------------------------------- #


@pytest.mark.asyncio
async def test_scenario_5_malformed_manifest_xml_raises_parse_error() -> None:
    """A docx whose grace/manifest_*.xml is truncated raises
    MANIFEST_PARSE_ERROR. Both xml_part_name AND document_path appear in
    the surfaced message."""
    docx = await _create_memo_docx()
    truncated_xml = "<manifest><audit_id>foo"  # mismatched / unterminated

    replaced_part = _replace_manifest_part(docx, truncated_xml)
    assert replaced_part.startswith("grace/manifest_") and replaced_part.endswith(
        ".xml"
    )

    ctx = FakeMCPContext(answers={})
    with pytest.raises(ManifestParseError, match="MANIFEST_PARSE_ERROR") as exc_info:
        await mint_read_grace_manifest(str(docx), ctx=ctx)
    msg = str(exc_info.value)
    assert replaced_part in msg
    assert str(docx) in msg or docx.name in msg


# --------------------------------------------------------------------------- #
# Scenario-6 — log markers on success and not-found
# --------------------------------------------------------------------------- #


@pytest.mark.asyncio
async def test_scenario_6_log_markers_on_success_and_not_found(
    tmp_path: Path, caplog_at_info: pytest.LogCaptureFixture
) -> None:
    """Success path emits BLOCK_READ_MANIFEST with the documented payload
    (document_path + manifest_count + selected_xml_part_name); not-found
    path emits BLOCK_MANIFEST_NOT_FOUND with document_path."""
    # Success.
    docx = await _create_memo_docx()
    caplog_at_info.clear()
    ctx = FakeMCPContext(answers={})
    with caplog_at_info.at_level(
        logging.INFO, logger="mint_python.mcp.manifest"
    ):
        await mint_read_grace_manifest(str(docx), ctx=ctx)
    success_msgs = [
        r.getMessage()
        for r in caplog_at_info.records
        if "BLOCK_READ_MANIFEST" in r.getMessage()
    ]
    assert success_msgs, "BLOCK_READ_MANIFEST log marker missing on success"
    msg = success_msgs[0]
    assert "document_path=" in msg
    assert "manifest_count=" in msg
    assert "selected_xml_part_name=" in msg

    # Not-found.
    bare = tmp_path / "no_grace.docx"
    try:
        from docx import Document as PyDocxDocument
        d = PyDocxDocument()
        d.add_paragraph("plain")
        d.save(str(bare))
    except ImportError:  # pragma: no cover
        with zipfile.ZipFile(bare, mode="w") as zf:
            zf.writestr("[Content_Types].xml", '<?xml version="1.0"?><Types/>')

    caplog_at_info.clear()
    with (
        caplog_at_info.at_level(logging.INFO, logger="mint_python.mcp.manifest"),
        pytest.raises(ManifestNotFound),
    ):
        await mint_read_grace_manifest(str(bare), ctx=ctx)
    notfound_msgs = [
        r.getMessage()
        for r in caplog_at_info.records
        if "BLOCK_MANIFEST_NOT_FOUND" in r.getMessage()
    ]
    assert notfound_msgs, "BLOCK_MANIFEST_NOT_FOUND log marker missing"
    assert "document_path=" in notfound_msgs[0]


# --------------------------------------------------------------------------- #
# VF-018 inv-2 READ-ONLY — sha256 byte equality before/after the read call.
# --------------------------------------------------------------------------- #


@pytest.mark.asyncio
async def test_vf_018_inv_2_read_only_byte_equality(
    tmp_path: Path,
    zip_byte_snapshot,
) -> None:
    """sha256(docx_bytes) before equals sha256(docx_bytes) after for both
    the success path AND the not-found path. Read-only by contract."""
    # Success path.
    docx = await _create_memo_docx()
    snap = zip_byte_snapshot(docx)
    ctx = FakeMCPContext(answers={})
    await mint_read_grace_manifest(str(docx), ctx=ctx)
    snap()  # raises AssertionError on byte drift

    # Not-found path — produce a non-GRACE docx and confirm the failure
    # path doesn't mutate either.
    bare = tmp_path / "no_grace.docx"
    try:
        from docx import Document as PyDocxDocument
        d = PyDocxDocument()
        d.add_paragraph("x")
        d.save(str(bare))
    except ImportError:  # pragma: no cover
        with zipfile.ZipFile(bare, mode="w") as zf:
            zf.writestr("[Content_Types].xml", '<?xml version="1.0"?><Types/>')
    snap2 = zip_byte_snapshot(bare)
    with pytest.raises(ManifestNotFound):
        await mint_read_grace_manifest(str(bare), ctx=ctx)
    snap2()


# --------------------------------------------------------------------------- #
# VF-018 forbidden-1 ZIP-MODE-PINNED — every ZipFile open on the read path
# uses mode='r' (or omits mode, which defaults to 'r').
# --------------------------------------------------------------------------- #


@pytest.mark.asyncio
async def test_vf_018_forbidden_1_zip_open_mode_is_read_only(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Patch zipfile.ZipFile.__init__ with a sentinel that asserts every
    open uses mode='r'. Run a successful read; the sentinel must not
    trip. Catches accidental 'a' / 'w' opens on the tool's call path."""
    real_init = zipfile.ZipFile.__init__
    observed_modes: list[str] = []

    def _mode_pinned_init(self, file, mode="r", *args, **kwargs):  # type: ignore[no-untyped-def]
        observed_modes.append(mode)
        # The contract is "read-only on the call path"; the test fixture
        # itself uses 'a'/'w' to craft fixtures BEFORE the patch is
        # active. Inside the patched window we accept only mode='r'.
        assert mode == "r", (
            f"VF-018 forbidden-1 violation: ZipFile opened with mode={mode!r}"
        )
        return real_init(self, file, mode, *args, **kwargs)

    docx = await _create_memo_docx()  # 'w' opens happen here pre-patch.

    monkeypatch.setattr(zipfile.ZipFile, "__init__", _mode_pinned_init)
    ctx = FakeMCPContext(answers={})
    await mint_read_grace_manifest(str(docx), ctx=ctx)
    # At least one read-mode open occurred (the tool's _read_all_manifests).
    assert observed_modes, "expected at least one ZipFile open during read"
    assert all(m == "r" for m in observed_modes)


# --------------------------------------------------------------------------- #
# Internal helper coverage — exercises the canonicalize / select_most_recent
# branches that the success scenarios above don't reach (model_identity +
# template_author present, missing-timestamp tie-break, empty
# fields_elicited string vs (none) sentinel).
# --------------------------------------------------------------------------- #


def test_canonicalize_surfaces_model_identity_and_template_author() -> None:
    """When the source manifest carries `model_identity=` /
    `template_author=` lines (forward-compat for future MINT additions),
    _canonicalize surfaces them on the canonical dict."""
    from mint_python.grace import GRACEManifest

    m = GRACEManifest(
        document_structure={},
        fingerprint="abc",
        instructions=[
            "audit_id=xyz",
            "model_identity=claude-opus-4-7",
            "template=memo.yaml",
            "template_version=1.1",
            "template_author=mevdokimov",
            "generated_at=2026-01-02T03:04:05+00:00",
            "fields_elicited=sender,body",
            "preset=klawd",
            "loose-line-without-equals",
        ],
        xml_part_name="grace/manifest_x.xml",
    )
    canon = manifest_module._canonicalize(m)
    assert canon["model_identity"] == "claude-opus-4-7"
    assert canon["template_author"] == "mevdokimov"
    assert canon["fields_elicited"] == ["sender", "body"]
    assert canon["timestamp"] == "2026-01-02T03:04:05+00:00"
    # Loose lines are passed through to instructions but ignored by k=v parser.
    assert "loose-line-without-equals" in canon["instructions"]


def test_select_most_recent_handles_missing_timestamp() -> None:
    """A manifest missing `generated_at` sorts as if it had timestamp=""
    — strictly older than any populated entry."""
    from mint_python.grace import GRACEManifest

    older = GRACEManifest(
        instructions=["audit_id=A"],  # no generated_at
        xml_part_name="grace/manifest_a.xml",
    )
    newer = GRACEManifest(
        instructions=["audit_id=B", "generated_at=2030-01-01T00:00:00+00:00"],
        xml_part_name="grace/manifest_b.xml",
    )
    chosen = manifest_module._select_most_recent([older, newer])
    assert chosen.xml_part_name == "grace/manifest_b.xml"
