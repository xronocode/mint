# FILE: tests/unit/test_mp_coverage_gaps.py
"""Targeted tests for branches not exercised by V-MP-* scenario suites.

This module exists solely to drive `src/mint_python/` coverage to 100%.
Each test is a surgical poke at a specific line of error-path or edge-case
code that the V-MP-* scenario suites do not happen to exercise.
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest
from PIL import Image as PILImage

from mint_python.core.content import (
    Image,
    Paragraph,
)
from mint_python.core.document import (
    Document,
    DocumentSaveIOError,
)
from mint_python.core.section import Section
from mint_python.core.style import (
    STYLE_PRESET_INVALID_SCHEMA,
    Style,
    load_preset,
)
from mint_python.core.table import (
    Cell,
    Table,
    TableInvalidDictKeysError,
    TableMarkdownParseError,
    TableRaggedRowsError,
)

# --------------------------------------------------------------------------- #
# MP-CONTENT
# --------------------------------------------------------------------------- #


def test_paragraph_empty_string_seeds_empty_runs():
    """content.py:143 — Paragraph('') yields empty _runs list (no implicit Run)."""
    p = Paragraph("")
    assert p._runs == []


def test_image_with_explicit_height_passes_inches_kwarg(tmp_path: Path):
    """content.py:312 — Image.from_path(height=...) flows through Inches on render."""
    png_path = tmp_path / "img.png"
    PILImage.new("RGB", (4, 4), "blue").save(png_path)
    from docx import Document as DocxDoc

    img = Image.from_path(png_path, width=2.0, height=1.0)
    doc = DocxDoc()
    img.render(doc)
    assert len(doc.inline_shapes) == 1


# --------------------------------------------------------------------------- #
# MP-SECTION
# --------------------------------------------------------------------------- #


def test_section_add_image_fluent_returns_self(tmp_path: Path):
    """section.py:140-141 — add_image appends to _blocks and returns self."""
    png_path = tmp_path / "img.png"
    PILImage.new("RGB", (4, 4), "red").save(png_path)
    img = Image.from_path(png_path)
    s = Section("Test", level=1)
    out = s.add_image(img)
    assert out is s
    assert s._blocks[-1] is img


# --------------------------------------------------------------------------- #
# MP-DOCUMENT
# --------------------------------------------------------------------------- #


def test_save_oserror_path_unlinks_partial_output(monkeypatch, tmp_path: Path):
    """document.py:311-317 — save() OSError unlinks partial; raises DocumentSaveIOError."""
    out = tmp_path / "out.docx"

    def boom(self, path):
        # python-docx Document.save raises; out file may exist mid-save.
        out.write_bytes(b"partial")  # simulate partial write
        raise OSError("simulated disk full")

    from docx.document import Document as DocxDocument

    monkeypatch.setattr(DocxDocument, "save", boom, raising=True)

    doc = Document(format="docx", title="X").with_style_preset("alga_corporate")
    doc.add_section(Section("S", level=1).add_paragraph("hi"))

    with pytest.raises(DocumentSaveIOError, match="save failed"):
        doc.save(out)
    assert not out.exists(), "partial output must be unlinked on OSError"


def test_save_with_cover_logo_image_renders(tmp_path: Path):
    """document.py:416 — add_cover(logo=path) reaches the doc.add_picture branch."""
    logo_path = tmp_path / "logo.png"
    PILImage.new("RGB", (8, 8), "green").save(logo_path)

    out = tmp_path / "out.docx"
    doc = Document(format="docx", title="WithLogo").with_style_preset("alga_corporate")
    doc.add_cover(title="Title", subtitle="Sub", logo=logo_path)
    doc.add_section(Section("Body", level=1).add_paragraph("ok"))
    doc.save(out)

    from docx import Document as DocxDoc

    reopened = DocxDoc(str(out))
    # Cover paragraphs + logo + page break + heading + body paragraph
    assert len(reopened.inline_shapes) >= 1


# --------------------------------------------------------------------------- #
# MP-STYLE — schema validator branches
# --------------------------------------------------------------------------- #


def _write_preset(tmp_path: Path, data: dict, name: str = "p.json") -> Path:
    p = tmp_path / name
    p.write_text(json.dumps(data))
    return p


def _base_preset() -> dict:
    return {
        "$schema": "x",
        "name": "x",
        "version": "1.0",
        "color_palette": {
            "primary": "#0F4C81",
            "secondary": "#5B8DBE",
            "accent": "#FFB400",
            "text": "#1A1A1A",
            "text_muted": "#6E6E6E",
            "background": "#FFFFFF",
            "border": "#D4D4D4",
        },
        "typography": {
            "heading1": {"font": "I", "size_pt": 24, "color": "#0F4C81"},
            "heading2": {"font": "I", "size_pt": 18, "color": "#0F4C81"},
            "heading3": {"font": "I", "size_pt": 14, "color": "#0F4C81"},
            "body": {"font": "I", "size_pt": 11, "color": "#1A1A1A"},
            "table_header": {"font": "I", "size_pt": 11, "color": "#FFFFFF"},
            "caption": {"font": "I", "size_pt": 9, "color": "#6E6E6E"},
        },
        "spacing": {
            "paragraph_default_before_pt": 0,
            "paragraph_default_after_pt": 6,
            "default_line_height": 1.15,
            "table_cell_padding_pt": 4,
        },
    }


def test_invalid_json_in_preset_file_raises(tmp_path: Path):
    """style.py:476-477 — malformed JSON in preset file raises STYLE_PRESET_INVALID_SCHEMA."""
    bad = tmp_path / "bad.json"
    bad.write_text("{this is not valid json")
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match="invalid JSON"):
        load_preset(path=bad)


def test_root_not_dict_raises(tmp_path: Path):
    """style.py:404 — preset root is not a dict."""
    bad = _write_preset(tmp_path, [], name="list.json")  # type: ignore[arg-type]
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match="expected object at root"):
        load_preset(path=bad)


def test_missing_top_level_key_raises(tmp_path: Path):
    """style.py:407 — missing top-level required key (e.g. 'spacing')."""
    d = _base_preset()
    del d["spacing"]
    p = _write_preset(tmp_path, d)
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match="/spacing"):
        load_preset(path=p)


def test_color_palette_not_dict_raises(tmp_path: Path):
    """style.py:263 — color_palette is not a dict."""
    d = _base_preset()
    d["color_palette"] = "not a dict"
    p = _write_preset(tmp_path, d)
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match=r"/color_palette.*expected object"):
        load_preset(path=p)


def test_color_field_neither_hex_nor_token_raises(tmp_path: Path):
    """style.py:289 — color value matches neither #RRGGBB nor @token regex."""
    d = _base_preset()
    d["typography"]["heading1"]["color"] = "blue"  # not hex, not @token
    p = _write_preset(tmp_path, d)
    with pytest.raises(
        STYLE_PRESET_INVALID_SCHEMA,
        match="expected hex #RRGGBB or @palette-token",
    ):
        load_preset(path=p)


def test_stylespec_not_dict_raises(tmp_path: Path):
    """style.py:297 — typography entry value is not a dict."""
    d = _base_preset()
    d["typography"]["heading1"] = "not a dict"
    p = _write_preset(tmp_path, d)
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match="/typography/heading1"):
        load_preset(path=p)


def test_stylespec_missing_required_key_raises(tmp_path: Path):
    """style.py:300 — StyleSpec missing one of font/size_pt/color."""
    d = _base_preset()
    del d["typography"]["heading1"]["font"]
    p = _write_preset(tmp_path, d)
    with pytest.raises(
        STYLE_PRESET_INVALID_SCHEMA, match="/typography/heading1/font"
    ):
        load_preset(path=p)


def test_typography_not_dict_raises(tmp_path: Path):
    """style.py:356 — typography is not a dict."""
    d = _base_preset()
    d["typography"] = "not a dict"
    p = _write_preset(tmp_path, d)
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match=r"/typography.*expected object"):
        load_preset(path=p)


def test_typography_missing_required_key_raises(tmp_path: Path):
    """style.py:359 — typography missing 'caption'."""
    d = _base_preset()
    del d["typography"]["caption"]
    p = _write_preset(tmp_path, d)
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match="/typography/caption"):
        load_preset(path=p)


def test_spacing_not_dict_raises(tmp_path: Path):
    """style.py:370 — spacing is not a dict."""
    d = _base_preset()
    d["spacing"] = "not a dict"
    p = _write_preset(tmp_path, d)
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match=r"/spacing.*expected object"):
        load_preset(path=p)


def test_spacing_missing_required_key_raises(tmp_path: Path):
    """style.py:373 — spacing missing required key."""
    d = _base_preset()
    del d["spacing"]["default_line_height"]
    p = _write_preset(tmp_path, d)
    with pytest.raises(
        STYLE_PRESET_INVALID_SCHEMA, match="/spacing/default_line_height"
    ):
        load_preset(path=p)


def test_version_unparseable_raises(tmp_path: Path):
    """style.py:392 — version not in MAJOR.MINOR shape (e.g. 'v1' or empty)."""
    d = _base_preset()
    d["version"] = "abc"
    p = _write_preset(tmp_path, d)
    with pytest.raises(
        STYLE_PRESET_INVALID_SCHEMA, match="expected SemVer-style"
    ):
        load_preset(path=p)


def test_require_number_nonneg_negative_raises(tmp_path: Path):
    """style.py:235 — _require_number_nonneg rejects negative values."""
    d = _base_preset()
    d["spacing"]["paragraph_default_before_pt"] = -5
    p = _write_preset(tmp_path, d)
    with pytest.raises(
        STYLE_PRESET_INVALID_SCHEMA, match="expected number >= 0"
    ):
        load_preset(path=p)


def test_require_number_nonneg_wrong_type_raises(tmp_path: Path):
    """style.py:233 — _require_number_nonneg rejects non-numeric types."""
    d = _base_preset()
    d["spacing"]["paragraph_default_before_pt"] = "5pt"  # str instead of number
    p = _write_preset(tmp_path, d)
    with pytest.raises(
        STYLE_PRESET_INVALID_SCHEMA, match="/spacing/paragraph_default_before_pt"
    ):
        load_preset(path=p)


def test_require_number_positive_zero_raises(tmp_path: Path):
    """style.py:242 — _require_number_positive rejects 0."""
    d = _base_preset()
    d["spacing"]["default_line_height"] = 0
    p = _write_preset(tmp_path, d)
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match="expected number > 0"):
        load_preset(path=p)


def test_require_number_positive_wrong_type_raises(tmp_path: Path):
    """style.py:233 / 240 — _require_number_positive rejects non-number."""
    d = _base_preset()
    d["typography"]["heading1"]["size_pt"] = None
    p = _write_preset(tmp_path, d)
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match="expected number"):
        load_preset(path=p)


def test_require_bool_wrong_type_raises(tmp_path: Path):
    """style.py:247 — _require_bool rejects non-bool."""
    d = _base_preset()
    d["typography"]["heading1"]["bold"] = "yes"
    p = _write_preset(tmp_path, d)
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match="expected boolean"):
        load_preset(path=p)


def test_require_string_wrong_type_raises(tmp_path: Path):
    """style.py:252 — _require_string rejects non-string."""
    d = _base_preset()
    d["typography"]["heading1"]["font"] = 42
    p = _write_preset(tmp_path, d)
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA, match="expected string"):
        load_preset(path=p)


def test_alignment_out_of_enum_raises(tmp_path: Path):
    """style.py — alignment validator rejects values outside the documented enum."""
    d = _base_preset()
    d["typography"]["body"]["alignment"] = "diagonal"
    p = _write_preset(tmp_path, d)
    with pytest.raises(STYLE_PRESET_INVALID_SCHEMA):
        load_preset(path=p)


# --------------------------------------------------------------------------- #
# MP-TABLE — error paths and helper branches
# --------------------------------------------------------------------------- #


def test_is_numeric_like_rejects_bool():
    """table.py:141 — bool is excluded from numeric detection."""
    from mint_python.core.table import _is_numeric_like

    assert _is_numeric_like(True) is False
    assert _is_numeric_like(False) is False


def test_is_numeric_like_rejects_other_types():
    """table.py:146 — neither int/float/str → False."""
    from mint_python.core.table import _is_numeric_like

    assert _is_numeric_like(None) is False
    assert _is_numeric_like([1, 2]) is False
    assert _is_numeric_like({"a": 1}) is False


def test_format_numeric_int_path():
    """table.py:155-156 — int yields no decimals."""
    from mint_python.core.table import _format_numeric

    assert _format_numeric(1500) == "1,500"


def test_format_numeric_float_path():
    """table.py:157-158 — float yields 2 decimals."""
    from mint_python.core.table import _format_numeric

    assert _format_numeric(1500.5) == "1,500.50"


def test_format_numeric_str_int():
    """table.py:159, 163 — string without dot parses as int."""
    from mint_python.core.table import _format_numeric

    assert _format_numeric("2000") == "2,000"


def test_format_numeric_str_float():
    """table.py:159, 161-162 — string with dot parses as float."""
    from mint_python.core.table import _format_numeric

    assert _format_numeric("2000.5") == "2,000.50"


def test_from_markdown_non_str_raises():
    """table.py:232 — from_markdown rejects non-str input."""
    with pytest.raises(TableMarkdownParseError, match="expected str input"):
        Table.from_markdown(42)  # type: ignore[arg-type]


def test_from_markdown_line_without_pipe_raises():
    """table.py:251 — _split_pipe rejects lines lacking '|'."""
    with pytest.raises(TableMarkdownParseError, match="expected pipe-delimited"):
        Table.from_markdown("no pipes here\n|---|\n| a |")


def test_from_list_of_dicts_row0_not_dict_raises():
    """table.py:303 — first element of rows is not a dict."""
    with pytest.raises(TableInvalidDictKeysError, match="row 0 is not a dict"):
        Table.from_list_of_dicts(["not a dict"])  # type: ignore[list-item]


def test_from_list_of_dicts_subsequent_row_not_dict_raises():
    """table.py:311 — second element is not a dict."""
    with pytest.raises(TableInvalidDictKeysError, match="row 1 is not a dict"):
        Table.from_list_of_dicts([{"a": 1}, "bad"])  # type: ignore[list-item]


def test_from_list_of_dicts_empty_returns_no_header_table():
    """table.py:301 — empty rows yield empty no-header table."""
    t = Table.from_list_of_dicts([])
    assert t._rows == []
    assert t._has_header is False


def test_financial_empty_rows_returns_empty():
    """table.py:343 — Table.financial([]) yields empty via from_list."""
    t = Table.financial([])
    assert t._rows == []


def test_financial_ragged_rows_raises():
    """table.py:350 — financial() validates rectangular shape."""
    with pytest.raises(TableRaggedRowsError, match="row 1 has"):
        Table.financial([["Q", "Rev"], ["Q1"]])


def test_financial_with_cell_input_numeric_promotes(tmp_path: Path):
    """table.py:366-378 — Cell-typed numeric value gets right-aligned + reformatted."""
    rows = [
        [Cell("Quarter"), Cell("Revenue")],
        [Cell("Q1"), Cell("1500")],  # str-numeric inside Cell
    ]
    t = Table.financial(rows)
    # Header row preserved as-is (idx 0 branch).
    assert t._rows[0][0].value == "Quarter"
    # Data row: Cell value normalized + aligned right.
    assert t._rows[1][1].value == "1,500"
    assert t._rows[1][1].align == "right"


def test_financial_with_cell_input_non_numeric_passes_through():
    """table.py — Cell with non-numeric value preserved as-is in financial()."""
    rows = [
        [Cell("Quarter"), Cell("Notes")],
        [Cell("Q1"), Cell("good")],  # non-numeric Cell → unchanged
    ]
    t = Table.financial(rows)
    assert t._rows[1][1].value == "good"
    # Original Cell preserved; align remains None (inherit).
    assert t._rows[1][1].align is None


def test_render_applies_effective_style_font(tmp_path: Path):
    """table.py:493 — render() applies effective_style.font when not None."""
    from docx import Document as DocxDoc

    body = Style(
        font="Inter",
        size_pt=11.0,
        color_hex="#1A1A1A",
    )
    t = Table.from_list([["a", "b"], ["1", "2"]]).apply_style(body)
    doc = DocxDoc()
    t.render(doc)
    # Inspect a data-row run; font.name should reflect the applied style.
    cell = doc.tables[-1].cell(1, 0)
    run = cell.paragraphs[0].runs[0]
    assert run.font.name == "Inter"
