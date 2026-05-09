# FILE: tests/unit/test_mp_section.py
# START_MODULE_CONTRACT
#   PURPOSE: V-MP-SECTION scenarios 1-4 + Phase-8 add_chart positive scenarios
#     — Section fluent surface, render order with heading, str/Paragraph
#     polymorphism on add_paragraph, level-range enforcement, and the Phase-8
#     add_chart positive contract (chart appended to _blocks; Section.render
#     walks Chart.render to produce inline_shape with correct EMU width).
#   SCOPE: Reuses central conftest fixtures (mp_clean_env, caplog_at_info,
#     marker_counter); imports sibling Paragraph + Table + Chart to validate
#     polymorphism + render ordering. No fixture redefinitions.
#   DEPENDS: pytest, mint_python.core.section, mint_python.core.content,
#     mint_python.core.table, mint_python.core.chart, python-docx.
#   LINKS: docs/verification-plan.xml#V-MP-SECTION,
#     docs/development-plan.xml#MP-SECTION
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   test_scenario_1_fluent_returns_self
#   test_scenario_2_render_emits_heading_then_blocks_in_order
#   test_scenario_3_add_paragraph_accepts_str_and_paragraph
#   test_scenario_4_level_out_of_range_raises
#   test_add_chart_appends_chart_and_returns_self - Phase-8 positive (replaces scenario-5)
#   test_section_render_walks_chart_render - Phase-8 positive (replaces scenario-6)
#   test_phase_guard_is_subclass_of_not_implemented_error - PHASE_GUARD class still subclass
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: Wave-8-2 (MP-SECTION): retire scenarios 5/6 (NotImplementedError +
#     BLOCK_PHASE_GUARD trace) — replace with positive scenarios verifying chart
#     appended + Section.render walks Chart.render.
#   PRIOR: Wave-7-3 (MP-SECTION): initial test suite.
# END_CHANGE_SUMMARY
from __future__ import annotations

import pytest
from docx import Document

from mint_python.core.content import Paragraph
from mint_python.core.section import (
    PhaseGuardNotImplementedError,
    Section,
    SectionLevelOutOfRangeError,
)
from mint_python.core.table import Table

# ---------------------------------------------------------------------------
# V-MP-SECTION scenario-1: fluent setters return self
# ---------------------------------------------------------------------------


def test_scenario_1_fluent_returns_self() -> None:
    section = Section("Title", level=1)
    assert section.add_paragraph("text") is section

    table = Table.from_list([["h"], ["a"]], header=True)
    assert section.add_table(table) is section

    # add_image accepts an Image; we don't construct one here to avoid a real
    # file dependency — scenario-1 only contracts "returns self" on
    # add_paragraph (chained with add_table to confirm both are fluent).


# ---------------------------------------------------------------------------
# V-MP-SECTION scenario-2: render emits heading then content blocks in order
# ---------------------------------------------------------------------------


def test_scenario_2_render_emits_heading_then_blocks_in_order() -> None:
    doc = Document()
    section = (
        Section("My Section", level=2)
        .add_paragraph("first")
        .add_paragraph("second")
    )
    before = len(doc.paragraphs)
    section.render(doc)

    new_paragraphs = doc.paragraphs[before:]
    # Expect: 1 heading paragraph + 2 body paragraphs.
    assert len(new_paragraphs) == 3

    heading_para = new_paragraphs[0]
    # python-docx assigns "Heading 2" style for level=2.
    assert heading_para.style.name == "Heading 2"
    assert heading_para.text == "My Section"

    # Body paragraphs preserve order.
    assert new_paragraphs[1].text == "first"
    assert new_paragraphs[2].text == "second"


# ---------------------------------------------------------------------------
# V-MP-SECTION scenario-3: add_paragraph accepts both str and Paragraph
# ---------------------------------------------------------------------------


def test_scenario_3_add_paragraph_accepts_str_and_paragraph() -> None:
    doc = Document()
    explicit = Paragraph("explicit")
    section = (
        Section("Polymorphism", level=3)
        .add_paragraph("shorthand")
        .add_paragraph(explicit)
    )
    before = len(doc.paragraphs)
    section.render(doc)

    new_paragraphs = doc.paragraphs[before:]
    # heading + 2 body paragraphs.
    assert len(new_paragraphs) == 3
    assert new_paragraphs[1].text == "shorthand"
    assert new_paragraphs[2].text == "explicit"


# ---------------------------------------------------------------------------
# V-MP-SECTION scenario-4: level out of range raises
# ---------------------------------------------------------------------------


@pytest.mark.parametrize("bad_level", [0, 7, 10, -1])
def test_scenario_4_level_out_of_range_raises(bad_level: int) -> None:
    with pytest.raises(SectionLevelOutOfRangeError) as exc_info:
        Section("X", level=bad_level)
    assert str(bad_level) in str(exc_info.value)


# ---------------------------------------------------------------------------
# Phase-8 positive add_chart scenarios (replace pre-Wave-8-2 scenarios 5 + 6
# which asserted the BLOCK_PHASE_GUARD stub behavior — now removed).
# ---------------------------------------------------------------------------


def test_add_chart_appends_chart_and_returns_self() -> None:
    """Phase-8: add_chart(chart: Chart) appends to _blocks and returns self."""
    from mint_python.core.chart import Chart

    chart = Chart.bar(["a", "b"], [1, 2], width_inches=3.0)
    s = Section("Title", level=1)
    out = s.add_chart(chart)
    assert out is s
    assert s._blocks[-1] is chart


def test_section_render_walks_chart_render() -> None:
    """Phase-8: Section.render iterates blocks; Chart.render is invoked."""
    from docx import Document as DocxDoc

    from mint_python.core.chart import Chart

    chart = Chart.line(["a", "b", "c"], [1, 2, 3], width_inches=4.0)
    s = Section("X", level=1).add_chart(chart)
    doc = DocxDoc()
    s.render(doc)
    assert len(doc.inline_shapes) == 1
    assert doc.inline_shapes[0].width.emu == round(4.0 * 914400)


def test_phase_guard_is_subclass_of_not_implemented_error() -> None:
    """PhaseGuardNotImplementedError class remains a NotImplementedError subclass.

    Phase-8 keeps the class in MP-SECTION's __all__ for surface compat (sibling
    modules MP-DOCUMENT + MP-CHART still raise their own PHASE_GUARD errors via
    the same-named class). The class is no longer raised from MP-SECTION
    itself after the add_chart unstub.
    """
    assert issubclass(PhaseGuardNotImplementedError, NotImplementedError)
