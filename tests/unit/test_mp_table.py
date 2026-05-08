# FILE: tests/unit/test_mp_table.py
# START_MODULE_CONTRACT
#   PURPOSE: V-MP-TABLE verification — covers scenarios 1-9 + BLOCK_RENDER_TABLE
#     trace assertion for mint_python.core.table (Table, Cell, factories,
#     shapers, render) per docs/verification-plan.xml#V-MP-TABLE.
#   SCOPE: Unit-level tests only. Reuses central fixtures from
#     tests/unit/conftest.py (mp_clean_env autouse, caplog_at_info,
#     marker_counter). Does NOT import from mint_python.core.content
#     (parallel sibling worker — Phase-7 Wave-7-2 race-collide guard).
#   DEPENDS: pytest, mint_python.core.table, mint_python.core.style,
#     python-docx (Document only — direct construction; no MP-DOCUMENT
#     dependency since that wave hasn't shipped yet).
#   LINKS: docs/verification-plan.xml#V-MP-TABLE,
#     docs/development-plan.xml#MP-TABLE
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   test_from_list_*               - V-MP-TABLE scenario-1
#   test_from_markdown_*           - scenario-2 + scenario-8
#   test_from_list_of_dicts_*      - scenario-3 + scenario-9
#   test_financial_*               - scenario-4
#   test_comparison_*              - scenario-5
#   test_apply_style_widths_autofit_* - scenario-6
#   test_ragged_rows_*             - scenario-7
#   test_render_block_marker_*     - BLOCK_RENDER_TABLE trace assertion
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: Wave-7-2 (MP-TABLE): initial test suite for scenarios 1-9.
# END_CHANGE_SUMMARY
from __future__ import annotations

from collections import Counter

import pytest
from docx import Document

from mint_python.core.style import Style
from mint_python.core.table import (
    Cell,
    Table,
    TableInvalidDictKeysError,
    TableMarkdownParseError,
    TableRaggedRowsError,
)
from tests.unit._mp_helpers import extract_marker

# ---------------------------------------------------------------------------
# scenario-1: Table.from_list — basic shape + header semantics
# ---------------------------------------------------------------------------


def test_from_list_renders_two_by_two():
    doc = Document()
    Table.from_list([["a", "b"], ["1", "2"]]).render(doc)
    docx_table = doc.tables[-1]
    assert len(docx_table.rows) == 2
    assert len(docx_table.columns) == 2
    assert docx_table.rows[0].cells[0].text == "a"
    assert docx_table.rows[0].cells[1].text == "b"
    assert docx_table.rows[1].cells[0].text == "1"
    assert docx_table.rows[1].cells[1].text == "2"


def test_from_list_header_row_is_bold_by_default():
    doc = Document()
    Table.from_list([["a", "b"], ["1", "2"]]).render(doc)
    header_run = doc.tables[-1].rows[0].cells[0].paragraphs[0].runs[0]
    body_run = doc.tables[-1].rows[1].cells[0].paragraphs[0].runs[0]
    assert header_run.font.bold is True
    # Body cell should NOT have bold forced by default.
    assert body_run.font.bold is not True


def test_from_list_header_disabled_keeps_first_row_plain():
    doc = Document()
    Table.from_list([["a", "b"], ["1", "2"]], header=False).render(doc)
    first_run = doc.tables[-1].rows[0].cells[0].paragraphs[0].runs[0]
    assert first_run.font.bold is not True


# ---------------------------------------------------------------------------
# scenario-2: Table.from_markdown — equivalent to from_list
# ---------------------------------------------------------------------------


def test_from_markdown_basic_pipe_table_matches_from_list():
    doc = Document()
    Table.from_markdown("| a | b |\n|---|---|\n| 1 | 2 |").render(doc)
    docx_table = doc.tables[-1]
    assert len(docx_table.rows) == 2
    assert len(docx_table.columns) == 2
    assert docx_table.rows[0].cells[0].text == "a"
    assert docx_table.rows[0].cells[1].text == "b"
    assert docx_table.rows[1].cells[0].text == "1"
    assert docx_table.rows[1].cells[1].text == "2"


def test_from_markdown_tolerates_blank_lines_and_extra_dashes():
    text = "\n| col1 | col2 |\n| ---- | ---- |\n| x | y |\n"
    doc = Document()
    Table.from_markdown(text).render(doc)
    docx_table = doc.tables[-1]
    assert docx_table.rows[0].cells[0].text == "col1"
    assert docx_table.rows[1].cells[1].text == "y"


# ---------------------------------------------------------------------------
# scenario-3: Table.from_list_of_dicts — preserves first-dict insertion order
# ---------------------------------------------------------------------------


def test_from_list_of_dicts_renders_three_by_two():
    doc = Document()
    Table.from_list_of_dicts([{"a": 1, "b": 2}, {"a": 3, "b": 4}]).render(doc)
    docx_table = doc.tables[-1]
    # 1 header row + 2 data rows = 3 rows; 2 columns.
    assert len(docx_table.rows) == 3
    assert len(docx_table.columns) == 2
    # Column order matches first dict's insertion order.
    assert docx_table.rows[0].cells[0].text == "a"
    assert docx_table.rows[0].cells[1].text == "b"
    assert docx_table.rows[1].cells[0].text == "1"
    assert docx_table.rows[1].cells[1].text == "2"
    assert docx_table.rows[2].cells[0].text == "3"
    assert docx_table.rows[2].cells[1].text == "4"


def test_from_list_of_dicts_preserves_non_alphabetical_insertion_order():
    # Insertion order: zeta, alpha — NOT alphabetical.
    doc = Document()
    Table.from_list_of_dicts([{"zeta": 1, "alpha": 2}]).render(doc)
    docx_table = doc.tables[-1]
    assert docx_table.rows[0].cells[0].text == "zeta"
    assert docx_table.rows[0].cells[1].text == "alpha"


# ---------------------------------------------------------------------------
# scenario-4: Table.financial — right-align numerics + thousands separator
# ---------------------------------------------------------------------------


def test_financial_formats_int_with_thousands_separator():
    doc = Document()
    Table.financial([["Q1", 1000], ["Q2", 1300.5]]).render(doc)
    docx_table = doc.tables[-1]
    # Header row 0 untouched.
    assert docx_table.rows[0].cells[0].text == "Q1"
    assert docx_table.rows[0].cells[1].text == "1000"
    # Data row int -> "1,300.50" only for floats; int "1000" header stays raw.
    # Row 1 is the only data row: ['Q2', 1300.5] -> col1 = "1,300.50".
    assert docx_table.rows[1].cells[0].text == "Q2"
    assert docx_table.rows[1].cells[1].text == "1,300.50"


def test_financial_int_value_in_data_row_uses_int_format():
    doc = Document()
    # Header + one data row that has an int.
    Table.financial([["Label", "Amount"], ["Sales", 1500]]).render(doc)
    docx_table = doc.tables[-1]
    assert docx_table.rows[1].cells[1].text == "1,500"


def test_financial_right_aligns_numeric_data_cells():
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    Table.financial([["Q1", "Q2"], ["Sales", 1000]]).render(doc)
    docx_table = doc.tables[-1]
    # Data row, col 1 is numeric -> right-aligned.
    para_numeric = docx_table.rows[1].cells[1].paragraphs[0]
    assert para_numeric.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    # Data row, col 0 is non-numeric ("Sales") -> default alignment (None).
    para_label = docx_table.rows[1].cells[0].paragraphs[0]
    assert para_label.paragraph_format.alignment is None


def test_financial_does_not_coerce_non_numeric_cells():
    doc = Document()
    Table.financial([["Q1", "Q2"], ["Sales", "high"]]).render(doc)
    # Non-numeric "high" is preserved verbatim.
    assert doc.tables[-1].rows[1].cells[1].text == "high"


# ---------------------------------------------------------------------------
# scenario-5: Table.comparison — 2-col side-by-side
# ---------------------------------------------------------------------------


def test_comparison_two_column_with_header_labels():
    doc = Document()
    Table.comparison("Q1", "Q2", [["Revenue", 1.0], ["Cost", 0.5]]).render(doc)
    docx_table = doc.tables[-1]
    assert len(docx_table.rows) == 3  # 1 header + 2 data
    assert len(docx_table.columns) == 2
    assert docx_table.rows[0].cells[0].text == "Q1"
    assert docx_table.rows[0].cells[1].text == "Q2"
    assert docx_table.rows[1].cells[0].text == "Revenue"
    assert docx_table.rows[1].cells[1].text == "1.0"
    assert docx_table.rows[2].cells[0].text == "Cost"
    assert docx_table.rows[2].cells[1].text == "0.5"


def test_comparison_rejects_non_two_column_data():
    with pytest.raises(TableRaggedRowsError) as excinfo:
        Table.comparison("L", "R", [["only-one"]])
    assert "row 0" in str(excinfo.value)
    assert "expected 2" in str(excinfo.value)


# ---------------------------------------------------------------------------
# scenario-6: apply_style + set_column_widths + autofit
# ---------------------------------------------------------------------------


def _make_demo_style(font: str = "Roboto", size_pt: float = 10.0) -> Style:
    return Style(font=font, size_pt=size_pt, color_hex="#112233")


def test_apply_style_overrides_cell_font():
    doc = Document()
    style = _make_demo_style(font="Roboto", size_pt=14)
    t = Table.from_list([["a", "b"], ["1", "2"]]).apply_style(style)
    t.render(doc)
    cell_run = doc.tables[-1].rows[1].cells[0].paragraphs[0].runs[0]
    assert cell_run.font.name == "Roboto"
    # python-docx Pt is opaque; compare via .pt attribute.
    assert cell_run.font.size.pt == 14


def test_set_column_widths_applies_inches_per_column():
    from docx.shared import Inches

    doc = Document()
    t = Table.from_list([["a", "b"], ["1", "2"]]).set_column_widths([2.0, 3.0])
    t.render(doc)
    docx_table = doc.tables[-1]
    assert docx_table.columns[0].width == Inches(2.0)
    assert docx_table.columns[1].width == Inches(3.0)


def test_set_column_widths_rejects_wrong_length():
    t = Table.from_list([["a", "b"], ["1", "2"]])
    with pytest.raises(ValueError):
        t.set_column_widths([1.0])


def test_autofit_sets_python_docx_flag():
    doc = Document()
    Table.from_list([["a", "b"], ["1", "2"]]).autofit().render(doc)
    assert doc.tables[-1].autofit is True


def test_apply_style_returns_self_for_chaining():
    style = _make_demo_style()
    t = Table.from_list([["a"], ["1"]])
    assert t.apply_style(style) is t
    assert t.set_column_widths([1.0]) is t
    assert t.autofit() is t


# ---------------------------------------------------------------------------
# scenario-7: ragged rows -> TableRaggedRowsError naming bad row index
# ---------------------------------------------------------------------------


def test_from_list_ragged_rows_raises_named_index():
    with pytest.raises(TableRaggedRowsError) as excinfo:
        Table.from_list([["a", "b"], ["1"]])
    msg = str(excinfo.value)
    assert "row 1" in msg
    assert "1 columns" in msg
    assert "expected 2" in msg


# ---------------------------------------------------------------------------
# scenario-8: malformed markdown -> TableMarkdownParseError
# ---------------------------------------------------------------------------


def test_from_markdown_not_a_table_raises():
    with pytest.raises(TableMarkdownParseError):
        Table.from_markdown("not a table")


def test_from_markdown_missing_separator_raises():
    # Header line has pipes, but the second non-blank line is not a "---" separator.
    with pytest.raises(TableMarkdownParseError) as excinfo:
        Table.from_markdown("| a | b |\n| 1 | 2 |")
    assert "separator" in str(excinfo.value).lower()


def test_from_markdown_separator_column_mismatch_raises():
    with pytest.raises(TableMarkdownParseError) as excinfo:
        Table.from_markdown("| a | b |\n|---|\n| 1 | 2 |")
    assert "separator" in str(excinfo.value).lower()


def test_from_markdown_data_row_column_mismatch_raises():
    with pytest.raises(TableMarkdownParseError) as excinfo:
        Table.from_markdown("| a | b |\n|---|---|\n| 1 |")
    assert "data row" in str(excinfo.value).lower()


# ---------------------------------------------------------------------------
# scenario-9: divergent dict key sets -> TableInvalidDictKeysError
# ---------------------------------------------------------------------------


def test_from_list_of_dicts_divergent_keys_raises_named_diff():
    with pytest.raises(TableInvalidDictKeysError) as excinfo:
        Table.from_list_of_dicts([{"a": 1}, {"b": 2}])
    msg = str(excinfo.value)
    # Names diverging keys: missing 'a', extra 'b' relative to row 0.
    assert "row 1" in msg
    assert "a" in msg
    assert "b" in msg


def test_from_list_of_dicts_extra_key_raises():
    with pytest.raises(TableInvalidDictKeysError):
        Table.from_list_of_dicts([{"a": 1, "b": 2}, {"a": 3, "b": 4, "c": 5}])


def test_from_list_of_dicts_missing_key_raises():
    with pytest.raises(TableInvalidDictKeysError):
        Table.from_list_of_dicts([{"a": 1, "b": 2}, {"a": 3}])


# ---------------------------------------------------------------------------
# BLOCK_RENDER_TABLE trace assertion
# ---------------------------------------------------------------------------


def test_render_emits_block_render_table_marker(caplog_at_info, marker_counter):
    doc = Document()
    Table.from_list([["a", "b"], ["1", "2"]]).render(doc)
    counts = marker_counter(caplog_at_info)
    assert counts["BLOCK_RENDER_TABLE"] == 1


def test_render_marker_payload_contains_rows_and_cols(caplog_at_info):
    doc = Document()
    Table.from_list([["a", "b", "c"], ["1", "2", "3"]]).render(doc)
    matched = [
        r.getMessage()
        for r in caplog_at_info.records
        if extract_marker(r.getMessage()) == "BLOCK_RENDER_TABLE"
    ]
    assert len(matched) == 1
    assert "rows=2" in matched[0]
    assert "cols=3" in matched[0]
    # Marker prefix is verbatim per spec.
    assert matched[0].startswith("[MP-Table][render][BLOCK_RENDER_TABLE]")


def test_render_marker_fires_before_mutation(caplog_at_info, marker_counter):
    # Two consecutive renders should produce two markers.
    doc = Document()
    Table.from_list([["a"], ["1"]]).render(doc)
    Table.from_list([["x"], ["y"]]).render(doc)
    counts = marker_counter(caplog_at_info)
    assert counts["BLOCK_RENDER_TABLE"] == 2


# ---------------------------------------------------------------------------
# Cell carrier coverage
# ---------------------------------------------------------------------------


def test_cell_is_frozen():
    from dataclasses import FrozenInstanceError

    c = Cell(value="x")
    with pytest.raises(FrozenInstanceError):
        c.value = "y"  # type: ignore[misc]


def test_cell_coerces_non_str_value_to_str():
    c = Cell(value=42)  # type: ignore[arg-type]
    assert c.value == "42"
    assert isinstance(c.value, str)


def test_cell_in_from_list_passes_through():
    custom = Cell(value="x", align="center")
    doc = Document()
    Table.from_list([["a", "b"], [custom, "y"]]).render(doc)
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    para = doc.tables[-1].rows[1].cells[0].paragraphs[0]
    assert para.text == "x"
    assert para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Negative coverage: empty inputs don't crash
# ---------------------------------------------------------------------------


def test_from_list_empty_rows_renders_placeholder():
    doc = Document()
    Table.from_list([]).render(doc)
    # Placeholder 1x1 table — render must not crash on empty input.
    assert len(doc.tables[-1].rows) == 1


def test_marker_counter_is_a_counter_instance(caplog_at_info, marker_counter):
    # Sanity check: the conftest fixture returns a Counter.
    doc = Document()
    Table.from_list([["a"], ["1"]]).render(doc)
    counts = marker_counter(caplog_at_info)
    assert isinstance(counts, Counter)
