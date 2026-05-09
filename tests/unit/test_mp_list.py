# FILE: tests/unit/test_mp_list.py
# START_MODULE_CONTRACT
#   PURPOSE: V-MP-LIST scenarios 1-7 + BLOCK_RENDER_LIST trace assertion.
#     Exercises List (bullet/numbered/checklist kinds, level indent, empty
#     items, render emits N styled paragraphs).
#   SCOPE: Reuses central conftest fixtures (caplog_at_info, marker_counter);
#     no local fixture redefinitions.
#   DEPENDS: pytest, mint_python.core.list_block, python-docx.
#   LINKS: docs/verification-plan.xml#V-MP-LIST,
#     docs/development-plan.xml#MP-LIST
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   test_scenario_1_bullet_list_renders
#   test_scenario_2_numbered_list_renders
#   test_scenario_3_checklist_prefixes_each_item
#   test_scenario_4_empty_items_emits_nothing
#   test_scenario_5_level_applies_indent
#   test_scenario_6_negative_level_raises
#   test_scenario_7_block_render_list_marker_payload
#   test_section_add_list_appends_block
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: v0.1.0 — initial test suite for MP-LIST.
# END_CHANGE_SUMMARY
from __future__ import annotations

import pytest
from docx import Document as DocxDocumentCtor
from docx.shared import Inches

from mint_python.core.list_block import (
    List,
    ListKind,
    ListLevelError,
)
from mint_python.core.section import Section

# ---------------------------------------------------------------------------
# V-MP-LIST scenario-1: bullet list renders N paragraphs with List Bullet style
# ---------------------------------------------------------------------------


def test_scenario_1_bullet_list_renders() -> None:
    doc = DocxDocumentCtor()
    List(["alpha", "beta", "gamma"]).render(doc)
    paragraphs = doc.paragraphs
    assert len(paragraphs) == 3
    for p, expected in zip(paragraphs, ["alpha", "beta", "gamma"], strict=True):
        assert p.text == expected
        assert p.style.name == "List Bullet"


# ---------------------------------------------------------------------------
# V-MP-LIST scenario-2: numbered list uses List Number style
# ---------------------------------------------------------------------------


def test_scenario_2_numbered_list_renders() -> None:
    doc = DocxDocumentCtor()
    List(["first", "second"], kind=ListKind.NUMBERED).render(doc)
    assert len(doc.paragraphs) == 2
    for p in doc.paragraphs:
        assert p.style.name == "List Number"


# ---------------------------------------------------------------------------
# V-MP-LIST scenario-3: checklist prefixes each item with ☐
# ---------------------------------------------------------------------------


def test_scenario_3_checklist_prefixes_each_item() -> None:
    doc = DocxDocumentCtor()
    List(["buy milk", "walk dog"], kind=ListKind.CHECKLIST).render(doc)
    paragraphs = doc.paragraphs
    assert len(paragraphs) == 2
    assert paragraphs[0].text == "☐ buy milk"
    assert paragraphs[1].text == "☐ walk dog"
    # Checklist uses bullet style (no native checkbox content-control in v0.1).
    assert paragraphs[0].style.name == "List Bullet"


# ---------------------------------------------------------------------------
# V-MP-LIST scenario-4: empty items emits zero paragraphs
# ---------------------------------------------------------------------------


def test_scenario_4_empty_items_emits_nothing() -> None:
    doc = DocxDocumentCtor()
    before = len(doc.paragraphs)
    List([]).render(doc)
    after = len(doc.paragraphs)
    assert after == before  # no paragraphs added


# ---------------------------------------------------------------------------
# V-MP-LIST scenario-5: level=N applies N * 0.25 in left indent
# ---------------------------------------------------------------------------


def test_scenario_5_level_applies_indent() -> None:
    doc = DocxDocumentCtor()
    List(["nested"], level=2).render(doc)
    p = doc.paragraphs[0]
    assert p.paragraph_format.left_indent == Inches(0.5)


# ---------------------------------------------------------------------------
# V-MP-LIST scenario-6: negative level raises ListLevelError
# ---------------------------------------------------------------------------


def test_scenario_6_negative_level_raises() -> None:
    with pytest.raises(ListLevelError, match=">= 0"):
        List(["x"], level=-1)


# ---------------------------------------------------------------------------
# V-MP-LIST scenario-7: BLOCK_RENDER_LIST marker payload
# ---------------------------------------------------------------------------


def test_scenario_7_block_render_list_marker_payload(
    caplog_at_info, marker_counter
) -> None:
    doc = DocxDocumentCtor()
    List(["a", "b"], kind=ListKind.NUMBERED, level=1).render(doc)
    counts = marker_counter(caplog_at_info)
    assert counts.get("BLOCK_RENDER_LIST") == 1

    msgs = [r.getMessage() for r in caplog_at_info.records]
    list_msgs = [m for m in msgs if "BLOCK_RENDER_LIST" in m]
    assert len(list_msgs) == 1
    payload = list_msgs[0]
    assert "kind=numbered" in payload
    assert "items=2" in payload
    assert "level=1" in payload


# ---------------------------------------------------------------------------
# Section integration: add_list appends List to _blocks; renders inline
# ---------------------------------------------------------------------------


def test_section_add_list_appends_block() -> None:
    sec = Section("Demo", level=1)
    returned = sec.add_list(List(["x", "y"]))
    # Fluent contract: add_list returns self.
    assert returned is sec
    assert len(sec._blocks) == 1
    assert isinstance(sec._blocks[0], List)

    doc = DocxDocumentCtor()
    sec.render(doc)
    # heading + 2 list items.
    assert any(p.text == "x" for p in doc.paragraphs)
    assert any(p.text == "y" for p in doc.paragraphs)
