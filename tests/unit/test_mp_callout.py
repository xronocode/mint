# FILE: tests/unit/test_mp_callout.py
# START_MODULE_CONTRACT
#   PURPOSE: V-MP-CALLOUT scenarios + BLOCK_RENDER_CALLOUT trace assertion.
#     Exercises Callout (info/warning/code kinds, optional title, render
#     emits a styled 1x1 table).
#   SCOPE: Reuses central conftest fixtures (caplog_at_info, marker_counter).
#   DEPENDS: pytest, mint_python.core.callout, python-docx, lxml.
#   LINKS: docs/verification-plan.xml#V-MP-CALLOUT,
#     docs/development-plan.xml#MP-CALLOUT
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   test_scenario_1_info_renders_one_table_with_shading
#   test_scenario_2_warning_uses_amber_palette
#   test_scenario_3_code_uses_monospace_font
#   test_scenario_4_explicit_title_overrides_default
#   test_scenario_5_block_render_callout_marker_payload
#   test_scenario_6_section_add_callout_appends_block
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: v0.1.0 — initial test suite for MP-CALLOUT.
# END_CHANGE_SUMMARY
from __future__ import annotations

from docx import Document as DocxDocumentCtor
from docx.oxml.ns import qn

from mint_python.core.callout import Callout, CalloutKind
from mint_python.core.section import Section


def _cell_shading_fill(cell) -> str | None:
    tc_pr = cell._tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    return None if shd is None else shd.get(qn("w:fill"))


def _cell_left_border_color(cell) -> str | None:
    tc_pr = cell._tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        return None
    left = borders.find(qn("w:left"))
    return None if left is None else left.get(qn("w:color"))


# ---------------------------------------------------------------------------
# V-MP-CALLOUT scenario-1: INFO renders a 1x1 table with light-blue shading
# ---------------------------------------------------------------------------


def test_scenario_1_info_renders_one_table_with_shading() -> None:
    doc = DocxDocumentCtor()
    Callout("Background note.", kind=CalloutKind.INFO).render(doc)
    assert len(doc.tables) == 1
    table = doc.tables[-1]
    assert len(table.rows) == 1
    assert len(table.columns) == 1
    cell = table.cell(0, 0)
    assert _cell_shading_fill(cell) == "EBF5FB"
    assert _cell_left_border_color(cell) == "2E75B6"
    # Body text appears in the cell.
    assert "Background note." in cell.text
    # Default title for info kind is "Info".
    assert "Info" in cell.text


# ---------------------------------------------------------------------------
# V-MP-CALLOUT scenario-2: WARNING uses amber palette
# ---------------------------------------------------------------------------


def test_scenario_2_warning_uses_amber_palette() -> None:
    doc = DocxDocumentCtor()
    Callout("Save first.", kind=CalloutKind.WARNING).render(doc)
    cell = doc.tables[-1].cell(0, 0)
    assert _cell_shading_fill(cell) == "FFF8E1"
    assert _cell_left_border_color(cell) == "E8A838"
    assert "Warning" in cell.text


# ---------------------------------------------------------------------------
# V-MP-CALLOUT scenario-3: CODE uses monospace + grey palette
# ---------------------------------------------------------------------------


def test_scenario_3_code_uses_monospace_font() -> None:
    doc = DocxDocumentCtor()
    Callout("def foo(): ...", kind=CalloutKind.CODE).render(doc)
    cell = doc.tables[-1].cell(0, 0)
    assert _cell_shading_fill(cell) == "F5F5F5"
    assert _cell_left_border_color(cell) == "DDDDDD"
    # The body run's font is Courier New.
    body_paragraph = cell.paragraphs[-1]
    body_run = body_paragraph.runs[-1]
    assert body_run.font.name == "Courier New"


# ---------------------------------------------------------------------------
# V-MP-CALLOUT scenario-4: explicit title overrides the kind default
# ---------------------------------------------------------------------------


def test_scenario_4_explicit_title_overrides_default() -> None:
    doc = DocxDocumentCtor()
    Callout("body", kind=CalloutKind.INFO, title="Custom Heading").render(doc)
    cell = doc.tables[-1].cell(0, 0)
    assert "Custom Heading" in cell.text
    assert "Info" not in cell.paragraphs[0].text  # default title not used


# ---------------------------------------------------------------------------
# V-MP-CALLOUT scenario-5: BLOCK_RENDER_CALLOUT marker payload
# ---------------------------------------------------------------------------


def test_scenario_5_block_render_callout_marker_payload(
    caplog_at_info, marker_counter
) -> None:
    doc = DocxDocumentCtor()
    Callout("hello", kind=CalloutKind.WARNING, title="!").render(doc)
    counts = marker_counter(caplog_at_info)
    assert counts.get("BLOCK_RENDER_CALLOUT") == 1
    msgs = [r.getMessage() for r in caplog_at_info.records]
    payload = next(m for m in msgs if "BLOCK_RENDER_CALLOUT" in m)
    assert "kind=warning" in payload
    assert "has_title=True" in payload


# ---------------------------------------------------------------------------
# V-MP-CALLOUT scenario-6: Section.add_callout appends + renders inline
# ---------------------------------------------------------------------------


def test_scenario_6_section_add_callout_appends_block() -> None:
    sec = Section("Demo", level=1)
    returned = sec.add_callout(Callout("note", kind=CalloutKind.INFO))
    assert returned is sec
    assert len(sec._blocks) == 1
    assert isinstance(sec._blocks[0], Callout)

    doc = DocxDocumentCtor()
    sec.render(doc)
    # heading paragraph + the callout's 1x1 table.
    assert len(doc.tables) == 1
