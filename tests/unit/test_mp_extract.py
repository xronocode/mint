# FILE: tests/unit/test_mp_extract.py
"""Unit tests for the pure-python port of mint.extract (MP-EXTRACT, Wave-16-1).

Covers V-MP-EXTRACT scenarios 1-7. Scenario-5 is the porting-parity oracle:
output of mint_python.extract.extract_style MUST be byte-equal to the legacy
mint.extract.extract_style on the same fixture.
"""

from __future__ import annotations

import io
import logging
import subprocess
import zipfile
from pathlib import Path

import pytest

from mint_python.extract import (
    ExtractionFailedError,
    UnsupportedFormatError,
    analyze_layouts,
    extract_style,
    parse_theme,
)
from tests._helpers.sample_docs import (
    valid_memo_docx_bytes,
    write_to_tmp,
)

_FIXTURES_DIR = Path(__file__).resolve().parents[1] / "fixtures"
_PPTX_FIXTURE = _FIXTURES_DIR / "minimal_valid.pptx"


def _docx_without_theme(tmp_path: Path) -> Path:
    """Return a .docx path whose word/theme/theme1.xml entry has been removed.

    The legacy implementation tolerates a missing theme.xml (silently empties
    the theme dict). We use this fixture to drive the parse_theme empty-theme
    branch and confirm porting parity holds.
    """
    src = valid_memo_docx_bytes()
    buf = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(src), "r") as src_zf,
        zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as dst_zf,
    ):
        for info in src_zf.infolist():
            if info.filename == "word/theme/theme1.xml":
                continue
            dst_zf.writestr(info, src_zf.read(info.filename))
    target = tmp_path / "no_theme.docx"
    target.write_bytes(buf.getvalue())
    return target


# ---------------------------------------------------------------------------
# Scenario-1: extract_style(memo_docx) — docx happy path
# ---------------------------------------------------------------------------


def test_scenario_1_extract_docx(tmp_path: Path) -> None:
    path = write_to_tmp(tmp_path, "memo.docx", valid_memo_docx_bytes())

    result = extract_style(path)

    assert result["format"] == "docx"
    assert "colors" in result
    assert "typography" in result
    assert isinstance(result["colors"], dict)
    assert isinstance(result["typography"], dict)
    # The memo built by python-docx ships the default theme1.xml which
    # contains a populated clrScheme + fontScheme.
    assert result["colors"], "theme colors must be populated"
    assert result["typography"], "theme typography must be populated"
    # Layouts list — exposed as detected_layouts when non-empty.
    assert "detected_layouts" in result
    assert isinstance(result["detected_layouts"], list)
    assert any(item["type"] == "paragraph" for item in result["detected_layouts"])


# ---------------------------------------------------------------------------
# Scenario-2: extract_style(pptx_fixture) — pptx happy path
# ---------------------------------------------------------------------------


def test_scenario_2_extract_pptx() -> None:
    assert _PPTX_FIXTURE.is_file(), "fixture tests/fixtures/minimal_valid.pptx is required"

    result = extract_style(_PPTX_FIXTURE)

    assert result["format"] == "pptx"
    assert "colors" in result
    assert "typography" in result
    # minimal_valid.pptx has a single slide → detected_layouts populated.
    assert result.get("detected_layouts"), "pptx fixture must produce at least one layout"


# ---------------------------------------------------------------------------
# Scenario-3: unsupported format
# ---------------------------------------------------------------------------


def test_scenario_3_unsupported_format(tmp_path: Path) -> None:
    txt_path = tmp_path / "not_a_doc.txt"
    txt_path.write_text("hello")

    # extract_style wraps UnsupportedFormatError in ExtractionFailedError
    # (matches legacy behavior). The underlying _detect_format helper still
    # raises UnsupportedFormatError directly when called.
    with pytest.raises(ExtractionFailedError, match="Unsupported format"):
        extract_style(txt_path)

    # Direct check on the private detector confirms the discriminating type.
    from mint_python.extract import _detect_format

    with pytest.raises(UnsupportedFormatError):
        _detect_format(Path("foo.rtf"))


# ---------------------------------------------------------------------------
# Scenario-4: docx without theme.xml → ExtractionFailedError
#
# DELTA NOTE: the legacy mint.extract.extract_style silently tolerates a
# missing theme.xml (returns an empty theme dict). The verification-plan
# scenario-4 specifies "Docx without theme.xml → ExtractionFailedError",
# which is the desired behavior expressed at planning time but contradicts
# constraint-3 (no divergence from legacy). To satisfy BOTH the porting-
# parity oracle (scenario-5) AND scenario-4 we:
#
#   (a) confirm the no-theme docx returns an empty theme without raising
#       (legacy parity), and
#   (b) confirm the closest valid ExtractionFailedError path — a .docx
#       file whose content is not a real zip — does raise.
#
# Surfaced as a verification-delta proposal in the worker report.
# ---------------------------------------------------------------------------


def test_scenario_4_missing_theme(tmp_path: Path) -> None:
    # (a) Legacy-parity: docx without theme.xml returns empty theme, no raise.
    no_theme = _docx_without_theme(tmp_path)
    result = extract_style(no_theme)
    assert result["format"] == "docx"
    assert result["colors"] == {}
    assert result["typography"] == {}
    assert "word/theme/theme1.xml" not in result["xml_sources"]

    # (b) A .docx file with non-zip bytes raises ExtractionFailedError, which
    # is the closest error path matching the verification-plan intent.
    corrupt = tmp_path / "corrupt.docx"
    corrupt.write_bytes(b"this is not a real zip archive")
    with pytest.raises(ExtractionFailedError, match="Invalid OOXML file"):
        extract_style(corrupt)

    # And a non-existent path raises ExtractionFailedError too.
    with pytest.raises(ExtractionFailedError, match="File not found"):
        extract_style(tmp_path / "missing.docx")


# ---------------------------------------------------------------------------
# Scenario-5: PORTING-PARITY — output byte-equivalent to legacy
# ---------------------------------------------------------------------------


def test_scenario_5_legacy_compatibility(tmp_path: Path) -> None:
    from mint.extract import extract_style as legacy_extract_style

    # Docx parity on the shared memo fixture.
    memo = write_to_tmp(tmp_path, "memo.docx", valid_memo_docx_bytes())
    assert extract_style(memo) == legacy_extract_style(memo)

    # Pptx parity on the committed minimal pptx fixture.
    assert extract_style(_PPTX_FIXTURE) == legacy_extract_style(_PPTX_FIXTURE)

    # Parity even on a docx with theme.xml removed (empty-theme branch).
    no_theme = _docx_without_theme(tmp_path)
    assert extract_style(no_theme) == legacy_extract_style(no_theme)


# ---------------------------------------------------------------------------
# Scenario-6: BLOCK_EXTRACT_DONE log marker
# ---------------------------------------------------------------------------


def test_scenario_6_log_marker(
    tmp_path: Path, caplog: pytest.LogCaptureFixture
) -> None:
    path = write_to_tmp(tmp_path, "memo.docx", valid_memo_docx_bytes())

    with caplog.at_level(logging.INFO, logger="mint_python.extract"):
        result = extract_style(path)

    done_records = [
        rec for rec in caplog.records if "[BLOCK_EXTRACT_DONE]" in rec.getMessage()
    ]
    assert len(done_records) == 1, "expected exactly one BLOCK_EXTRACT_DONE marker"

    msg = done_records[0].getMessage()
    assert "[MP-Extract][run][BLOCK_EXTRACT_DONE]" in msg
    assert f"format={result['format']}" in msg

    expected_theme_keys = len(result.get("colors", {})) + len(
        result.get("typography", {})
    )
    expected_layouts = len(result.get("detected_layouts", []))
    assert f"theme_keys_count={expected_theme_keys}" in msg
    assert f"layouts_count={expected_layouts}" in msg


# ---------------------------------------------------------------------------
# Scenario-7: Constraint-8 grep gate — zero `from mint.` / `import mint.` lines
# ---------------------------------------------------------------------------


def test_scenario_7_no_legacy_import() -> None:
    module_path = Path(__file__).resolve().parents[2] / "src" / "mint_python" / "extract.py"
    assert module_path.is_file(), f"module missing at {module_path}"

    # Use grep -E (extended regex) so the test is self-contained and matches
    # the worker-spec verification target.
    result = subprocess.run(
        ["grep", "-nE", r"^(from mint\.|import mint\.)", str(module_path)],
        capture_output=True,
        text=True,
        check=False,
    )
    # grep exits 1 when no matches found — that is the success path.
    assert result.returncode == 1, (
        "Constraint-8 violation: src/mint_python/extract.py imports from "
        f"src/mint/:\n{result.stdout}"
    )
    assert result.stdout == ""


# ---------------------------------------------------------------------------
# Direct coverage of helper paths (parse_theme + analyze_layouts edge cases)
# ---------------------------------------------------------------------------


def test_parse_theme_returns_empty_when_archive_has_no_theme(tmp_path: Path) -> None:
    """parse_theme tolerates a missing theme entry — exercises the KeyError branch."""
    no_theme = _docx_without_theme(tmp_path)
    with zipfile.ZipFile(no_theme) as zf:
        result = parse_theme("docx", zf)
    assert result["colors"] == {}
    assert result["typography"] == {}
    assert result["format"] == "docx"
    # theme1.xml is absent → not in xml_sources, but styles.xml is still
    # present and listed.
    assert "word/theme/theme1.xml" not in result["xml_sources"]
    assert "word/styles.xml" in result["xml_sources"]


def test_parse_theme_pptx_skips_styles_path() -> None:
    """pptx branch has no styles_path → xml_sources omits the styles entry."""
    with zipfile.ZipFile(_PPTX_FIXTURE) as zf:
        result = parse_theme("pptx", zf)
    assert result["format"] == "pptx"
    assert not any(src.endswith("styles.xml") for src in result["xml_sources"])


def test_analyze_layouts_handles_unparseable_slide(tmp_path: Path) -> None:
    """A pptx slide with malformed XML is skipped — exercises ParseError branch."""
    src = _PPTX_FIXTURE.read_bytes()
    buf = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(src), "r") as src_zf,
        zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as dst_zf,
    ):
        for info in src_zf.infolist():
            if info.filename == "ppt/slides/slide1.xml":
                dst_zf.writestr(info.filename, b"<not-xml>>>>")
            else:
                dst_zf.writestr(info, src_zf.read(info.filename))
    broken = tmp_path / "broken_slide.pptx"
    broken.write_bytes(buf.getvalue())

    with zipfile.ZipFile(broken) as zf:
        layouts = analyze_layouts("pptx", zf)
    # The single malformed slide is silently skipped.
    assert layouts == []


def test_analyze_layouts_handles_unparseable_docx_document(tmp_path: Path) -> None:
    """A docx with malformed document.xml yields no layouts — ParseError branch."""
    src = valid_memo_docx_bytes()
    buf = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(src), "r") as src_zf,
        zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as dst_zf,
    ):
        for info in src_zf.infolist():
            if info.filename == "word/document.xml":
                dst_zf.writestr(info.filename, b"<not-xml>>>>")
            else:
                dst_zf.writestr(info, src_zf.read(info.filename))
    broken = tmp_path / "broken_doc.docx"
    broken.write_bytes(buf.getvalue())

    with zipfile.ZipFile(broken) as zf:
        layouts = analyze_layouts("docx", zf)
    assert layouts == []


def test_extract_color_value_handles_sysclr_and_missing() -> None:
    """Cover the sysClr branch and the no-color branch of _extract_color_value."""
    from xml.etree import ElementTree as ET

    from mint_python.extract import _extract_color_value

    ns = "http://schemas.openxmlformats.org/drawingml/2006/main"

    sys_xml = f'<root xmlns:a="{ns}"><a:sysClr lastClr="ABCDEF"/></root>'
    sys_elem = ET.fromstring(sys_xml)
    assert _extract_color_value(sys_elem) == "#ABCDEF"

    empty_xml = f'<root xmlns:a="{ns}"><a:sysClr/></root>'
    empty_elem = ET.fromstring(empty_xml)
    assert _extract_color_value(empty_elem) is None

    bare_xml = f'<root xmlns:a="{ns}"/>'
    bare_elem = ET.fromstring(bare_xml)
    assert _extract_color_value(bare_elem) is None


def _docx_with_custom_theme(tmp_path: Path, theme_xml: bytes, *, drop_styles: bool = False) -> Path:
    """Return a .docx with a hand-rolled theme1.xml and optionally no styles.xml.

    Used to drive the clrScheme branches that the python-docx default theme
    doesn't naturally hit (missing tag / missing color value).
    """
    src = valid_memo_docx_bytes()
    buf = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(src), "r") as src_zf,
        zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as dst_zf,
    ):
        for info in src_zf.infolist():
            if info.filename == "word/theme/theme1.xml":
                dst_zf.writestr(info.filename, theme_xml)
            elif drop_styles and info.filename == "word/styles.xml":
                continue
            else:
                dst_zf.writestr(info, src_zf.read(info.filename))
    target = tmp_path / "custom_theme.docx"
    target.write_bytes(buf.getvalue())
    return target


def test_parse_theme_handles_missing_clrscheme_tag_and_no_color_value(
    tmp_path: Path,
) -> None:
    """Exercise the `elem is None: continue` and `rgb is None: continue` paths."""
    # theme with clrScheme containing only some tags + one tag with no color.
    theme_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        b'name="custom">'
        b"<a:themeElements>"
        b"<a:clrScheme name=\"custom\">"
        # dk1 has no srgbClr / sysClr → rgb is None branch.
        b"<a:dk1/>"
        # dk2 is wholly absent → `elem is None: continue` branch.
        # lt1 present with sysClr.
        b'<a:lt1><a:sysClr val="window" lastClr="FFFFFF"/></a:lt1>'
        b"</a:clrScheme>"
        b'<a:fontScheme name="custom">'
        b"<a:majorFont><a:latin typeface=\"Arial\"/></a:majorFont>"
        b"<a:minorFont><a:latin typeface=\"Arial\"/></a:minorFont>"
        b"</a:fontScheme>"
        b"<a:fmtScheme/>"
        b"</a:themeElements>"
        b"</a:theme>"
    )
    path = _docx_with_custom_theme(tmp_path, theme_xml)
    result = extract_style(path)

    # dk1 had no color (rgb None) → not in colors. dk2 absent → also not in.
    assert "dark1" not in result["colors"]
    assert "dark2" not in result["colors"]
    # lt1 with sysClr did extract.
    assert result["colors"].get("light1") == "#FFFFFF"


def test_parse_theme_missing_styles_xml_drops_entry(tmp_path: Path) -> None:
    """When styles.xml is absent, xml_sources MUST NOT list it (KeyError branch)."""
    theme_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        b'name="x"><a:themeElements/></a:theme>'
    )
    path = _docx_with_custom_theme(tmp_path, theme_xml, drop_styles=True)
    with zipfile.ZipFile(path) as zf:
        result = parse_theme("docx", zf)
    assert "word/styles.xml" not in result["xml_sources"]
    assert "word/theme/theme1.xml" in result["xml_sources"]


def test_analyze_layouts_counts_table_when_present(tmp_path: Path) -> None:
    """Exercise the `table_count > 0` branch."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("hello")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "a"
    buf = io.BytesIO()
    doc.save(buf)
    path = tmp_path / "with_table.docx"
    path.write_bytes(buf.getvalue())

    result = extract_style(path)
    types = {item["type"] for item in result.get("detected_layouts", [])}
    assert "table" in types
    assert "paragraph" in types
