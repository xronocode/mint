# FILE: tests/unit/test_page_layout.py
# VERSION: 0.1.0
"""Unit tests for MP-PAGE_LAYOUT (Margins, PageLayout, apply_to_docx_section).

Covers:
- Construction validation (margins range, columns range, spacing > 0,
  orientation literal).
- Section.with_page_layout fluent return.
- Render-time sectPr emission: orientation flag + dimension swap, margins
  in twentieths-of-a-point, w:cols num/space, header/footer text override.
- Idempotence: re-applying overwrites the previous w:cols rather than
  appending duplicates.
"""
from __future__ import annotations

import tempfile
import zipfile
from pathlib import Path

import pytest
from lxml import etree

from mint_python.sdk import Document, Margins, PageLayout, Section

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _save_and_parse(doc: Document) -> etree._Element:
    """Save doc to a temp .docx and return the parsed word/document.xml root."""
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "out.docx"
        doc.save(out)
        with zipfile.ZipFile(out) as z:
            return etree.fromstring(z.read("word/document.xml"))


# --------------------------------------------------------------------------- #
# Margins validation
# --------------------------------------------------------------------------- #


class TestMargins:
    def test_defaults_are_one_inch(self) -> None:
        m = Margins()
        assert m.top == m.bottom == m.left == m.right == 1.0

    def test_zero_or_negative_rejected(self) -> None:
        with pytest.raises(ValueError, match="top"):
            Margins(top=0)
        with pytest.raises(ValueError, match="bottom"):
            Margins(bottom=-0.1)

    def test_excessive_rejected(self) -> None:
        with pytest.raises(ValueError, match="left"):
            Margins(left=22.5)


# --------------------------------------------------------------------------- #
# PageLayout validation
# --------------------------------------------------------------------------- #


class TestPageLayoutValidation:
    def test_columns_must_be_positive(self) -> None:
        with pytest.raises(ValueError, match="columns"):
            PageLayout(columns=0)
        with pytest.raises(ValueError, match="columns"):
            PageLayout(columns=13)

    def test_column_spacing_must_be_positive(self) -> None:
        with pytest.raises(ValueError, match="spacing"):
            PageLayout(column_spacing_inches=0)

    def test_orientation_literal_enforced(self) -> None:
        with pytest.raises(ValueError, match="orientation"):
            PageLayout(orientation="diagonal")  # type: ignore[arg-type]


# --------------------------------------------------------------------------- #
# Section.with_page_layout fluent
# --------------------------------------------------------------------------- #


def test_with_page_layout_returns_self() -> None:
    sec = Section("S", level=1)
    returned = sec.with_page_layout(PageLayout(columns=2))
    assert returned is sec
    assert sec._page_layout is not None
    assert sec._page_layout.columns == 2


# --------------------------------------------------------------------------- #
# Render-time sectPr emission
# --------------------------------------------------------------------------- #


class TestRenderEmission:
    def test_landscape_swaps_dimensions(self) -> None:
        doc = Document(format="docx", title="T")
        doc.add_section(
            Section("Wide", level=1)
            .with_page_layout(PageLayout(orientation="landscape"))
            .add_paragraph("x")
        )
        root = _save_and_parse(doc)
        sect_prs = root.findall(".//w:sectPr", NS)
        # First sectPr is the default (cover/no-layout body); ours is the second.
        assert len(sect_prs) >= 2
        pg_sz = sect_prs[1].find("w:pgSz", NS)
        assert pg_sz is not None
        w, h = int(pg_sz.get(_w("w"))), int(pg_sz.get(_w("h")))
        assert w > h, "landscape page must be wider than tall"
        assert pg_sz.get(_w("orient")) == "landscape"

    def test_portrait_after_landscape_swaps_back(self) -> None:
        doc = Document(format="docx", title="T")
        doc.add_section(
            Section("Wide", level=1)
            .with_page_layout(PageLayout(orientation="landscape"))
            .add_paragraph("x")
        )
        doc.add_section(
            Section("Tall", level=1)
            .with_page_layout(PageLayout(orientation="portrait"))
            .add_paragraph("y")
        )
        root = _save_and_parse(doc)
        sect_prs = root.findall(".//w:sectPr", NS)
        # sect_prs: [default, landscape, portrait]
        portrait = sect_prs[-1].find("w:pgSz", NS)
        assert portrait is not None
        w, h = int(portrait.get(_w("w"))), int(portrait.get(_w("h")))
        assert w < h, "portrait page must be taller than wide after a landscape predecessor"

    def test_margins_emitted_in_twips(self) -> None:
        doc = Document(format="docx", title="T")
        doc.add_section(
            Section("M", level=1)
            .with_page_layout(
                PageLayout(margins=Margins(top=2, bottom=2, left=1.5, right=1.5))
            )
            .add_paragraph("x")
        )
        root = _save_and_parse(doc)
        sect_prs = root.findall(".//w:sectPr", NS)
        pg_mar = sect_prs[1].find("w:pgMar", NS)
        assert pg_mar is not None
        # 1 inch = 1440 twentieths-of-a-point.
        assert pg_mar.get(_w("top")) == "2880"
        assert pg_mar.get(_w("bottom")) == "2880"
        assert pg_mar.get(_w("left")) == "2160"
        assert pg_mar.get(_w("right")) == "2160"

    def test_columns_emitted(self) -> None:
        doc = Document(format="docx", title="T")
        doc.add_section(
            Section("C", level=1)
            .with_page_layout(PageLayout(columns=3, column_spacing_inches=0.25))
            .add_paragraph("x")
        )
        root = _save_and_parse(doc)
        sect_prs = root.findall(".//w:sectPr", NS)
        cols = sect_prs[1].find("w:cols", NS)
        assert cols is not None
        assert cols.get(_w("num")) == "3"
        # 0.25 inch = 360 twips
        assert cols.get(_w("space")) == "360"

    def test_header_footer_override(self) -> None:
        doc = Document(format="docx", title="T")
        doc.add_section(
            Section("H", level=1)
            .with_page_layout(PageLayout(header="Confidential", footer="Page 1"))
            .add_paragraph("x")
        )
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "h.docx"
            doc.save(out)
            with zipfile.ZipFile(out) as z:
                names = z.namelist()
                # python-docx writes header/footer parts under word/header*.xml,
                # word/footer*.xml — verify our override produced at least one
                # of each (the document has cover+TOC sections too, so there
                # may already be one default header/footer).
                header_parts = [n for n in names if n.startswith("word/header")]
                footer_parts = [n for n in names if n.startswith("word/footer")]
                assert header_parts, "no header part emitted"
                assert footer_parts, "no footer part emitted"
                joined_header_text = b"".join(z.read(p) for p in header_parts)
                joined_footer_text = b"".join(z.read(p) for p in footer_parts)
        assert b"Confidential" in joined_header_text
        assert b"Page 1" in joined_footer_text

    def test_continuous_break_when_page_break_before_false(self) -> None:
        doc = Document(format="docx", title="T")
        doc.add_section(
            Section("C", level=1)
            .with_page_layout(PageLayout(columns=2, page_break_before=False))
            .add_paragraph("x")
        )
        root = _save_and_parse(doc)
        sect_prs = root.findall(".//w:sectPr", NS)
        # The continuous-type marker lives at <w:type w:val="continuous"/>
        # under the sectPr the section ADDS — that's the second-to-last
        # sectPr in this document (last is the trailing default).
        type_el = sect_prs[1].find("w:type", NS)
        assert type_el is not None
        assert type_el.get(_w("val")) == "continuous"

    def test_idempotent_columns_application(self) -> None:
        """Re-applying a layout must not stack duplicate w:cols entries."""
        from mint_python.core.page_layout import apply_to_docx_section

        doc = Document(format="docx", title="T")
        doc.add_section(Section("S", level=1).add_paragraph("x"))

        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "x.docx"
            doc.save(out)
            # Re-apply twice on the saved docx's last section to confirm
            # idempotence at the helper level.
            from docx import Document as DocxDocument

            d = DocxDocument(str(out))
            apply_to_docx_section(d.sections[0], PageLayout(columns=2))
            apply_to_docx_section(d.sections[0], PageLayout(columns=4))
            sect_pr = d.sections[0]._sectPr
            cols = sect_pr.findall(_w("cols"))
            assert len(cols) == 1
            assert cols[0].get(_w("num")) == "4"
