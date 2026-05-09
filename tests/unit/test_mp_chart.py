# FILE: tests/unit/test_mp_chart.py
# START_MODULE_CONTRACT
#   PURPOSE: Module-level tests for MP-CHART covering V-MP-CHART scenarios 1-16
#     (factories, render, theming, chartjunk, error paths) plus 100% coverage
#     closures on chart.py.
#   SCOPE: Reuses central conftest fixtures (mp_clean_env autouse,
#     mpl_figure_cleanup autouse, caplog_at_info, marker_counter,
#     mpl_rcparams_snapshot opt-in) and helpers from tests.unit._mp_helpers
#     (extract_marker, assert_seaborn_not_imported,
#     assert_matplotlib_backend_is_agg). Does NOT exercise Section.add_chart
#     (Wave-8-2) or the SDK Chart re-export (Wave-8-2).
#   DEPENDS: pytest, mint_python.core.chart, mint_python.core.style, docx,
#     matplotlib (for figure assertions in scenarios 14-15).
#   LINKS: docs/verification-plan.xml#V-MP-CHART
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   test_bar_factory_smoke               - V-MP-CHART scenario-1
#   test_line_factory_smoke              - V-MP-CHART scenario-2
#   test_stacked_bar_accepts_dict_series - V-MP-CHART scenario-3
#   test_pie_factory_smoke               - V-MP-CHART scenario-4
#   test_heatmap_shape_mismatch_raises   - V-MP-CHART scenario-5
#   test_waterfall_smoke                 - V-MP-CHART scenario-6
#   test_gantt_smoke                     - V-MP-CHART scenario-7
#   test_from_matplotlib_wraps_figure    - V-MP-CHART scenario-8
#   test_from_seaborn_*                  - V-MP-CHART scenario-9 (3 cases)
#   test_from_plotly_creates_chart        - V-MP-CHART scenario-10a (mocked)
#   test_from_plotly_no_plotly_raises      - V-MP-CHART scenario-10b (mocked)
#   test_from_plotly_no_kaleido_raises     - V-MP-CHART scenario-10c (mocked)
#   test_factory_marker_count            - V-MP-CHART scenario-11
#   test_render_marker_payload           - V-MP-CHART scenario-12
#   test_render_inline_shape_emu         - V-MP-CHART scenario-13
#   test_corporate_theming_no_leak       - V-MP-CHART scenario-14
#   test_chartjunk_rc_baseline           - V-MP-CHART scenario-15
#   test_invalid_data_*                  - V-MP-CHART scenario-16 + coverage
#   test_seaborn_lazy_import             - V-MP-CHART forbidden-3
#   test_backend_is_agg                  - V-MP-CHART forbidden-2
#   coverage closures                    - savefig error, gantt empty, etc.
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: Phase-12 — replace from_plotly phase-guard test with 3
#     mock-based tests for the concrete from_plotly implementation
#     (with-mock, no-plotly, no-kaleido). Previous: Wave-8-1 initial tests.
# END_CHANGE_SUMMARY

from __future__ import annotations

import builtins
import sys
from typing import Any

import pytest
from docx import Document as DocxDocument

from mint_python.core.chart import (
    _CORPORATE_RC,
    Chart,
    ChartFigureRenderFailedError,
    ChartInvalidDataError,
    _apply_corporate_theme,
)
from mint_python.core.style import Style
from tests.unit._mp_helpers import (
    assert_matplotlib_backend_is_agg,
)

# ---------------------------------------------------------------------------
# Scenario 1: Chart.bar
# ---------------------------------------------------------------------------


def test_bar_factory_smoke(caplog_at_info: pytest.LogCaptureFixture) -> None:
    chart = Chart.bar(
        ["a", "b", "c"],
        [1.0, 2.0, 3.0],
        title="Bar title",
        caption="bar caption",
        width_inches=4.0,
    )
    assert chart.chart_type == "bar"
    assert isinstance(chart._png_bytes, bytes)
    assert len(chart._png_bytes) > 0
    assert chart.width_inches == 4.0
    assert chart.caption == "bar caption"


# ---------------------------------------------------------------------------
# Scenario 2: Chart.line
# ---------------------------------------------------------------------------


def test_line_factory_smoke(caplog_at_info: pytest.LogCaptureFixture) -> None:
    chart = Chart.line(
        ["t1", "t2", "t3"],
        [10.0, 20.0, 15.0],
        title="Line",
    )
    assert chart.chart_type == "line"
    assert len(chart._png_bytes) > 0


# ---------------------------------------------------------------------------
# Scenario 3: Chart.stacked_bar
# ---------------------------------------------------------------------------


def test_stacked_bar_accepts_dict_series(
    caplog_at_info: pytest.LogCaptureFixture,
) -> None:
    chart = Chart.stacked_bar(
        ["Q1", "Q2", "Q3"],
        {"NA": [1.0, 2.0, 3.0], "EU": [0.5, 0.6, 0.7]},
        title="Stacked",
    )
    assert chart.chart_type == "stacked_bar"
    assert len(chart._png_bytes) > 0


def test_stacked_bar_series_length_mismatch() -> None:
    with pytest.raises(ChartInvalidDataError, match="length mismatch"):
        Chart.stacked_bar(
            ["Q1", "Q2", "Q3"],
            {"bad": [1.0, 2.0]},  # 2 vs 3 labels
        )


# ---------------------------------------------------------------------------
# Scenario 4: Chart.pie
# ---------------------------------------------------------------------------


def test_pie_factory_smoke(caplog_at_info: pytest.LogCaptureFixture) -> None:
    chart = Chart.pie(
        ["A", "B", "C"],
        [10.0, 20.0, 70.0],
        title="Pie",
    )
    assert chart.chart_type == "pie"
    assert len(chart._png_bytes) > 0


def test_pie_length_mismatch_raises() -> None:
    with pytest.raises(ChartInvalidDataError, match="length mismatch"):
        Chart.pie(["A", "B"], [1.0])


# ---------------------------------------------------------------------------
# Scenario 5: Chart.heatmap shape mismatch
# ---------------------------------------------------------------------------


def test_heatmap_smoke() -> None:
    chart = Chart.heatmap(
        [[1.0, 2.0, 3.0], [4.0, 5.0, 6.0]],
        ["r1", "r2"],
        ["c1", "c2", "c3"],
        title="Heat",
    )
    assert chart.chart_type == "heatmap"
    assert len(chart._png_bytes) > 0


def test_heatmap_row_count_mismatch_raises() -> None:
    with pytest.raises(ChartInvalidDataError, match="row count mismatch"):
        Chart.heatmap(
            [[1.0, 2.0]],  # 1 row but 2 row_labels
            ["r1", "r2"],
            ["c1", "c2"],
        )


def test_heatmap_col_count_mismatch_raises() -> None:
    with pytest.raises(ChartInvalidDataError, match="column count mismatch"):
        Chart.heatmap(
            [[1.0, 2.0], [3.0, 4.0, 5.0]],  # row 1 has 3 cols, expected 2
            ["r1", "r2"],
            ["c1", "c2"],
        )


# ---------------------------------------------------------------------------
# Scenario 6: Chart.waterfall sum-of-deltas annotation
# ---------------------------------------------------------------------------


def test_waterfall_smoke_with_mixed_signs() -> None:
    chart = Chart.waterfall(
        ["Start", "Add", "Subtract", "End"],
        [100.0, 50.0, -30.0, 20.0],
        title="Waterfall",
    )
    assert chart.chart_type == "waterfall"
    assert len(chart._png_bytes) > 0


def test_waterfall_length_mismatch_raises() -> None:
    with pytest.raises(ChartInvalidDataError, match="length mismatch"):
        Chart.waterfall(["a", "b"], [1.0])


# ---------------------------------------------------------------------------
# Scenario 7: Chart.gantt
# ---------------------------------------------------------------------------


def test_gantt_smoke() -> None:
    chart = Chart.gantt(
        [("design", 0.0, 5.0), ("build", 5.0, 10.0), ("test", 12.0, 3.0)],
        title="Gantt",
    )
    assert chart.chart_type == "gantt"
    assert len(chart._png_bytes) > 0


def test_gantt_empty_tasks_raises() -> None:
    with pytest.raises(ChartInvalidDataError, match="non-empty"):
        Chart.gantt([])


def test_gantt_bad_task_shape_raises() -> None:
    with pytest.raises(ChartInvalidDataError, match="3-tuple"):
        Chart.gantt([("design", 0.0)])  # type: ignore[list-item]


# ---------------------------------------------------------------------------
# Scenario 8: Chart.from_matplotlib
# ---------------------------------------------------------------------------


def test_from_matplotlib_wraps_user_figure(
    caplog_at_info: pytest.LogCaptureFixture,
) -> None:
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(3, 2), dpi=100)
    ax.plot([1, 2, 3], [1, 4, 9])

    chart = Chart.from_matplotlib(
        fig, caption="user-built", width_inches=4.0, vector=True
    )
    assert chart.chart_type == "matplotlib"
    assert chart.caption == "user-built"
    assert chart.vector is True
    assert chart.width_inches == 4.0
    assert len(chart._png_bytes) > 0


# ---------------------------------------------------------------------------
# Scenario 9: Chart.from_seaborn — installed + missing variants
# ---------------------------------------------------------------------------


def test_from_seaborn_aliases_to_matplotlib_when_present(
    monkeypatch: pytest.MonkeyPatch, caplog_at_info: pytest.LogCaptureFixture
) -> None:
    """When seaborn is import-able, from_seaborn aliases to from_matplotlib.

    We monkeypatch a fake seaborn module so the test does not require seaborn
    as a real dependency.
    """
    import types as _types

    import matplotlib.pyplot as plt

    fake_sns = _types.ModuleType("seaborn")
    monkeypatch.setitem(sys.modules, "seaborn", fake_sns)

    fig, ax = plt.subplots(figsize=(3, 2), dpi=100)
    ax.bar(["a", "b"], [1, 2])

    chart = Chart.from_seaborn(fig, caption="sb")
    assert chart.chart_type == "seaborn"
    assert chart.caption == "sb"
    assert len(chart._png_bytes) > 0


def test_from_seaborn_raises_with_install_hint_when_missing(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """When seaborn is NOT importable, from_seaborn raises ImportError
    with a 'pip install seaborn' hint (V-MP-CHART scenario-9 negative case).
    """
    import builtins

    monkeypatch.delitem(sys.modules, "seaborn", raising=False)

    real_import = builtins.__import__

    def _fail_seaborn_import(
        name: str, globals_: Any = None, locals_: Any = None, fromlist: Any = (), level: int = 0
    ) -> Any:
        if name == "seaborn":
            raise ImportError("No module named 'seaborn'")
        return real_import(name, globals_, locals_, fromlist, level)

    monkeypatch.setattr(builtins, "__import__", _fail_seaborn_import)

    with pytest.raises(ImportError, match="pip install seaborn"):
        Chart.from_seaborn(object())


# ---------------------------------------------------------------------------
# Scenario 10: Chart.from_plotly (Phase-12 concrete, mock-based)
# ---------------------------------------------------------------------------


class TestFromPlotly:
    def test_from_plotly_creates_chart(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """from_plotly wraps plotly Figure into Chart."""
        from unittest.mock import MagicMock

        mock_fig = MagicMock()
        mock_fig.to_image.return_value = b"fake_png_bytes"
        monkeypatch.setitem(sys.modules, "plotly", MagicMock())

        chart = Chart.from_plotly(mock_fig, caption="Plot", width_inches=5.0)
        assert chart.chart_type == "plotly"
        assert chart._png_bytes == b"fake_png_bytes"
        assert chart.caption == "Plot"
        mock_fig.to_image.assert_called_once_with(format="png", scale=2)

    def test_from_plotly_no_plotly_raises(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Missing plotly → ImportError with install hint."""
        from unittest.mock import MagicMock

        mock_fig = MagicMock()
        original_import = builtins.__import__

        def _fail_plotly(name: str, *args: Any, **kwargs: Any) -> Any:
            if name == "plotly":
                raise ImportError("no plotly")
            return original_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", _fail_plotly)

        with pytest.raises(ImportError, match="pip install plotly kaleido"):
            Chart.from_plotly(mock_fig)

    def test_from_plotly_no_kaleido_raises(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Plotly installed but kaleido missing → ImportError."""
        from unittest.mock import MagicMock

        mock_fig = MagicMock()
        mock_fig.to_image.side_effect = ValueError("kaleido required")
        monkeypatch.setitem(sys.modules, "plotly", MagicMock())

        with pytest.raises(ImportError, match="kaleido"):
            Chart.from_plotly(mock_fig)


# ---------------------------------------------------------------------------
# Scenario 11: factory marker count
# ---------------------------------------------------------------------------


def test_factory_marker_count(
    caplog_at_info: pytest.LogCaptureFixture, marker_counter: Any
) -> None:
    Chart.bar(["a", "b"], [1.0, 2.0])
    Chart.line(["a", "b"], [1.0, 2.0])
    Chart.pie(["a", "b"], [1.0, 1.0])
    counts = marker_counter(caplog_at_info)
    # Each factory emits BLOCK_BUILD_CHART exactly once.
    assert counts["BLOCK_BUILD_CHART"] == 3


def test_factory_marker_payload_format(
    caplog_at_info: pytest.LogCaptureFixture,
) -> None:
    Chart.bar(["a", "b", "c"], [1.0, 2.0, 3.0])
    bar_msgs = [
        r.getMessage()
        for r in caplog_at_info.records
        if "[MP-Chart][bar][BLOCK_BUILD_CHART]" in r.getMessage()
    ]
    assert len(bar_msgs) == 1
    assert "chart_type=bar" in bar_msgs[0]
    assert "data_shape=3x1" in bar_msgs[0]


# ---------------------------------------------------------------------------
# Scenario 12: render BLOCK_RENDER_CHART marker payload
# ---------------------------------------------------------------------------


def test_render_emits_marker_with_chart_type_and_emu_width(
    caplog_at_info: pytest.LogCaptureFixture,
) -> None:
    chart = Chart.bar(["a", "b"], [1.0, 2.0], width_inches=3.0)
    caplog_at_info.clear()  # focus on render markers

    doc = DocxDocument()
    chart.render(doc)

    render_msgs = [
        r.getMessage()
        for r in caplog_at_info.records
        if "BLOCK_RENDER_CHART" in r.getMessage()
    ]
    assert len(render_msgs) == 1
    msg = render_msgs[0]
    expected_emu = round(3.0 * 914400)
    assert "chart_type=bar" in msg
    assert f"emu_width={expected_emu}" in msg


# ---------------------------------------------------------------------------
# Scenario 13: render produces inline_shape with correct EMU width
# ---------------------------------------------------------------------------


def test_render_inline_shape_emu_matches() -> None:
    chart = Chart.bar(["a", "b"], [1.0, 2.0], width_inches=3.0)
    doc = DocxDocument()
    chart.render(doc)

    assert len(doc.inline_shapes) == 1
    shape = doc.inline_shapes[0]
    assert shape.width.emu == round(3.0 * 914400)


def test_render_with_height_inches_passes_kwarg() -> None:
    chart = Chart.bar(["a", "b"], [1.0, 2.0], width_inches=3.0, height_inches=2.0)
    doc = DocxDocument()
    chart.render(doc)
    shape = doc.inline_shapes[0]
    assert shape.width.emu == round(3.0 * 914400)
    assert shape.height.emu == round(2.0 * 914400)


# ---------------------------------------------------------------------------
# Scenario 14: corporate theming — rcParams must NOT leak past the call
# ---------------------------------------------------------------------------


def test_corporate_theming_does_not_leak_rcparams(
    mpl_rcparams_snapshot: Any,
) -> None:
    """Build a Chart under a Style preset; rcParams must be unchanged after.

    The mpl_rcparams_snapshot fixture asserts no leak at teardown — V-MP-CHART
    forbidden-5. If chart.py introduced a module-level rcParams mutation or
    forgot the rc_context wrapper, this test would fail at fixture teardown.
    """
    style = Style(font="Inter", size_pt=11.0, color_hex="#0F4C81")
    chart = Chart.bar(["a", "b"], [1.0, 2.0], style=style)
    assert chart.chart_type == "bar"


def test_apply_corporate_theme_with_style_sets_font_and_color() -> None:
    style = Style(font="Inter", size_pt=11.0, color_hex="#0F4C81")
    rc = _apply_corporate_theme(style)
    assert rc["font.family"] == ["Inter", "DejaVu Sans"]
    assert rc["axes.labelcolor"] == "#0F4C81"
    assert rc["axes.titlecolor"] == "#0F4C81"
    assert rc["xtick.color"] == "#0F4C81"
    assert rc["ytick.color"] == "#0F4C81"


def test_apply_corporate_theme_without_style_returns_baseline_only() -> None:
    rc = _apply_corporate_theme(None)
    assert "font.family" not in rc
    assert rc["axes.spines.top"] is False
    assert rc["axes.spines.right"] is False


# ---------------------------------------------------------------------------
# Scenario 15: chartjunk — corporate baseline turns top/right spines off
# ---------------------------------------------------------------------------


def test_chartjunk_corporate_rc_baseline_disables_top_right_spines() -> None:
    """V-MP-CHART scenario-15 (decision-level evidence per forbidden-15).

    Asserting on the rcParams baseline that we apply via rc_context is the
    deterministic surface — Figure-internal inspection after savefig is
    fragile across matplotlib versions.
    """
    assert _CORPORATE_RC["axes.spines.top"] is False
    assert _CORPORATE_RC["axes.spines.right"] is False
    assert _CORPORATE_RC["axes.grid"] is True


# ---------------------------------------------------------------------------
# Scenario 16: ChartInvalidDataError on length mismatch
# ---------------------------------------------------------------------------


def test_bar_invalid_data_length_mismatch_message() -> None:
    with pytest.raises(ChartInvalidDataError) as excinfo:
        Chart.bar(["a", "b"], [1.0])
    msg = str(excinfo.value)
    assert "len(labels)=2" in msg
    assert "len(values)=1" in msg


def test_line_invalid_data_length_mismatch() -> None:
    with pytest.raises(ChartInvalidDataError, match="length mismatch"):
        Chart.line(["a"], [1.0, 2.0])


# ---------------------------------------------------------------------------
# Forbidden invariants: backend pin + seaborn lazy import
# ---------------------------------------------------------------------------


def test_matplotlib_backend_is_agg() -> None:
    """V-MP-CHART forbidden-2: backend MUST be 'Agg' after chart import."""
    assert_matplotlib_backend_is_agg()


def test_seaborn_not_imported_at_module_level() -> None:
    """V-MP-CHART forbidden-3: seaborn MUST be lazy-imported.

    Inspect the chart module's source directly: scan top-level statements
    for `import seaborn` / `from seaborn`. This is more reliable than
    sys.modules introspection since it cannot be polluted by sibling tests
    that monkeypatch sys.modules['seaborn'] for from_seaborn coverage.
    Source-level static guard is the canonical V-MP-CHART forbidden-3 check.
    """
    import ast
    import inspect

    import mint_python.core.chart as chart_module

    src = inspect.getsource(chart_module)
    tree = ast.parse(src)
    for node in tree.body:  # top-level only, NOT recursive
        if isinstance(node, ast.Import):
            for alias in node.names:
                assert alias.name != "seaborn", (
                    "seaborn imported at module level (V-MP-CHART forbidden-3)"
                )
        elif isinstance(node, ast.ImportFrom):
            assert node.module != "seaborn", (
                "seaborn imported at module level (V-MP-CHART forbidden-3)"
            )


# ---------------------------------------------------------------------------
# Coverage closures
# ---------------------------------------------------------------------------


def test_fig_to_png_bytes_wraps_savefig_failure() -> None:
    """ChartFigureRenderFailedError wraps the underlying matplotlib error."""

    class _BrokenFig:
        def savefig(self, *_a: Any, **_k: Any) -> None:
            raise RuntimeError("disk on fire")

    with pytest.raises(ChartFigureRenderFailedError, match="disk on fire"):
        Chart._fig_to_png_bytes(_BrokenFig())


def test_render_default_height_omits_height_kwarg() -> None:
    """When height_inches is None, render passes only width to add_picture."""
    chart = Chart.bar(["a", "b"], [1.0, 2.0], width_inches=3.0)
    doc = DocxDocument()
    chart.render(doc)
    assert chart.height_inches is None
    # Shape exists; height was set by python-docx auto-aspect.
    assert len(doc.inline_shapes) == 1


def test_bar_without_title_omits_title() -> None:
    """Cover the `if title` False branch."""
    chart = Chart.bar(["a", "b"], [1.0, 2.0])  # no title kwarg
    assert chart.chart_type == "bar"


def test_line_without_title() -> None:
    chart = Chart.line(["a", "b"], [1.0, 2.0])
    assert chart.chart_type == "line"


def test_stacked_bar_without_title_or_series() -> None:
    """Empty series dict — no validation failure, no legend, no title branch."""
    chart = Chart.stacked_bar(["a", "b"], {})
    assert chart.chart_type == "stacked_bar"


def test_pie_without_title() -> None:
    chart = Chart.pie(["A", "B"], [1.0, 1.0])
    assert chart.chart_type == "pie"


def test_heatmap_without_title() -> None:
    chart = Chart.heatmap([[1.0]], ["r"], ["c"])
    assert chart.chart_type == "heatmap"


def test_waterfall_without_title() -> None:
    chart = Chart.waterfall(["a", "b"], [1.0, -0.5])
    assert chart.chart_type == "waterfall"


def test_gantt_without_title() -> None:
    chart = Chart.gantt([("a", 0.0, 1.0)])
    assert chart.chart_type == "gantt"
