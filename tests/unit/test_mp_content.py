# FILE: tests/unit/test_mp_content.py
# START_MODULE_CONTRACT
#   PURPOSE: V-MP-CONTENT scenarios 1-7 + BLOCK_RENDER_CONTENT trace assertion.
#     Exercises Paragraph (fluent + style application), Run (data-carrier
#     immutability), and Image (from_path / from_bytes / pillow validation /
#     width-in-EMU correctness).
#   SCOPE: Reuses central conftest fixtures (mp_clean_env, caplog_at_info,
#     marker_counter); imports extract_marker from tests.unit._mp_helpers;
#     no local fixture redefinitions. Generates a 4x4 red PNG into tmp_path
#     via pillow for the image scenarios.
#   DEPENDS: pytest, mint_python.core.content, mint_python.core.style,
#     python-docx, pillow.
#   LINKS: docs/verification-plan.xml#V-MP-CONTENT,
#     docs/development-plan.xml#MP-CONTENT
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   test_scenario_1_paragraph_basic_render
#   test_scenario_2_paragraph_fluent_chain
#   test_scenario_3_paragraph_inherits_style
#   test_scenario_4_image_from_path_render
#   test_scenario_5_image_from_path_missing
#   test_scenario_6_image_from_bytes_invalid
#   test_scenario_7_image_width_emu
#   test_block_render_content_marker_payload
#   test_run_immutability_and_reuse
#   test_paragraph_render_no_filesystem_writes
#   test_per_run_style_overrides_paragraph_style
#   test_scenario_9_run_bold_override
#   test_scenario_10_run_italic_color_size_overrides
#   test_scenario_11_run_underline_no_style_fallback
#   test_scenario_12_run_explicit_false_beats_style_true
#   test_scenario_13_add_run_kwargs_mirror_run_fields
#   test_run_color_validation_rejects_bad_hex
#   test_run_font_size_pt_validation_rejects_nonpositive
# END_MODULE_MAP
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: v0.1.0 — added scenarios 9-13 covering per-run
#     formatting overrides (bold/italic/underline/color/font_size_pt) +
#     2 validation tests for Run.color hex and Run.font_size_pt > 0.
# END_CHANGE_SUMMARY
from __future__ import annotations

import dataclasses
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image as PilImage

from mint_python.core.content import (
    Image,
    ImageFileNotFoundError,
    ImageFormatUnsupportedError,
    Paragraph,
    Run,
)
from mint_python.core.style import Style, load_preset
from tests.unit._mp_helpers import extract_marker

# ---------------------------------------------------------------------------
# Fixtures local to this module (none redefine conftest fixtures)
# ---------------------------------------------------------------------------


@pytest.fixture
def red_png(tmp_path: Path) -> Path:
    """Tiny 4x4 red PNG written to tmp_path — V-MP-CONTENT scenario-4 source."""
    p = tmp_path / "red.png"
    PilImage.new("RGB", (4, 4), "red").save(p)
    return p


@pytest.fixture
def alga_styles():
    return load_preset(name="alga_corporate")


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-1: Paragraph.render appends w:p with single w:r
# ---------------------------------------------------------------------------


def test_scenario_1_paragraph_basic_render() -> None:
    doc = Document()
    Paragraph("hello").render(doc)
    assert doc.paragraphs[-1].text == "hello"
    assert len(doc.paragraphs[-1].runs) == 1
    assert doc.paragraphs[-1].runs[0].text == "hello"


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-2: Fluent chain + per-run bold rPr distinction
# ---------------------------------------------------------------------------


def test_scenario_2_paragraph_fluent_chain(alga_styles) -> None:
    doc = Document()
    para = Paragraph("a")
    returned = para.add_run("b", style=alga_styles.heading1)
    # Fluent contract: add_run returns self.
    assert returned is para

    para.render(doc)

    runs = doc.paragraphs[-1].runs
    assert len(runs) == 2
    assert runs[0].text == "a"
    assert runs[1].text == "b"
    # First run had no style → bold attr unset (falsy in python-docx).
    assert not runs[0].bold
    # Second run inherits heading1.bold = True.
    assert runs[1].bold is True


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-3: Paragraph inherits Style when no per-run style
# ---------------------------------------------------------------------------


def test_scenario_3_paragraph_inherits_style(alga_styles) -> None:
    doc = Document()
    Paragraph("body text", style=alga_styles.body).render(doc)

    docx_run = doc.paragraphs[-1].runs[0]
    assert docx_run.font.name == alga_styles.body.font
    # font.size is a python-docx Emu length; convert pt -> emu via Inches/Pt math.
    # Pt(11) -> 11 * 12700 emu per pt = 139_700.
    assert docx_run.font.size is not None
    assert docx_run.font.size.pt == pytest.approx(alga_styles.body.size_pt)
    # body.color is text color from palette: alga primary text "#1A1A1A".
    expected_hex = alga_styles.body.color_hex.lstrip("#").upper()
    assert str(docx_run.font.color.rgb) == expected_hex


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-4: Image.from_path.render appends w:drawing
# ---------------------------------------------------------------------------


def test_scenario_4_image_from_path_render(red_png: Path) -> None:
    doc = Document()
    before = len(doc.inline_shapes)
    Image.from_path(red_png).render(doc)
    after = len(doc.inline_shapes)
    assert after == before + 1
    # rId should be unique within the part — python-docx assigns auto.
    inline = doc.inline_shapes[-1]
    # Sanity: shape carries a non-empty xml fragment naming a:blip.
    inline_xml = inline._inline.xml
    assert "blip" in inline_xml
    assert "rId" in inline_xml or "r:embed" in inline_xml


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-5: Image.from_path on missing file raises
# ---------------------------------------------------------------------------


def test_scenario_5_image_from_path_missing(tmp_path: Path) -> None:
    missing = tmp_path / "does_not_exist.png"
    with pytest.raises(ImageFileNotFoundError) as exc_info:
        Image.from_path(missing)
    # Error message names the path so callers can debug.
    assert str(missing) in str(exc_info.value)


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-6: Image.from_bytes on invalid payload raises
# ---------------------------------------------------------------------------


def test_scenario_6_image_from_bytes_invalid() -> None:
    with pytest.raises(ImageFormatUnsupportedError) as exc_info:
        Image.from_bytes(b"\x00\x01bad", format="png")
    msg = str(exc_info.value)
    assert "png" in msg.lower()


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-7: width=3.0 -> 3 * 914400 EMU
# ---------------------------------------------------------------------------


def test_scenario_7_image_width_emu(red_png: Path) -> None:
    doc = Document()
    Image.from_path(red_png, width=3.0).render(doc)
    inline = doc.inline_shapes[-1]
    # python-docx Length comparison + raw EMU check both agree.
    assert inline.width == Inches(3)
    assert inline.width.emu == 3 * 914400


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-8 (extra): BLOCK_RENDER_CONTENT marker payload
# ---------------------------------------------------------------------------


def test_block_render_content_marker_payload(
    caplog_at_info, marker_counter, red_png: Path
) -> None:
    doc = Document()
    Paragraph("hello").render(doc)
    Paragraph("a").add_run("b").render(doc)
    Image.from_path(red_png).render(doc)
    bio = BytesIO()
    PilImage.new("RGB", (4, 4), "blue").save(bio, format="png")
    Image.from_bytes(bio.getvalue(), format="png").render(doc)

    counts = marker_counter(caplog_at_info)
    # 2 paragraph renders + 2 image renders -> 4 total emits.
    assert counts.get("BLOCK_RENDER_CONTENT") == 4

    # Drill into the records to verify kind=paragraph vs kind=image discrim.
    msgs = [r.getMessage() for r in caplog_at_info.records]
    content_msgs = [m for m in msgs if extract_marker(m) == "BLOCK_RENDER_CONTENT"]
    para_msgs = [m for m in content_msgs if "kind=paragraph" in m]
    image_msgs = [m for m in content_msgs if "kind=image" in m]
    assert len(para_msgs) == 2
    assert len(image_msgs) == 2

    # Paragraph payloads include runs=N.
    assert any("runs=1" in m for m in para_msgs)  # Paragraph("hello")
    assert any("runs=2" in m for m in para_msgs)  # Paragraph("a").add_run("b")

    # Image payloads carry source=path|bytes.
    assert any("source=path" in m for m in image_msgs)
    assert any("source=bytes" in m for m in image_msgs)


# ---------------------------------------------------------------------------
# Forbidden-2 proof: Run is immutable; sharing across paragraphs is harmless
# ---------------------------------------------------------------------------


def test_run_immutability_and_reuse() -> None:
    r = Run("shared", style=None)
    # Run is a frozen dataclass: mutation raises FrozenInstanceError.
    with pytest.raises(dataclasses.FrozenInstanceError):
        r.text = "mutated"  # type: ignore[misc]

    # Sharing the same Run instance across two Paragraphs is safe — both
    # render correctly with no cross-corruption.
    p1 = Paragraph([r])
    p2 = Paragraph([r, Run(" + tail")])
    doc = Document()
    p1.render(doc)
    p2.render(doc)
    assert doc.paragraphs[-2].text == "shared"
    assert doc.paragraphs[-1].text == "shared + tail"


# ---------------------------------------------------------------------------
# Forbidden-1 proof: Paragraph.render does not write to the filesystem
# ---------------------------------------------------------------------------


def test_paragraph_render_no_filesystem_writes(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Sentinel check: render() must mutate parent_doc only — no fs writes.

    We monkeypatch ``builtins.open`` to refuse any *write* mode while
    Paragraph.render runs. Any disk write attempt during render -> the
    test fails. (Read-mode opens by python-docx internals, if any, are
    allowed; we only block write-mode side effects.)
    """
    import builtins

    real_open = builtins.open

    def guarded_open(path, mode="r", *args, **kwargs):  # type: ignore[no-untyped-def]
        if any(flag in mode for flag in ("w", "a", "x", "+")):
            raise AssertionError(
                f"forbidden-1 violated: render() opened {path!r} in mode {mode!r}"
            )
        return real_open(path, mode, *args, **kwargs)

    monkeypatch.setattr(builtins, "open", guarded_open)

    doc = Document()
    Paragraph("safe").render(doc)
    assert doc.paragraphs[-1].text == "safe"


# ---------------------------------------------------------------------------
# Style-override sanity: per-run style wins over paragraph style
# ---------------------------------------------------------------------------


def test_per_run_style_overrides_paragraph_style(alga_styles) -> None:
    doc = Document()
    # Paragraph style = body (not bold); explicit run = heading1 (bold).
    bold_style: Style = alga_styles.heading1
    Paragraph("a", style=alga_styles.body).add_run(
        "B", style=bold_style
    ).render(doc)
    runs = doc.paragraphs[-1].runs
    # Run 0 inherits paragraph (body -> not bold).
    assert runs[0].bold in (False, None)
    # Run 1 has explicit heading1 -> bold.
    assert runs[1].bold is True


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-9: Run.bold override beats inherited Style.bold=False
# ---------------------------------------------------------------------------


def test_scenario_9_run_bold_override(alga_styles) -> None:
    doc = Document()
    # body Style has bold=False. add_run(bold=True) must produce bold run
    # without needing a derived Style.
    (
        Paragraph("plain ", style=alga_styles.body)
        .add_run("BOLD", bold=True)
        .add_run(" tail")
        .render(doc)
    )
    runs = doc.paragraphs[-1].runs
    assert runs[0].bold in (False, None)  # inherits body
    assert runs[1].bold is True            # override
    assert runs[2].bold in (False, None)  # inherits body


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-10: italic + color + font_size overrides on one Run
# ---------------------------------------------------------------------------


def test_scenario_10_run_italic_color_size_overrides(alga_styles) -> None:
    doc = Document()
    Paragraph("base ", style=alga_styles.body).add_run(
        "fancy",
        italic=True,
        color="#FF0000",
        font_size_pt=18.0,
    ).render(doc)
    runs = doc.paragraphs[-1].runs

    # Run 0 inherits body unchanged.
    assert runs[0].italic in (False, None)
    expected_body_hex = alga_styles.body.color_hex.lstrip("#").upper()
    assert str(runs[0].font.color.rgb) == expected_body_hex
    assert runs[0].font.size.pt == pytest.approx(alga_styles.body.size_pt)

    # Run 1 overrides italic, color, size.
    assert runs[1].italic is True
    assert str(runs[1].font.color.rgb) == "FF0000"
    assert runs[1].font.size.pt == pytest.approx(18.0)


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-11: underline override has no Style fallback
# ---------------------------------------------------------------------------


def test_scenario_11_run_underline_no_style_fallback(alga_styles) -> None:
    doc = Document()
    Paragraph("plain ", style=alga_styles.body).add_run(
        "under", underline=True
    ).render(doc)
    runs = doc.paragraphs[-1].runs
    # Run 0 has no underline (Style has no underline → None passthrough).
    assert runs[0].underline in (False, None)
    # Run 1 explicitly underlines.
    assert runs[1].underline is True


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-12: explicit Run.bold=False beats Style.bold=True
# ---------------------------------------------------------------------------


def test_scenario_12_run_explicit_false_beats_style_true(alga_styles) -> None:
    doc = Document()
    # heading1 has bold=True; per-run bold=False must win.
    Paragraph(style=alga_styles.heading1).add_run(
        "not bold here", bold=False
    ).render(doc)
    docx_run = doc.paragraphs[-1].runs[0]
    assert docx_run.bold is False


# ---------------------------------------------------------------------------
# V-MP-CONTENT scenario-13: add_run kwargs reach Run via construction
# ---------------------------------------------------------------------------


def test_scenario_13_add_run_kwargs_mirror_run_fields() -> None:
    para = Paragraph().add_run(
        "x",
        bold=True,
        italic=True,
        underline=True,
        color="#0F4C81",
        font_size_pt=14.0,
    )
    # add_run returns the Paragraph; inspect the appended Run via repr.
    last = para._runs[-1]
    assert last.text == "x"
    assert last.bold is True
    assert last.italic is True
    assert last.underline is True
    assert last.color == "#0F4C81"
    assert last.font_size_pt == 14.0


def test_run_color_validation_rejects_bad_hex() -> None:
    with pytest.raises(ValueError, match="#RRGGBB"):
        Run("x", color="red")
    with pytest.raises(ValueError, match="hex digits"):
        Run("x", color="#GGHHII")


def test_run_font_size_pt_validation_rejects_nonpositive() -> None:
    with pytest.raises(ValueError, match="> 0"):
        Run("x", font_size_pt=0)
    with pytest.raises(ValueError, match="> 0"):
        Run("x", font_size_pt=-1.5)
